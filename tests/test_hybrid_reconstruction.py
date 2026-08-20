from __future__ import annotations

import io
import re
import urllib.request
import zipfile
from pathlib import Path
from xml.etree import ElementTree

import pytest
from docx import Document as WordDocument
from PIL import Image

from docreconstruct.evaluation.hybrid_validation import validate_hybrid
from docreconstruct.exceptions import UnsupportedInputError
from docreconstruct.reconstruction import (
    finalize_hybrid_reconstruction,
    prepare_markdown_layout_sources,
    prepare_markdown_pdf_sources,
    reconstruct_hybrid,
)
from docreconstruct.reconstruction.asset_matching import (
    AssetMatch,
    _download_remote,
    _SafeAssetRedirectHandler,
    _write_asset_cache,
    match_markdown_assets,
    resolve_markdown_asset,
)
from docreconstruct.reconstruction.hybrid_docx import (
    _balanced_editable_column_streams,
    _dialogue_story_boundary,
    _duplicate_figure_annotation_ids,
    _new_paragraph,
    _render_image_table_pair,
    _source_column_ink_capacities,
    _source_figure_bytes,
    _wrap_column_blocks,
    render_hybrid_docx,
)
from docreconstruct.reconstruction.hybrid_planner import (
    HybridBlockPlacement,
    HybridLayoutPlan,
    HybridPagePlan,
    apply_page_vertical_fit_budget,
    build_page_vertical_fit_budget,
    equation_layout_units,
)
from docreconstruct.reconstruction.markdown_content import (
    MarkdownBlock,
    MarkdownBlockKind,
    MarkdownContent,
    parse_markdown_content,
)
from docreconstruct.reconstruction.markdown_inline import parse_markdown_inline
from docreconstruct.reconstruction.math_omml import build_omml
from docreconstruct.reconstruction.scan_layout import (
    PixelBox,
    ScanDocumentLayout,
    ScanPageLayout,
    ScanRegion,
    ScanRegionKind,
    analyze_scan_page,
    analyze_scan_source,
)
from docreconstruct.reconstruction.table_matching import match_markdown_tables


def test_tall_single_row_masthead_uses_source_scale_as_editable_text(tmp_path: Path) -> None:
    page = ScanPageLayout(
        number=1,
        width=600,
        height=800,
        pdf_width=600,
        pdf_height=800,
        content_bbox=PixelBox(x0=0, y0=0, x1=600, y1=800),
        line_pitch=14,
        image=Image.new("RGB", (600, 800), "white"),
        metadata={"source_kind": "image"},
    )
    layout = ScanDocumentLayout(source=str(tmp_path / "masthead.png"), pages=[page])
    masthead = PixelBox(x0=30, y0=35, x1=530, y1=165)
    placement = HybridBlockPlacement(
        block_id="md-1",
        block_index=0,
        page_number=1,
        source_bbox=masthead,
        source_rows=[PixelBox(x0=30, y0=top, x1=530, y1=top + 10) for top in range(40, 160, 12)],
        geometry_source="json_consensus",
    )
    document = WordDocument()

    paragraph = _new_paragraph(
        document,
        "ĐÀN-BÀ",
        size=9.0,
        line_height=14.0,
        kind=MarkdownBlockKind.HEADING,
        available_width_points=600.0,
        placement=placement,
        layout=layout,
    )

    assert paragraph.runs[0].font.size is not None
    assert paragraph.runs[0].font.size.pt >= 90.0


def test_hybrid_sources_are_fingerprinted_without_fetching_remote_assets(tmp_path: Path) -> None:
    markdown = tmp_path / "content.md"
    markdown.write_text(
        "Question text\n\n![figure](https://example.invalid/mutable.png)\n",
        encoding="utf-8",
    )
    pdf = tmp_path / "layout.pdf"
    pdf.write_bytes(b"%PDF-1.7\nlocal-layout-evidence")

    manifest = prepare_markdown_pdf_sources(markdown, pdf)

    assert manifest.content.path == str(markdown.resolve())
    assert manifest.layout.path == str(pdf.resolve())
    assert manifest.content_policy == "verbatim_markdown"
    assert manifest.layout_policy == "pdf_geometry_and_original_figures_only"
    assert manifest.external_references == ["https://example.invalid/mutable.png"]
    assert len(manifest.content.sha256) == 64
    assert len(manifest.layout.sha256) == 64

    output = tmp_path / "result.docx"
    output.write_bytes(b"editable-local-result")
    result = finalize_hybrid_reconstruction(manifest, output)
    assert result.output.path == str(output.resolve())
    assert result.output.size == len(b"editable-local-result")


def test_hybrid_manifest_fingerprints_repeatable_saved_ocr_evidence(tmp_path: Path) -> None:
    markdown = tmp_path / "content.md"
    markdown.write_text("Exact Markdown wording", encoding="utf-8")
    layout = tmp_path / "layout.png"
    Image.new("RGB", (100, 160), "white").save(layout)
    paddle = tmp_path / "paddle.json"
    paddle.write_text('{"res":{"rec_texts":["Exact Markdown wording"]}}', encoding="utf-8")
    mineru = tmp_path / "mineru.jsonl"
    mineru.write_text('{"text":"Exact Markdown wording"}\n', encoding="utf-8")

    manifest = prepare_markdown_layout_sources(
        markdown,
        layout,
        evidence=[paddle, mineru],
    )

    assert [item.path for item in manifest.evidence] == [
        str(paddle.resolve()),
        str(mineru.resolve()),
    ]
    assert [item.media_type for item in manifest.evidence] == [
        "application/json",
        "application/x-ndjson",
    ]
    assert manifest.evidence_policy == "geometry_style_confidence_only_never_text_authority"


def test_hybrid_manifest_rejects_duplicate_or_non_json_evidence(tmp_path: Path) -> None:
    markdown = tmp_path / "content.md"
    markdown.write_text("content", encoding="utf-8")
    layout = tmp_path / "layout.png"
    Image.new("RGB", (100, 160), "white").save(layout)
    evidence = tmp_path / "saved.json"
    evidence.write_text("{}", encoding="utf-8")

    with pytest.raises(ValueError, match="more than once"):
        prepare_markdown_layout_sources(
            markdown,
            layout,
            evidence=[evidence, evidence],
        )

    unsupported = tmp_path / "saved.txt"
    unsupported.write_text("not structured evidence", encoding="utf-8")
    with pytest.raises(UnsupportedInputError, match="JSON"):
        prepare_markdown_layout_sources(markdown, layout, evidence=[unsupported])


def test_hybrid_source_roles_reject_swapped_formats(tmp_path: Path) -> None:
    markdown = tmp_path / "content.txt"
    markdown.write_text("content", encoding="utf-8")
    pdf = tmp_path / "layout.pdf"
    pdf.write_bytes(b"%PDF-1.7\nlayout")

    with pytest.raises(UnsupportedInputError, match="Markdown"):
        prepare_markdown_pdf_sources(markdown, pdf)


def test_raster_layout_authority_is_validated_and_fingerprinted(tmp_path: Path) -> None:
    markdown = tmp_path / "content.md"
    markdown.write_text("Editable content", encoding="utf-8")
    layout = tmp_path / "layout.jpg"
    Image.new("RGB", (120, 180), "white").save(layout)

    manifest = prepare_markdown_layout_sources(markdown, layout)

    assert manifest.layout.media_type == "image/jpeg"
    assert manifest.layout_policy == "image_geometry_and_original_figures_only"


def test_tinted_paper_becomes_native_word_background_not_a_page_scan(tmp_path: Path) -> None:
    pytest.importorskip("numpy")
    from PIL import ImageDraw

    markdown = tmp_path / "tinted.md"
    markdown.write_text("Editable content", encoding="utf-8")
    content = parse_markdown_content(markdown)
    image = Image.new("RGB", (600, 820), (205, 225, 240))
    draw = ImageDraw.Draw(image)
    for top in range(80, 220, 28):
        draw.rectangle((70, top, 530, top + 8), fill="black")
    page = analyze_scan_page(
        image,
        number=1,
        pdf_width=595,
        pdf_height=813,
        metadata={"source_kind": "image", "rectified": False},
    )
    color = page.metadata["paper_color"]
    channels = tuple(int(color[index : index + 2], 16) for index in (0, 2, 4))
    assert channels == pytest.approx((205, 225, 240), abs=2)
    scan = ScanDocumentLayout(source=str(tmp_path / "layout.png"), pages=[page])
    placement = HybridBlockPlacement(
        block_id=content.blocks[0].id,
        block_index=0,
        page_number=1,
    )
    plan = HybridLayoutPlan(
        content_source=content.source,
        layout_source=scan.source,
        pages=[
            HybridPagePlan(
                number=1,
                pdf_width=page.pdf_width,
                pdf_height=page.pdf_height,
                raster_width=page.width,
                raster_height=page.height,
                content_bbox=page.content_bbox,
                line_pitch=page.line_pitch,
                placements=[placement],
            )
        ],
    )

    payload = render_hybrid_docx(content, scan, plan, [])
    with zipfile.ZipFile(io.BytesIO(payload)) as package:
        root = ElementTree.fromstring(package.read("word/document.xml"))
        assert not any(name.startswith("word/media/") for name in package.namelist())
    word = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    background = root.find(f"{word}background")
    assert background is not None
    assert background.get(f"{word}color") == color


def test_saved_figure_with_wrong_aspect_recrops_authoritative_source_box(
    tmp_path: Path,
) -> None:
    page = ScanPageLayout(
        number=1,
        width=600,
        height=800,
        pdf_width=595,
        pdf_height=793,
        content_bbox=PixelBox(x0=30, y0=20, x1=570, y1=780),
        line_pitch=30,
        image=Image.new("RGB", (600, 800), "white"),
        metadata={"source_kind": "pdf"},
    )
    layout = ScanDocumentLayout(source=str(tmp_path / "layout.pdf"), pages=[page])
    wrong = io.BytesIO()
    Image.new("RGB", (40, 120), "black").save(wrong, format="PNG")
    bbox = PixelBox(x0=100, y0=200, x1=340, y1=300)

    selected = _source_figure_bytes(
        {"figure": wrong.getvalue()},
        "figure",
        layout,
        1,
        bbox,
    )

    with Image.open(io.BytesIO(selected)) as image:
        assert image.size == (bbox.width, bbox.height)


def test_coarse_provider_paragraph_bbox_does_not_inflate_native_line_height(
    tmp_path: Path,
) -> None:
    page = ScanPageLayout(
        number=1,
        width=1200,
        height=1697,
        pdf_width=595,
        pdf_height=842,
        content_bbox=PixelBox(x0=80, y0=60, x1=1120, y1=1620),
        line_pitch=40,
        image=Image.new("RGB", (1200, 1697), "white"),
        metadata={"source_kind": "pdf"},
    )
    layout = ScanDocumentLayout(source=str(tmp_path / "layout.pdf"), pages=[page])
    coarse = PixelBox(x0=100, y0=200, x1=1050, y1=420)
    placement = HybridBlockPlacement(
        block_id="paragraph",
        block_index=0,
        page_number=1,
        source_bbox=coarse,
        source_rows=[coarse],
        source_gap_before=0,
        geometry_source="json_consensus",
    )
    document = WordDocument()

    paragraph = _new_paragraph(
        document,
        "Editable prose represented by a multi-line provider envelope.",
        size=12,
        line_height=18,
        kind=MarkdownBlockKind.PARAGRAPH,
        available_width_points=500,
        placement=placement,
        layout=layout,
    )

    assert paragraph.paragraph_format.line_spacing is not None
    assert paragraph.paragraph_format.line_spacing.pt == pytest.approx(18)


def test_coarse_tall_inline_math_bbox_retains_source_line_height_floor(
    tmp_path: Path,
) -> None:
    page = ScanPageLayout(
        number=1,
        width=1200,
        height=1697,
        pdf_width=595,
        pdf_height=842,
        content_bbox=PixelBox(x0=80, y0=60, x1=1120, y1=1620),
        line_pitch=40,
        image=Image.new("RGB", (1200, 1697), "white"),
        metadata={"source_kind": "pdf"},
    )
    layout = ScanDocumentLayout(source=str(tmp_path / "layout.pdf"), pages=[page])
    coarse = PixelBox(x0=100, y0=200, x1=1050, y1=420)
    placement = HybridBlockPlacement(
        block_id="paragraph",
        block_index=0,
        page_number=1,
        source_bbox=coarse,
        source_rows=[coarse],
        source_gap_before=0,
        geometry_source="json_consensus",
    )
    document = WordDocument()

    paragraph = _new_paragraph(
        document,
        r"From $\int_0^x e^{t^2}\,dt=\frac{x^3}{3}$ we obtain the limit.",
        size=12,
        line_height=18,
        kind=MarkdownBlockKind.PARAGRAPH,
        available_width_points=500,
        placement=placement,
        layout=layout,
    )

    assert paragraph.paragraph_format.line_spacing is not None
    assert paragraph.paragraph_format.line_spacing.pt == pytest.approx(18 * 1.35)


def test_markdown_parser_retains_solution_groups_lists_and_display_math(tmp_path: Path) -> None:
    markdown = tmp_path / "math.md"
    markdown.write_text(
        "# Title\n\nA-1 Exact prose.\n\n$$ \\frac{x^{2}}{y} $$\n\n"
        "- First condition\n\nA–2 Next proof.\n",
        encoding="utf-8",
    )

    content = parse_markdown_content(markdown)

    assert [block.kind for block in content.blocks] == [
        MarkdownBlockKind.HEADING,
        MarkdownBlockKind.PARAGRAPH,
        MarkdownBlockKind.EQUATION,
        MarkdownBlockKind.LIST_ITEM,
        MarkdownBlockKind.PARAGRAPH,
    ]
    assert content.blocks[1].starts_group
    assert content.blocks[2].group_id == content.blocks[1].group_id
    assert content.blocks[4].starts_group
    assert content.blocks[4].group_id != content.blocks[1].group_id


def test_consecutive_list_items_stay_separate_blocks(tmp_path: Path) -> None:
    """Each marker starts its own block instead of one run-on paragraph.

    Every list line used to fall through to the paragraph buffer, which joined
    them with spaces, so a four-item list reached the renderer as one justified
    Word paragraph classified by its first marker alone.
    """

    markdown = tmp_path / "lists.md"
    markdown.write_text(
        "- First item\n"
        "- Second item\n"
        "  wrapped onto another line\n"
        "  - Nested item\n"
        "- Third item\n"
        "\n"
        "1. Numbered one\n"
        "2. Numbered two\n",
        encoding="utf-8",
    )

    blocks = parse_markdown_content(markdown).blocks

    assert [block.text for block in blocks] == [
        "- First item",
        "- Second item wrapped onto another line",
        "- Nested item",
        "- Third item",
        "1. Numbered one",
        "2. Numbered two",
    ]
    assert all(block.kind is MarkdownBlockKind.LIST_ITEM for block in blocks[:4])
    # Consecutive numbered items each open their own group; merging them left
    # only the first number visible to `_group_label`.
    assert blocks[4].starts_group
    assert blocks[5].starts_group
    assert blocks[4].group_id != blocks[5].group_id


def test_prose_around_a_list_is_not_absorbed_into_it(tmp_path: Path) -> None:
    markdown = tmp_path / "mixed.md"
    markdown.write_text(
        "Intro prose.\n- Only item\nBack to prose.\n\n3.14159 stays prose.\n",
        encoding="utf-8",
    )

    blocks = parse_markdown_content(markdown).blocks

    assert [(block.kind, block.text) for block in blocks] == [
        (MarkdownBlockKind.PARAGRAPH, "Intro prose."),
        (MarkdownBlockKind.LIST_ITEM, "- Only item"),
        (MarkdownBlockKind.PARAGRAPH, "Back to prose."),
        (MarkdownBlockKind.PARAGRAPH, "3.14159 stays prose."),
    ]


def test_latex_math_becomes_editable_office_math() -> None:
    equation = build_omml(r"\int_{0}^{\infty}\frac{x^{2}}{\sqrt{y}}")
    payload = ElementTree.tostring(equation, encoding="unicode")
    tags = {element.tag.rsplit("}", 1)[-1] for element in equation.iter()}

    assert "oMath" in payload
    assert "f" in tags
    assert "rad" in tags
    assert "∞" in payload


def test_cjk_numbered_method_labels_become_section_headings(tmp_path: Path) -> None:
    markdown = tmp_path / "cjk-methods.md"
    markdown.write_text(
        "方法二\n\n由泰勒公式得 $ e^x=1+x $\n\n方法三\n\n从而有 $ x=1 $\n",
        encoding="utf-8",
    )

    content = parse_markdown_content(markdown)

    assert [block.kind for block in content.blocks] == [
        MarkdownBlockKind.HEADING,
        MarkdownBlockKind.PARAGRAPH,
        MarkdownBlockKind.HEADING,
        MarkdownBlockKind.PARAGRAPH,
    ]
    assert content.blocks[0].metadata["role"] == "section_heading"
    assert content.blocks[2].metadata["role"] == "section_heading"


def test_cjk_font_profile_and_multirow_equation_spacing_are_native(tmp_path: Path) -> None:
    markdown = tmp_path / "cjk-math.md"
    markdown.write_text(
        "方法二\n\n$$ \\begin{aligned}&=\\frac{x}{2}\\\\&=1\\end{aligned} $$\n\n"
        "由泰勒公式得 $e^x=1$\n",
        encoding="utf-8",
    )
    content = parse_markdown_content(markdown)
    content_box = PixelBox(x0=35, y0=25, x1=565, y1=815)
    scan_page = ScanPageLayout(
        number=1,
        width=600,
        height=840,
        pdf_width=595,
        pdf_height=842,
        content_bbox=content_box,
        line_pitch=22,
        image=Image.new("RGB", (600, 840), "white"),
    )
    scan = ScanDocumentLayout(source=str(tmp_path / "layout.png"), pages=[scan_page])
    placements = [
        HybridBlockPlacement(
            block_id=block.id,
            block_index=block.index,
            page_number=1,
        )
        for block in content.blocks
    ]
    plan = HybridLayoutPlan(
        content_source=content.source,
        layout_source=scan.source,
        pages=[
            HybridPagePlan(
                number=1,
                pdf_width=595,
                pdf_height=842,
                raster_width=600,
                raster_height=840,
                content_bbox=content_box,
                line_pitch=22,
                placements=placements,
            )
        ],
    )

    payload = render_hybrid_docx(content, scan, plan, [])
    with zipfile.ZipFile(io.BytesIO(payload)) as package:
        root = ElementTree.fromstring(package.read("word/document.xml"))
        assert not any(name.startswith("word/media/") for name in package.namelist())

    word = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    math = "{http://schemas.openxmlformats.org/officeDocument/2006/math}"
    heading_run = next(
        run
        for run in root.iter(f"{word}r")
        if "方法二" in "".join(node.text or "" for node in run.iter(f"{word}t"))
    )
    fonts = heading_run.find(f"{word}rPr/{word}rFonts")
    language = heading_run.find(f"{word}rPr/{word}lang")
    assert fonts is not None
    assert fonts.get(f"{word}ascii") == "Times New Roman"
    assert fonts.get(f"{word}eastAsia") == "SimHei"
    assert language is not None
    assert language.get(f"{word}eastAsia") == "zh-CN"
    assert language.get(f"{word}val") != "vi-VN"
    body_run = next(
        run
        for run in root.iter(f"{word}r")
        if "由泰勒公式得" in "".join(node.text or "" for node in run.iter(f"{word}t"))
    )
    body_fonts = body_run.find(f"{word}rPr/{word}rFonts")
    assert body_fonts is not None
    assert body_fonts.get(f"{word}eastAsia") == "SimSun"

    equation_paragraph = next(
        paragraph
        for paragraph in root.iter(f"{word}p")
        if paragraph.find(f".//{math}oMath") is not None
    )
    spacing = equation_paragraph.find(f"{word}pPr/{word}spacing")
    assert spacing is not None
    assert spacing.get(f"{word}lineRule") == "atLeast"
    assert equation_layout_units(r"\begin{aligned}x\\y\\z\end{aligned}") > (
        equation_layout_units("x") * 2
    )
    math_run_sizes = {
        size.get(f"{word}val")
        for run in root.iter(f"{math}r")
        for size in run.findall(f"{word}rPr/{word}sz")
    }
    math_control_sizes = {
        size.get(f"{word}val")
        for control in root.iter(f"{math}ctrlPr")
        for size in control.findall(f"{word}rPr/{word}sz")
    }
    # The renderer uses one base half-point size for inline and display math;
    # Word remains responsible for the natural scaling of scripts/operators.
    assert len(math_run_sizes) == 1
    assert math_control_sizes == math_run_sizes
    mark_size = equation_paragraph.find(f"{word}pPr/{word}rPr/{word}sz")
    mark_font = equation_paragraph.find(f"{word}pPr/{word}rPr/{word}rFonts")
    assert mark_size is not None
    assert mark_size.get(f"{word}val") in math_run_sizes
    assert mark_font is not None
    assert mark_font.get(f"{word}ascii") == "Cambria Math"


def test_source_geometry_drives_eighteen_visual_slots_without_clipping(
    tmp_path: Path,
) -> None:
    """Retain a dense mixed math/text page without one giant equation line box."""

    equation_rows = (
        r"\begin{aligned}&=a\\&=b\\&=c\\&=d\\&=e\end{aligned}",
        "x=1",
        r"\begin{aligned}a&=b\\&=c\end{aligned}",
        "y=1",
        r"\begin{aligned}a&=b\\&=c\\&=d\\&=e\end{aligned}",
    )
    specifications = [
        (MarkdownBlockKind.EQUATION, equation_rows[0], None),
        (MarkdownBlockKind.HEADING, "方法二", "section_heading"),
        (MarkdownBlockKind.EQUATION, equation_rows[1], None),
        (MarkdownBlockKind.PARAGRAPH, "由 $x=1$", None),
        (MarkdownBlockKind.EQUATION, equation_rows[2], None),
        (MarkdownBlockKind.EQUATION, equation_rows[3], None),
        (MarkdownBlockKind.HEADING, "方法三", "section_heading"),
        (MarkdownBlockKind.PARAGRAPH, "由泰勒公式得 $e^x=1+x$", None),
        (MarkdownBlockKind.PARAGRAPH, "从而 $x=1$，于是有", None),
        (MarkdownBlockKind.EQUATION, equation_rows[4], None),
    ]
    blocks = [
        MarkdownBlock(
            id=f"md-{index + 1}",
            index=index,
            kind=kind,
            text=text,
            metadata={"role": role} if role else {},
        )
        for index, (kind, text, role) in enumerate(specifications)
    ]
    content = MarkdownContent(source=str(tmp_path / "content.md"), blocks=blocks)
    content_box = PixelBox(x0=86, y0=27, x1=958, y1=1347)
    source_rows = [
        [
            PixelBox(x0=355, y0=32, x1=853, y1=100),
            PixelBox(x0=356, y0=132, x1=853, y1=193),
            PixelBox(x0=356, y0=228, x1=704, y1=289),
            PixelBox(x0=356, y0=313, x1=747, y1=381),
            PixelBox(x0=355, y0=401, x1=722, y1=448),
        ],
        [PixelBox(x0=90, y0=465, x1=157, y1=484)],
        [PixelBox(x0=91, y0=499, x1=698, y1=576)],
        [PixelBox(x0=92, y0=603, x1=370, y1=672)],
        [
            PixelBox(x0=113, y0=691, x1=742, y1=741),
            PixelBox(x0=320, y0=760, x1=829, y1=807),
        ],
        [PixelBox(x0=89, y0=828, x1=518, y1=897)],
        [PixelBox(x0=89, y0=921, x1=156, y1=940)],
        [PixelBox(x0=91, y0=955, x1=439, y1=978)],
        [PixelBox(x0=89, y0=999, x1=477, y1=1041)],
        [
            PixelBox(x0=89, y0=1066, x1=958, y1=1145),
            PixelBox(x0=353, y0=1163, x1=713, y1=1210),
            PixelBox(x0=353, y0=1230, x1=855, y1=1278),
            PixelBox(x0=353, y0=1298, x1=857, y1=1343),
        ],
    ]
    assert sum(len(rows) for rows in source_rows) == 18
    source_boxes = [
        PixelBox(
            x0=min(row.x0 for row in rows),
            y0=min(row.y0 for row in rows),
            x1=max(row.x1 for row in rows),
            y1=max(row.y1 for row in rows),
        )
        for rows in source_rows
    ]
    previous_bottom = content_box.y0
    placements = []
    for block, box, rows in zip(blocks, source_boxes, source_rows, strict=True):
        placements.append(
            HybridBlockPlacement(
                block_id=block.id,
                block_index=block.index,
                page_number=1,
                source_bbox=box,
                source_rows=rows,
                source_gap_before=max(0, box.y0 - previous_bottom),
                match_score=1.0,
            )
        )
        previous_bottom = box.y1
    page = ScanPageLayout(
        number=1,
        width=980,
        height=1400,
        pdf_width=595.28,
        pdf_height=841.89,
        content_bbox=content_box,
        line_pitch=22.5,
        image=Image.new("RGB", (980, 1400), "white"),
        metadata={"source_kind": "image", "column_count": 1},
    )
    scan = ScanDocumentLayout(source=str(tmp_path / "layout.jpg"), pages=[page])
    plan = HybridLayoutPlan(
        content_source=content.source,
        layout_source=scan.source,
        pages=[
            HybridPagePlan(
                number=1,
                pdf_width=page.pdf_width,
                pdf_height=page.pdf_height,
                raster_width=page.width,
                raster_height=page.height,
                content_bbox=content_box,
                line_pitch=page.line_pitch,
                placements=placements,
            )
        ],
    )

    vertical_scale = page.pdf_height / page.height
    printable_height = page.content_bbox.height * vertical_scale
    body_size = max(8.6, min(12.0, page.line_pitch * vertical_scale * 0.76))
    vertical_budget = build_page_vertical_fit_budget(
        page,
        placements,
        printable_height_points=printable_height,
        font_size_points=body_size,
    )
    fitted_placements = apply_page_vertical_fit_budget(page, placements, vertical_budget)
    assert vertical_budget.fits
    assert vertical_budget.block_gap_scale < 1
    assert vertical_budget.row_gap_scale == 1
    assert fitted_placements[0].source_gap_before == placements[0].source_gap_before
    assert fitted_placements[1].source_gap_before < placements[1].source_gap_before

    payload = render_hybrid_docx(content, scan, plan, [])
    with zipfile.ZipFile(io.BytesIO(payload)) as package:
        root = ElementTree.fromstring(package.read("word/document.xml"))

    word = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    math = "{http://schemas.openxmlformats.org/officeDocument/2006/math}"
    body = root.find(f"{word}body")
    assert body is not None
    paragraphs = [child for child in body if child.tag == f"{word}p"]
    assert len(paragraphs) == len(blocks)
    display_paragraphs = [
        paragraph for paragraph in paragraphs if paragraph.find(f".//{math}oMathPara") is not None
    ]
    assert len(display_paragraphs) == 5
    assert all(
        paragraph.find(f"{word}pPr/{word}keepLines") is not None for paragraph in display_paragraphs
    )
    # Fitting is achieved by source whitespace allocation, never by silently
    # shrinking the document-wide math font.
    display_math_sizes = {
        size.get(f"{word}val")
        for paragraph in display_paragraphs
        for run in paragraph.iter(f"{math}r")
        for size in run.findall(f"{word}rPr/{word}sz")
    }
    assert display_math_sizes == {str(round(body_size * 2))}

    horizontal_scale = page.pdf_width / page.width
    vertical_scale = page.pdf_height / page.height
    first_spacing = display_paragraphs[0].find(f"{word}pPr/{word}spacing")
    first_indent = display_paragraphs[0].find(f"{word}pPr/{word}ind")
    assert first_spacing is not None
    assert first_indent is not None
    assert first_spacing.get(f"{word}lineRule") == "atLeast"
    assert int(first_spacing.get(f"{word}before", "0")) == pytest.approx(
        round((32 - content_box.y0) * vertical_scale * 20),
        abs=1,
    )
    assert int(first_indent.get(f"{word}left", "0")) == pytest.approx(
        round((355 - content_box.x0) * horizontal_scale * 20),
        abs=1,
    )
    assert int(first_indent.get(f"{word}right", "0")) == pytest.approx(
        round((content_box.x1 - 853) * horizontal_scale * 20),
        abs=1,
    )
    # The line floor follows the tallest source row, not the 250-point bbox.
    assert int(first_spacing.get(f"{word}line", "0")) >= round(68 * vertical_scale * 20) - 1
    assert int(first_spacing.get(f"{word}line", "0")) < round(
        source_boxes[0].height * vertical_scale * 20 / 2
    )
    second_spacing = paragraphs[1].find(f"{word}pPr/{word}spacing")
    assert second_spacing is not None
    assert int(second_spacing.get(f"{word}before", "0")) == pytest.approx(
        round(int(fitted_placements[1].source_gap_before or 0) * vertical_scale * 20),
        abs=1,
    )
    assert int(second_spacing.get(f"{word}before", "0")) < round(
        int(placements[1].source_gap_before or 0) * vertical_scale * 20
    )
    first_math_properties = display_paragraphs[0].find(f".//{math}oMathParaPr/{math}jc")
    assert first_math_properties is not None
    assert first_math_properties.get(f"{math}val") == "left"
    row_spacing_rule = display_paragraphs[0].find(f".//{math}eqArrPr/{math}rSpRule")
    row_spacing = display_paragraphs[0].find(f".//{math}eqArrPr/{math}rSp")
    assert row_spacing_rule is not None
    assert row_spacing is not None
    assert row_spacing_rule.get(f"{math}val") == "3"
    assert row_spacing.get(f"{math}val") == "17"

    final_indent = display_paragraphs[-1].find(f"{word}pPr/{word}ind")
    assert final_indent is not None
    assert int(final_indent.get(f"{word}left", "0")) == pytest.approx(
        round((89 - content_box.x0) * horizontal_scale * 20),
        abs=1,
    )
    assert int(final_indent.get(f"{word}right", "0")) == 0

    # Exact leading smaller than a native font can clip glyphs.  Every run of
    # editable prose in this fixture retains at least its encoded font height.
    for paragraph in paragraphs:
        spacing = paragraph.find(f"{word}pPr/{word}spacing")
        if spacing is None:
            continue
        line = int(spacing.get(f"{word}line", "0"))
        sizes = [
            int(size.get(f"{word}val", "0")) for size in paragraph.findall(f".//{word}rPr/{word}sz")
        ]
        if sizes:
            assert line >= max(sizes) * 10


def test_four_medium_math_options_render_as_native_two_by_two_grid(tmp_path: Path) -> None:
    markdown = tmp_path / "options.md"
    markdown.write_text(
        "Câu 1. Chọn phương án đúng.\n\n"
        "A. $\\frac{x-1}{2}=\\frac{y+3}{-1}=\\frac{z-5}{3}$.\n\n"
        "B. $\\frac{x-1}{2}=\\frac{y+3}{-1}=\\frac{z+5}{3}$.\n\n"
        "C. $\\frac{x-1}{2}=\\frac{y-3}{-1}=\\frac{z+5}{3}$.\n\n"
        "D. $\\frac{x-1}{2}=\\frac{y-3}{-1}=\\frac{z-5}{3}$.\n",
        encoding="utf-8",
    )
    content = parse_markdown_content(markdown)
    content_box = PixelBox(x0=40, y0=30, x1=560, y1=790)
    scan = ScanDocumentLayout(
        source=str(tmp_path / "layout.png"),
        pages=[
            ScanPageLayout(
                number=1,
                width=600,
                height=820,
                pdf_width=595,
                pdf_height=813,
                content_bbox=content_box,
                line_pitch=24,
                image=Image.new("RGB", (600, 820), "white"),
            )
        ],
    )
    placements = [
        HybridBlockPlacement(block_id=block.id, block_index=block.index, page_number=1)
        for block in content.blocks
    ]
    plan = HybridLayoutPlan(
        content_source=content.source,
        layout_source=scan.source,
        pages=[
            HybridPagePlan(
                number=1,
                pdf_width=595,
                pdf_height=813,
                raster_width=600,
                raster_height=820,
                content_bbox=content_box,
                line_pitch=24,
                placements=placements,
            )
        ],
    )

    payload = render_hybrid_docx(content, scan, plan, [])
    with zipfile.ZipFile(io.BytesIO(payload)) as package:
        root = ElementTree.fromstring(package.read("word/document.xml"))
    word = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    table = next(root.iter(f"{word}tbl"))
    rows = table.findall(f"{word}tr")

    assert len(rows) == 2
    assert all(len(row.findall(f"{word}tc")) == 2 for row in rows)
    assert all(
        (height := row.find(f"{word}trPr/{word}trHeight")) is not None
        and height.get(f"{word}hRule") == "atLeast"
        for row in rows
    )
    math = "{http://schemas.openxmlformats.org/officeDocument/2006/math}"
    math_paragraphs = [
        paragraph
        for paragraph in table.iter(f"{word}p")
        if paragraph.find(f".//{math}oMath") is not None
    ]
    assert len(math_paragraphs) == 4
    assert all(
        (spacing := paragraph.find(f"{word}pPr/{word}spacing")) is not None
        and spacing.get(f"{word}lineRule") == "atLeast"
        for paragraph in math_paragraphs
    )


def test_source_rows_choose_two_by_two_nary_options_without_clipping(tmp_path: Path) -> None:
    markdown = tmp_path / "integral-options.md"
    markdown.write_text(
        "Câu 6. Chọn đẳng thức đúng.\n\n"
        "A. $\\int 7f(x)dx=7+\\int f(x)dx$.\n\n"
        "B. $\\int 7f(x)dx=7\\int f(x)dx$.\n\n"
        "C. $\\int 7f(x)dx=7-\\int f(x)dx$.\n\n"
        "D. $\\int 7f(x)dx=-7\\int f(x)dx$.\n",
        encoding="utf-8",
    )
    content = parse_markdown_content(markdown)
    content_box = PixelBox(x0=45, y0=25, x1=555, y1=795)
    option_rows = [
        PixelBox(x0=70, y0=180, x1=530, y1=211),
        PixelBox(x0=70, y0=220, x1=530, y1=253),
    ]
    page = ScanPageLayout(
        number=1,
        width=600,
        height=820,
        pdf_width=595,
        pdf_height=813,
        content_bbox=content_box,
        line_pitch=28,
        image=Image.new("RGB", (600, 820), "white"),
        metadata={
            "source_kind": "image",
            "render_content_bbox": content_box.model_dump(),
        },
    )
    scan = ScanDocumentLayout(source=str(tmp_path / "layout.png"), pages=[page])
    placements = []
    for block in content.blocks:
        if block.kind is MarkdownBlockKind.OPTION:
            placements.append(
                HybridBlockPlacement(
                    block_id=block.id,
                    block_index=block.index,
                    page_number=1,
                    source_bbox=PixelBox(x0=70, y0=180, x1=530, y1=253),
                    source_rows=option_rows,
                    source_gap_before=4,
                )
            )
        else:
            placements.append(
                HybridBlockPlacement(
                    block_id=block.id,
                    block_index=block.index,
                    page_number=1,
                    source_bbox=PixelBox(x0=70, y0=140, x1=530, y1=168),
                    source_rows=[PixelBox(x0=70, y0=140, x1=530, y1=168)],
                    source_gap_before=0,
                )
            )
    plan = HybridLayoutPlan(
        content_source=content.source,
        layout_source=scan.source,
        pages=[
            HybridPagePlan(
                number=1,
                pdf_width=page.pdf_width,
                pdf_height=page.pdf_height,
                raster_width=page.width,
                raster_height=page.height,
                content_bbox=content_box,
                line_pitch=page.line_pitch,
                placements=placements,
            )
        ],
    )

    payload = render_hybrid_docx(content, scan, plan, [])
    with zipfile.ZipFile(io.BytesIO(payload)) as package:
        root = ElementTree.fromstring(package.read("word/document.xml"))
    word = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    table = next(root.iter(f"{word}tbl"))
    rows = table.findall(f"{word}tr")

    assert len(rows) == 2
    assert all(len(row.findall(f"{word}tc")) == 2 for row in rows)
    for row in rows:
        height = row.find(f"{word}trPr/{word}trHeight")
        assert height is not None
        assert height.get(f"{word}hRule") == "atLeast"
        assert int(height.get(f"{word}val", "0")) >= 300


def test_missing_source_rows_still_limit_nary_options_to_two_columns(tmp_path: Path) -> None:
    markdown = tmp_path / "integral-options-no-geometry.md"
    markdown.write_text(
        "Câu 6. Chọn đẳng thức đúng.\n\n"
        "A. $\\int 7f(x)dx=7+\\int f(x)dx$.\n\n"
        "B. $\\int 7f(x)dx=7\\int f(x)dx$.\n\n"
        "C. $\\int 7f(x)dx=7-\\int f(x)dx$.\n\n"
        "D. $\\int 7f(x)dx=-7\\int f(x)dx$.\n",
        encoding="utf-8",
    )
    content = parse_markdown_content(markdown)
    content_box = PixelBox(x0=45, y0=25, x1=555, y1=795)
    page = ScanPageLayout(
        number=1,
        width=600,
        height=820,
        pdf_width=595,
        pdf_height=813,
        content_bbox=content_box,
        line_pitch=28,
        image=Image.new("RGB", (600, 820), "white"),
        metadata={"source_kind": "image", "render_content_bbox": content_box.model_dump()},
    )
    scan = ScanDocumentLayout(source=str(tmp_path / "layout.png"), pages=[page])
    placements = [
        HybridBlockPlacement(
            block_id=block.id,
            block_index=block.index,
            page_number=1,
        )
        for block in content.blocks
    ]
    plan = HybridLayoutPlan(
        content_source=content.source,
        layout_source=scan.source,
        pages=[
            HybridPagePlan(
                number=1,
                pdf_width=page.pdf_width,
                pdf_height=page.pdf_height,
                raster_width=page.width,
                raster_height=page.height,
                content_bbox=content_box,
                line_pitch=page.line_pitch,
                placements=placements,
            )
        ],
    )

    payload = render_hybrid_docx(content, scan, plan, [])
    with zipfile.ZipFile(io.BytesIO(payload)) as package:
        root = ElementTree.fromstring(package.read("word/document.xml"))
    word = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    table = next(root.iter(f"{word}tbl"))
    rows = table.findall(f"{word}tr")

    assert len(rows) == 2
    assert all(len(row.findall(f"{word}tc")) == 2 for row in rows)


def test_native_table_uses_source_height_as_editable_row_floor(tmp_path: Path) -> None:
    markdown = tmp_path / "table.md"
    markdown.write_text(
        "<table>\n"
        "<tr><th>Mức lương</th><th>[5;6)</th><th>[6;7)</th></tr>\n"
        "<tr><td>Tần số</td><td>4</td><td>5</td></tr>\n"
        "</table>\n",
        encoding="utf-8",
    )
    content = parse_markdown_content(markdown)
    assert len(content.blocks) == 1
    block = content.blocks[0]
    page = ScanPageLayout(
        number=1,
        width=600,
        height=820,
        pdf_width=595,
        pdf_height=813,
        content_bbox=PixelBox(x0=45, y0=25, x1=555, y1=795),
        line_pitch=24,
        image=Image.new("RGB", (600, 820), "white"),
        metadata={"source_kind": "image"},
    )
    scan = ScanDocumentLayout(source=str(tmp_path / "layout.png"), pages=[page])
    placement = HybridBlockPlacement(
        block_id=block.id,
        block_index=block.index,
        page_number=1,
        source_bbox=PixelBox(x0=70, y0=180, x1=530, y1=250),
        source_rows=[
            PixelBox(x0=70, y0=180, x1=530, y1=213),
            PixelBox(x0=70, y0=217, x1=530, y1=250),
        ],
        source_gap_before=0,
    )
    plan = HybridLayoutPlan(
        content_source=content.source,
        layout_source=scan.source,
        pages=[
            HybridPagePlan(
                number=1,
                pdf_width=page.pdf_width,
                pdf_height=page.pdf_height,
                raster_width=page.width,
                raster_height=page.height,
                content_bbox=page.content_bbox,
                line_pitch=page.line_pitch,
                placements=[placement],
            )
        ],
    )

    payload = render_hybrid_docx(content, scan, plan, [])
    with zipfile.ZipFile(io.BytesIO(payload)) as package:
        root = ElementTree.fromstring(package.read("word/document.xml"))
    word = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    rows = next(root.iter(f"{word}tbl")).findall(f"{word}tr")

    assert len(rows) == 2
    for row in rows:
        height = row.find(f"{word}trPr/{word}trHeight")
        assert height is not None
        assert height.get(f"{word}hRule") == "atLeast"
        assert int(height.get(f"{word}val", "0")) >= 250


def test_flat_raster_scan_detects_two_columns_without_ocr(tmp_path: Path) -> None:
    pytest.importorskip("numpy")
    layout = tmp_path / "two-columns.png"
    image = Image.new("RGB", (600, 840), "white")
    from PIL import ImageDraw

    draw = ImageDraw.Draw(image)
    draw.rectangle((120, 70, 480, 82), fill="black")
    draw.rectangle((230, 100, 370, 108), fill="black")
    for top in range(180, 740, 24):
        draw.rectangle((55, top, 270, top + 7), fill="black")
        draw.rectangle((330, top, 545, top + 7), fill="black")
    image.save(layout)

    scan = analyze_scan_source(layout)

    assert len(scan.pages) == 1
    assert scan.pages[0].metadata["source_kind"] == "image"
    assert scan.pages[0].metadata["column_count"] == 2
    assert len(scan.pages[0].metadata["column_boxes"]) == 2


def test_geometry_driven_three_column_page_renders_native_no_group_flow(
    tmp_path: Path,
) -> None:
    markdown = tmp_path / "newspaper.md"
    long_first_column = "Column one opening paragraph. " + " ".join(
        f"geometry-owned-word-{index}" for index in range(90)
    )
    markdown.write_text(
        "![Masthead](masthead.png)\n\n"
        "# BROADSHEET HEADLINE\n\n"
        f"{long_first_column}\n\n"
        "Column one continuation.\n\n"
        "Column two opening paragraph.\n\n"
        "Column two continuation.\n\n"
        "Column three opening paragraph.\n\n"
        "Column three continuation.\n\n"
        "Full-width continuation below the columns.\n",
        encoding="utf-8",
    )
    content = parse_markdown_content(markdown)
    assert not any(block.starts_group for block in content.blocks)

    page = ScanPageLayout(
        number=1,
        width=620,
        height=840,
        pdf_width=595,
        pdf_height=842,
        content_bbox=PixelBox(x0=30, y0=20, x1=590, y1=810),
        line_pitch=24,
        image=Image.new("RGB", (620, 840), "white"),
        metadata={
            "source_kind": "image",
            "column_count": 3,
            "column_boxes": [
                [40, 220, 210, 700],
                [225, 220, 395, 700],
                [410, 220, 580, 700],
            ],
            "column_content_bottoms": [320, 320, 320],
            "render_content_bbox": {"x0": 40, "y0": 20, "x1": 580, "y1": 810},
        },
    )
    boxes = [
        PixelBox(x0=25, y0=0, x1=595, y1=150),
        PixelBox(x0=50, y0=170, x1=400, y1=198),
        PixelBox(x0=45, y0=230, x1=205, y1=244),
        PixelBox(x0=45, y0=275, x1=205, y1=289),
        PixelBox(x0=230, y0=230, x1=390, y1=244),
        PixelBox(x0=230, y0=275, x1=390, y1=289),
        PixelBox(x0=415, y0=230, x1=575, y1=244),
        PixelBox(x0=415, y0=275, x1=575, y1=289),
        PixelBox(x0=50, y0=350, x1=570, y1=364),
    ]
    placements = [
        HybridBlockPlacement(
            block_id=block.id,
            block_index=block.index,
            page_number=1,
            source_bbox=boxes[index],
            source_rows=[] if block.kind is MarkdownBlockKind.IMAGE else [boxes[index]],
            source_gap_before=0,
            match_score=1.0,
            geometry_source="json_consensus",
            evidence_providers=("provider",),
        )
        for index, block in enumerate(content.blocks)
    ]
    scan = ScanDocumentLayout(source=str(tmp_path / "newspaper.png"), pages=[page])
    plan = HybridLayoutPlan(
        content_source=content.source,
        layout_source=scan.source,
        pages=[
            HybridPagePlan(
                number=1,
                pdf_width=page.pdf_width,
                pdf_height=page.pdf_height,
                raster_width=page.width,
                raster_height=page.height,
                content_bbox=page.content_bbox,
                line_pitch=page.line_pitch,
                placements=placements,
            )
        ],
    )

    payload = render_hybrid_docx(content, scan, plan, [])
    with zipfile.ZipFile(io.BytesIO(payload)) as package:
        root = ElementTree.fromstring(package.read("word/document.xml"))
        media = [
            package.read(name) for name in package.namelist() if name.startswith("word/media/")
        ]
    word = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    body = root.find(f"{word}body")
    assert body is not None
    body_children = list(body)
    column_table = next(
        table
        for table in root.iter(f"{word}tbl")
        if (caption := table.find(f"{word}tblPr/{word}tblCaption")) is not None
        and caption.get(f"{word}val") == "docreconstruct:body-columns-3"
    )
    page_margins = root.find(f".//{word}sectPr/{word}pgMar")
    assert page_margins is not None
    assert page_margins.get(f"{word}top") == "0"
    cells = column_table.findall(f"{word}tr/{word}tc")
    assert len(cells) == 5
    assert column_table.find(f"{word}tr/{word}trPr/{word}cantSplit") is None
    cell_text = [" ".join(node.text or "" for node in cell.iter(f"{word}t")) for cell in cells]
    assert "Column one opening" in cell_text[0]
    assert not cell_text[1]
    assert "Column two opening" in cell_text[2]
    assert not cell_text[3]
    assert "Column three opening" in cell_text[4]
    table_index = body_children.index(column_table)
    assert any(child.find(f".//{word}drawing") is not None for child in body_children[:table_index])
    assert "BROADSHEET HEADLINE" in " ".join(
        node.text or "" for child in body_children[:table_index] for node in child.iter(f"{word}t")
    )
    assert "Full-width continuation below" in " ".join(
        node.text or ""
        for child in body_children[table_index + 1 :]
        for node in child.iter(f"{word}t")
    )
    assert len(media) == 1
    with Image.open(io.BytesIO(media[0])) as embedded:
        assert embedded.size == (570, 150)


def test_long_editable_paragraph_can_flow_across_native_column_boundaries() -> None:
    text = " ".join(f"word{index}" for index in range(150))
    block = MarkdownBlock(
        id="md-1",
        index=0,
        kind=MarkdownBlockKind.PARAGRAPH,
        text=text,
    )

    streams = _balanced_editable_column_streams(
        [block],
        widths=[2.0, 2.0, 2.0],
        target_heights=[1.0, 1.0, 1.0],
    )

    assert len(streams) == 3
    assert all(stream for stream in streams)
    assert " ".join(part.text for stream in streams for part in stream) == text
    assert all(stream[0].id == block.id for stream in streams)


def test_short_byline_before_sustained_dialogue_starts_a_new_story_flow() -> None:
    blocks = [
        MarkdownBlock(
            id="md-1",
            index=0,
            kind=MarkdownBlockKind.PARAGRAPH,
            text="Previous article conclusion.",
        ),
        MarkdownBlock(
            id="md-2",
            index=1,
            kind=MarkdownBlockKind.PARAGRAPH,
            text="A. Byline",
        ),
        *[
            MarkdownBlock(
                id=f"md-{index + 3}",
                index=index + 2,
                kind=MarkdownBlockKind.PARAGRAPH,
                text=f"— Dialogue turn {index + 1}.",
            )
            for index in range(4)
        ],
    ]

    assert _dialogue_story_boundary(blocks) == 1


def test_source_column_foreground_mass_guides_relative_story_capacity() -> None:
    from PIL import ImageDraw

    image = Image.new("RGB", (240, 180), "white")
    draw = ImageDraw.Draw(image)
    for top in range(30, 150, 12):
        draw.rectangle((15, top, 95, top + 5), fill="black")
    for top in range(30, 105, 18):
        draw.rectangle((135, top, 205, top + 4), fill="black")
    page = ScanPageLayout(
        number=1,
        width=240,
        height=180,
        pdf_width=240,
        pdf_height=180,
        content_bbox=PixelBox(x0=10, y0=20, x1=210, y1=160),
        line_pitch=12,
        image=image,
        metadata={"paper_color": "FFFFFF"},
    )
    boxes = [
        PixelBox(x0=10, y0=20, x1=105, y1=160),
        PixelBox(x0=125, y0=20, x1=215, y1=160),
    ]

    capacities = _source_column_ink_capacities(
        page,
        boxes,
        body_top=20,
        bottoms=[160, 160],
    )

    assert len(capacities) == 2
    assert capacities[0] > capacities[1] * 1.5


def test_flat_raster_scan_detects_three_columns_below_full_width_banner(
    tmp_path: Path,
) -> None:
    pytest.importorskip("numpy")
    layout = tmp_path / "three-columns.png"
    image = Image.new("RGB", (640, 900), "white")
    from PIL import ImageDraw

    draw = ImageDraw.Draw(image)
    draw.rectangle((30, 20, 610, 135), fill="black")
    draw.rectangle((55, 160, 585, 174), fill="black")
    for top in range(260, 800, 22):
        draw.rectangle((42, top, 214, top + 7), fill="black")
        draw.rectangle((234, top, 406, top + 7), fill="black")
        draw.rectangle((426, top, 598, top + 7), fill="black")
    image.save(layout)

    page = analyze_scan_source(layout).pages[0]

    assert page.metadata["column_count"] == 3
    assert len(page.metadata["column_boxes"]) == 3
    assert page.metadata["column_boxes"][0][1] >= 240
    assert len(page.metadata["column_gutters"]) == 2


def test_fragmented_full_width_rows_do_not_create_false_columns(tmp_path: Path) -> None:
    pytest.importorskip("numpy")
    from PIL import ImageDraw

    layout = tmp_path / "fragmented-full-width-lines.png"
    image = Image.new("RGB", (640, 900), "white")
    draw = ImageDraw.Draw(image)
    draw.rectangle((40, 45, 600, 58), fill="black")
    groups = ((42, 150), (190, 300), (340, 450), (490, 598))
    for top in range(180, 800, 54):
        for left, right in groups:
            draw.rectangle((left, top, right, top + 7), fill="black")
        draw.rectangle((42, top + 18, 598, top + 25), fill="black")
        draw.rectangle((42, top + 36, 598, top + 43), fill="black")
    image.save(layout)

    page = analyze_scan_source(layout).pages[0]

    assert page.metadata["column_count"] == 1
    assert "column_gutters" not in page.metadata


def test_flat_raster_scan_detects_four_persistent_columns(tmp_path: Path) -> None:
    pytest.importorskip("numpy")
    from PIL import ImageDraw

    layout = tmp_path / "four-columns.png"
    image = Image.new("RGB", (640, 900), "white")
    draw = ImageDraw.Draw(image)
    draw.rectangle((30, 20, 610, 135), fill="black")
    draw.rectangle((55, 160, 585, 174), fill="black")
    for top in range(260, 800, 22):
        draw.rectangle((42, top, 170, top + 7), fill="black")
        draw.rectangle((186, top, 314, top + 7), fill="black")
        draw.rectangle((330, top, 458, top + 7), fill="black")
        draw.rectangle((474, top, 602, top + 7), fill="black")
    image.save(layout)

    page = analyze_scan_source(layout).pages[0]

    assert page.metadata["column_count"] == 4
    assert len(page.metadata["column_boxes"]) == 4
    assert len(page.metadata["column_gutters"]) == 3


def test_photographed_skewed_grid_is_detected_as_a_ruled_table(tmp_path: Path) -> None:
    pytest.importorskip("numpy")
    from PIL import ImageDraw

    layout = tmp_path / "skewed-table.png"
    image = Image.new("RGB", (800, 1132), "white")
    draw = ImageDraw.Draw(image)
    for top in range(70, 300, 34):
        for left in range(70, 690, 70):
            draw.rectangle((left, top, left + 45, top + 8), fill="black")
    left, right = 75, 725
    tops = (360, 430, 500)
    for top in tops:
        draw.line((left, top, right, top + 15), fill="black", width=3)
    for column in range(8):
        x = left + round((right - left) * column / 7)
        offset = round(15 * column / 7)
        draw.line((x, tops[0] + offset, x, tops[-1] + offset), fill="black", width=3)
    for top in range(560, 990, 34):
        for left_text in range(70, 690, 70):
            draw.rectangle((left_text, top, left_text + 45, top + 8), fill="black")
    image.save(layout)

    page = analyze_scan_source(layout).pages[0]
    tables = [region for region in page.regions if region.kind is ScanRegionKind.TABLE]

    assert tables
    assert max(int(region.metadata["horizontal_rules"]) for region in tables) >= 3
    assert max(int(region.metadata["vertical_rules"]) for region in tables) >= 6


def test_markdown_table_rejects_image_chart_and_uses_compatible_grid(
    tmp_path: Path,
) -> None:
    markdown = tmp_path / "table.md"
    markdown.write_text(
        "<table><tr>"
        + "".join(f"<td>H{index}</td>" for index in range(7))
        + "</tr><tr>"
        + "".join(f"<td>{index}</td>" for index in range(7))
        + "</tr></table>\n",
        encoding="utf-8",
    )
    content = parse_markdown_content(markdown)
    chart = PixelBox(x0=100, y0=500, x1=420, y1=570)
    grid = PixelBox(x0=60, y0=180, x1=540, y1=270)
    page = ScanPageLayout(
        number=1,
        width=600,
        height=800,
        pdf_width=595,
        pdf_height=794,
        content_bbox=PixelBox(x0=40, y0=30, x1=560, y1=760),
        line_pitch=24,
        regions=[
            ScanRegion(
                kind=ScanRegionKind.TABLE,
                bbox=chart,
                confidence=0.8,
                metadata={"horizontal_rules": 2, "vertical_rules": 1},
            ),
            ScanRegion(
                kind=ScanRegionKind.TABLE,
                bbox=grid,
                confidence=0.9,
                metadata={"horizontal_rules": 3, "vertical_rules": 8},
            ),
        ],
        image=Image.new("RGB", (600, 800), "white"),
    )
    layout = ScanDocumentLayout(source=str(tmp_path / "layout.png"), pages=[page])
    image_match = AssetMatch(
        block_id="figure",
        source="https://assets.invalid/chart.png",
        page_number=1,
        bbox=chart,
        score=0.8,
        resolved=False,
    )

    matches = match_markdown_tables(content, layout, [image_match])

    assert len(matches) == 1
    assert matches[0].bbox == grid


def test_markdown_parser_does_not_split_lowercase_markers_inside_prose(tmp_path: Path) -> None:
    markdown = tmp_path / "content.md"
    markdown.write_text(
        "Câu 1: Trình tự gồm a) chuẩn bị mẫu, b) quan sát và c) ghi kết quả.\n",
        encoding="utf-8",
    )

    content = parse_markdown_content(markdown)

    assert len(content.blocks) == 1
    assert content.blocks[0].kind is MarkdownBlockKind.PARAGRAPH
    assert "a) chuẩn bị mẫu" in content.blocks[0].text


def test_exam_html_wrappers_and_combined_section_question_are_structured(
    tmp_path: Path,
) -> None:
    markdown = tmp_path / "exam.md"
    markdown.write_text(
        "PHAN I: Thí sinh trả lời từ câu 1 đến câu 12 Câu 1. Chọn đáp án.\n\n"
        '<div style="text-align: center;">Câu 2. Dữ liệu như sau.</div>\n\n'
        '<div style="text-align: center;"><img src="https://assets.example.test/'
        'img_in_image_box_10_20_110_80.jpg" alt="Biểu đồ" width="37%" /></div>\n',
        encoding="utf-8",
    )

    content = parse_markdown_content(markdown)

    assert [block.kind for block in content.blocks] == [
        MarkdownBlockKind.HEADING,
        MarkdownBlockKind.PARAGRAPH,
        MarkdownBlockKind.PARAGRAPH,
        MarkdownBlockKind.IMAGE,
    ]
    assert content.blocks[0].metadata["role"] == "section_heading"
    assert content.blocks[1].starts_group
    assert content.blocks[1].text == "Câu 1. Chọn đáp án."
    assert content.blocks[2].starts_group
    assert content.blocks[2].metadata["alignment"] == "center"
    assert content.blocks[3].source is not None
    assert content.blocks[3].metadata["alignment"] == "center"
    assert content.blocks[3].metadata["width_percent"] == 37


def test_remote_markdown_images_are_reused_from_disk_cache(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    payload = io.BytesIO()
    Image.new("RGB", (8, 6), "white").save(payload, format="PNG")
    calls = 0

    class Headers:
        @staticmethod
        def get(name: str) -> str | None:
            return None

        @staticmethod
        def get_content_type() -> str:
            return "image/png"

    class Response(io.BytesIO):
        headers = Headers()

        def __enter__(self) -> Response:
            return self

        def __exit__(self, *args: object) -> None:
            self.close()

    def open_once(*args: object, **kwargs: object) -> Response:
        nonlocal calls
        calls += 1
        if calls > 1:
            raise AssertionError("the cached URL should not be downloaded twice")
        return Response(payload.getvalue())

    monkeypatch.setenv("DOCRECONSTRUCT_ASSET_CACHE", str(tmp_path / "cache"))
    monkeypatch.setattr(
        "docreconstruct.reconstruction.asset_matching.socket.getaddrinfo",
        lambda *args, **kwargs: [
            (2, 1, 6, "", ("93.184.216.34", 443)),
        ],
    )
    monkeypatch.setattr(
        "docreconstruct.reconstruction.asset_matching._open_remote_asset",
        open_once,
    )
    block = MarkdownBlock(
        id="md-1",
        index=0,
        kind=MarkdownBlockKind.IMAGE,
        source="https://assets.example.test/figure.png",
    )
    _download_remote.cache_clear()
    first = resolve_markdown_asset(block, markdown_directory=tmp_path)
    _download_remote.cache_clear()
    second = resolve_markdown_asset(block, markdown_directory=tmp_path)

    assert first is not None and second is not None
    assert first.data == second.data == payload.getvalue()
    assert calls == 1


def test_local_markdown_assets_are_confined_to_markdown_directory(tmp_path: Path) -> None:
    markdown_directory = tmp_path / "authority"
    markdown_directory.mkdir()
    safe_image = markdown_directory / "figure.png"
    outside_image = tmp_path / "private.png"
    Image.new("RGB", (8, 6), "white").save(safe_image)
    Image.new("RGB", (8, 6), "black").save(outside_image)

    safe = MarkdownBlock(
        id="safe",
        index=0,
        kind=MarkdownBlockKind.IMAGE,
        source="figure.png",
    )
    resolved = resolve_markdown_asset(safe, markdown_directory=markdown_directory)
    assert resolved is not None
    assert resolved.source == "figure.png"

    for source in ("../private.png", str(outside_image.resolve())):
        escaped = safe.model_copy(update={"id": f"escaped-{source}", "source": source})
        with pytest.raises(ValueError, match="relative path|escapes the Markdown directory"):
            resolve_markdown_asset(escaped, markdown_directory=markdown_directory)


def test_remote_markdown_asset_rejects_private_dns_before_fetch(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    monkeypatch.setattr(
        "docreconstruct.reconstruction.asset_matching.socket.getaddrinfo",
        lambda *args, **kwargs: [(2, 1, 6, "", ("127.0.0.1", 443))],
    )
    monkeypatch.setattr(
        "docreconstruct.reconstruction.asset_matching._open_remote_asset",
        lambda *args, **kwargs: pytest.fail("private targets must be rejected before HTTP"),
    )
    block = MarkdownBlock(
        id="private-remote",
        index=0,
        kind=MarkdownBlockKind.IMAGE,
        source="https://assets.example.test/private.png",
    )
    _download_remote.cache_clear()
    with pytest.raises(ValueError, match="private, loopback, link-local, or reserved"):
        resolve_markdown_asset(block, markdown_directory=tmp_path)


def test_remote_markdown_asset_revalidates_redirect_targets(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    monkeypatch.setattr(
        "docreconstruct.reconstruction.asset_matching.socket.getaddrinfo",
        lambda *args, **kwargs: [(2, 1, 6, "", ("169.254.169.254", 443))],
    )
    request = urllib.request.Request("https://public.example.test/figure.png")
    handler = _SafeAssetRedirectHandler()

    with pytest.raises(ValueError, match="private, loopback, link-local, or reserved"):
        handler.redirect_request(
            request,
            None,
            302,
            "Found",
            {},
            "https://169.254.169.254/latest/meta-data/",
        )


def test_data_uri_obeys_asset_byte_limit(tmp_path: Path) -> None:
    block = MarkdownBlock(
        id="large-data-uri",
        index=0,
        kind=MarkdownBlockKind.IMAGE,
        source="data:image/png;base64,QUJDREVGRw==",
    )
    with pytest.raises(ValueError, match="safety limit"):
        resolve_markdown_asset(block, markdown_directory=tmp_path, maximum_bytes=4)


def test_remote_asset_disk_cache_prunes_to_operator_limit(
    tmp_path: Path,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    cache = tmp_path / "cache"
    monkeypatch.setenv("DOCRECONSTRUCT_ASSET_CACHE_MAX_MB", "1")
    first = cache / "first.bin"
    second = cache / "second.bin"
    payload = b"x" * (600 * 1024)

    _write_asset_cache(cache, first, cache / "first.type", payload, "image/png")
    _write_asset_cache(cache, second, cache / "second.type", payload, "image/png")

    assert not first.exists()
    assert second.stat().st_size == len(payload)
    assert sum(path.stat().st_size for path in cache.glob("*.bin")) <= 1024 * 1024


def test_provider_crop_filename_is_an_offline_source_figure_fallback(
    tmp_path: Path,
) -> None:
    markdown = tmp_path / "content.md"
    markdown.write_text(
        '<div><img src="https://assets.invalid/img_in_image_box_10_20_110_80.jpg" '
        'alt="Figure" /></div>\n',
        encoding="utf-8",
    )
    content = parse_markdown_content(markdown)
    page = ScanPageLayout(
        number=1,
        width=200,
        height=160,
        pdf_width=595,
        pdf_height=476,
        content_bbox=PixelBox(x0=0, y0=0, x1=200, y1=160),
        line_pitch=20,
        image=Image.new("RGB", (200, 160), "white"),
        metadata={"source_kind": "image", "rectified": False},
    )
    layout = ScanDocumentLayout(source=str(tmp_path / "layout.png"), pages=[page])

    matches = match_markdown_assets(content, layout, allow_remote=False)

    assert len(matches) == 1
    assert matches[0].bbox == PixelBox(x0=10, y0=20, x1=110, y1=80)
    assert not matches[0].resolved


def test_provider_page_folder_selects_explicit_page_across_equal_sized_pages(
    tmp_path: Path,
) -> None:
    markdown = tmp_path / "content.md"
    sources = [
        "https://assets.invalid/export/"
        f"markdown_{page_index}/imgs/img_in_image_box_10_20_110_80.jpg"
        for page_index in range(4)
    ]
    sources.extend(
        [
            # An explicit page behind the monotonic cursor must not move backwards.
            "https://assets.invalid/export/markdown_1/imgs/img_in_image_box_20_30_120_90.jpg",
            # An explicit page must not fall through to another page when out of bounds.
            "https://assets.invalid/export/markdown_3/imgs/img_in_image_box_10_20_210_80.jpg",
        ]
    )
    markdown.write_text(
        "\n".join(f'<img src="{source}" alt="Figure" />' for source in sources) + "\n",
        encoding="utf-8",
    )
    content = parse_markdown_content(markdown)
    pages = [
        ScanPageLayout(
            number=page_number,
            width=200,
            height=160,
            pdf_width=595,
            pdf_height=476,
            content_bbox=PixelBox(x0=0, y0=0, x1=200, y1=160),
            line_pitch=20,
            image=Image.new("RGB", (200, 160), "white"),
            metadata={"source_kind": "pdf", "rectified": False},
        )
        for page_number in range(1, 5)
    ]
    layout = ScanDocumentLayout(source=str(tmp_path / "layout.pdf"), pages=pages)

    matches = match_markdown_assets(content, layout, allow_remote=False)

    assert [match.page_number for match in matches] == [1, 2, 3, 4]
    assert all(match.bbox == PixelBox(x0=10, y0=20, x1=110, y1=80) for match in matches)
    assert all(not match.resolved for match in matches)


def test_docx_side_layout_keeps_valid_trailing_paragraph_after_nested_table(
    tmp_path: Path,
) -> None:
    figure = tmp_path / "figure.png"
    Image.new("RGB", (120, 80), "white").save(figure)
    group = "section-1:item-1"
    blocks = [
        MarkdownBlock(
            id="md-1",
            index=0,
            kind=MarkdownBlockKind.PARAGRAPH,
            text="Question 1: Editable prompt",
            group_id=group,
            starts_group=True,
        ),
        *[
            MarkdownBlock(
                id=f"md-{index + 2}",
                index=index + 1,
                kind=MarkdownBlockKind.OPTION,
                text=f"{label}. choice",
                group_id=group,
            )
            for index, label in enumerate("ABCD")
        ],
        MarkdownBlock(
            id="md-6",
            index=5,
            kind=MarkdownBlockKind.IMAGE,
            source=str(figure),
            group_id=group,
        ),
    ]
    content = MarkdownContent(source=str(tmp_path / "content.md"), blocks=blocks)
    content_box = PixelBox(x0=30, y0=30, x1=570, y1=770)
    figure_box = PixelBox(x0=390, y0=90, x1=540, y1=250)
    scan_page = ScanPageLayout(
        number=1,
        width=600,
        height=800,
        pdf_width=595,
        pdf_height=842,
        content_bbox=content_box,
        line_pitch=20,
        image=Image.new("RGB", (600, 800), "white"),
    )
    scan = ScanDocumentLayout(source=str(tmp_path / "layout.pdf"), pages=[scan_page])
    placements = [
        HybridBlockPlacement(
            block_id=block.id,
            block_index=block.index,
            page_number=1,
            source_bbox=figure_box if block.kind is MarkdownBlockKind.IMAGE else None,
        )
        for block in blocks
    ]
    plan = HybridLayoutPlan(
        content_source=content.source,
        layout_source=scan.source,
        pages=[
            HybridPagePlan(
                number=1,
                pdf_width=595,
                pdf_height=842,
                raster_width=600,
                raster_height=800,
                content_bbox=content_box,
                line_pitch=20,
                placements=placements,
            )
        ],
    )
    match = AssetMatch(
        block_id="md-6",
        source=str(figure),
        page_number=1,
        bbox=figure_box,
        score=1.0,
    )

    payload_bytes = render_hybrid_docx(content, scan, plan, [match])
    with zipfile.ZipFile(io.BytesIO(payload_bytes)) as package:
        root = ElementTree.fromstring(package.read("word/document.xml"))
    namespace = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    nested_table_cells = []
    for cell in root.iter(f"{namespace}tc"):
        children = [child for child in cell if child.tag != f"{namespace}tcPr"]
        if any(child.tag == f"{namespace}tbl" for child in children):
            nested_table_cells.append(children)

    assert nested_table_cells
    assert all(children[-1].tag == f"{namespace}p" for children in nested_table_cells)


def test_same_group_disjoint_images_render_in_one_native_horizontal_row(
    tmp_path: Path,
) -> None:
    group = "section-1:question-1"
    blocks = [
        MarkdownBlock(
            id="prompt",
            index=0,
            kind=MarkdownBlockKind.PARAGRAPH,
            text="Editable prompt above two figures.",
            group_id=group,
            starts_group=True,
        ),
        MarkdownBlock(
            id="left",
            index=1,
            kind=MarkdownBlockKind.IMAGE,
            source="left.png",
            group_id=group,
        ),
        MarkdownBlock(
            id="right",
            index=2,
            kind=MarkdownBlockKind.IMAGE,
            source="right.png",
            group_id=group,
        ),
        MarkdownBlock(
            id="answer",
            index=3,
            kind=MarkdownBlockKind.PARAGRAPH,
            text="Editable answer below both figures.",
            group_id=group,
        ),
    ]
    content = MarkdownContent(source=str(tmp_path / "content.md"), blocks=blocks)
    content_box = PixelBox(x0=40, y0=25, x1=560, y1=775)
    prompt_box = PixelBox(x0=60, y0=80, x1=540, y1=105)
    left_box = PixelBox(x0=80, y0=180, x1=270, y1=360)
    right_box = PixelBox(x0=330, y0=160, x1=530, y1=370)
    answer_box = PixelBox(x0=60, y0=420, x1=540, y1=445)
    page_image = Image.new("RGB", (600, 800), "white")
    scan_page = ScanPageLayout(
        number=1,
        width=600,
        height=800,
        pdf_width=595,
        pdf_height=793,
        content_bbox=content_box,
        line_pitch=30,
        image=page_image,
        metadata={"source_kind": "pdf", "column_count": 1},
    )
    scan = ScanDocumentLayout(source=str(tmp_path / "layout.pdf"), pages=[scan_page])
    boxes = [prompt_box, left_box, right_box, answer_box]
    placements = [
        HybridBlockPlacement(
            block_id=block.id,
            block_index=block.index,
            page_number=1,
            source_bbox=box,
            source_rows=[] if block.kind is MarkdownBlockKind.IMAGE else [box],
            source_gap_before=0,
        )
        for block, box in zip(blocks, boxes, strict=True)
    ]
    plan = HybridLayoutPlan(
        content_source=content.source,
        layout_source=scan.source,
        pages=[
            HybridPagePlan(
                number=1,
                pdf_width=595,
                pdf_height=793,
                raster_width=600,
                raster_height=800,
                content_bbox=content_box,
                line_pitch=30,
                placements=placements,
            )
        ],
    )

    payload = render_hybrid_docx(content, scan, plan, [])
    with zipfile.ZipFile(io.BytesIO(payload)) as package:
        root = ElementTree.fromstring(package.read("word/document.xml"))
    word = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    drawing = "{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}"
    visual_row = next(
        table
        for table in root.iter(f"{word}tbl")
        if (caption := table.find(f"{word}tblPr/{word}tblCaption")) is not None
        and caption.get(f"{word}val") == "docreconstruct:horizontal-visual-row"
    )

    assert len(list(visual_row.iter(f"{drawing}inline"))) == 2


def _mixed_masthead_fixture(tmp_path: Path) -> tuple[Path, Path]:
    markdown = tmp_path / "mixed-layout.md"
    markdown.write_text(
        "# ORGANIZATION NAME ANNUAL GENERAL EXAMINATION\n\n"
        "OFFICIAL COPY\n\nSubject: EDITABLE TEXT\n\n(2 pages)\n\n"
        "Time allowed: 90 minutes\n\nName:\n\nCandidate ID:\n\n"
        "I. READING (4 points)\n\nRead the passage:\n\n"
        "(1) First quoted paragraph remains editable and uses source wording.\n\n"
        "(2) Second quoted paragraph remains editable and uses source wording.\n\n"
        "(3) Third quoted paragraph remains editable and uses source wording.\n\n"
        "(Source publication, page 5)\n\nAnswer the questions:\n\n"
        "Question 1. Identify the main claim.\n",
        encoding="utf-8",
    )
    layout = tmp_path / "mixed-layout.png"
    image = Image.new("RGB", (600, 840), "white")
    from PIL import ImageDraw

    draw = ImageDraw.Draw(image)

    def text_strokes(left: int, right: int, top: int) -> None:
        cursor = left
        while cursor < right:
            end = min(right, cursor + 42)
            draw.rectangle((cursor, top, end, top + 8), fill="black")
            cursor = end + 8

    for top in (28, 55, 82):
        text_strokes(35, 225, top)
        text_strokes(295, 560, top)
    text_strokes(35, 350, 122)
    text_strokes(35, 350, 148)
    text_strokes(35, 250, 186)
    for top in range(218, 720, 24):
        left = 70 if top in {218, 314, 410} else 35
        text_strokes(left, 560, top)
    image.save(layout)
    return markdown, layout


def test_scan_analysis_recovers_dense_lines_and_split_masthead(tmp_path: Path) -> None:
    pytest.importorskip("numpy")
    _, layout = _mixed_masthead_fixture(tmp_path)

    page = analyze_scan_source(layout).pages[0]

    assert page.metadata["column_count"] == 1
    assert page.metadata["header_column_count"] == 2
    assert len(page.text_lines) >= 20
    assert len(page.line_bands) == len(page.text_lines)
    assert page.regions == []


# ECMA-376 child order for the property containers this renderer writes into.
_PROPERTY_SEQUENCES = {
    "w:pPr": (
        "w:pStyle",
        "w:keepNext",
        "w:keepLines",
        "w:pageBreakBefore",
        "w:framePr",
        "w:widowControl",
        "w:numPr",
        "w:suppressLineNumbers",
        "w:pBdr",
        "w:shd",
        "w:tabs",
        "w:suppressAutoHyphens",
        "w:kinsoku",
        "w:wordWrap",
        "w:overflowPunct",
        "w:topLinePunct",
        "w:autoSpaceDE",
        "w:autoSpaceDN",
        "w:bidi",
        "w:adjustRightInd",
        "w:snapToGrid",
        "w:spacing",
        "w:ind",
        "w:contextualSpacing",
        "w:mirrorIndents",
        "w:suppressOverlap",
        "w:jc",
        "w:textDirection",
        "w:textAlignment",
        "w:textboxTightWrap",
        "w:outlineLvl",
        "w:divId",
        "w:cnfStyle",
        "w:rPr",
        "w:sectPr",
        "w:pPrChange",
    ),
    "w:tcPr": (
        "w:cnfStyle",
        "w:tcW",
        "w:gridSpan",
        "w:hMerge",
        "w:vMerge",
        "w:tcBorders",
        "w:shd",
        "w:noWrap",
        "w:tcMar",
        "w:textDirection",
        "w:tcFitText",
        "w:vAlign",
        "w:hideMark",
        "w:headers",
        "w:cellIns",
        "w:cellDel",
        "w:cellMerge",
        "w:tcPrChange",
    ),
    "w:tblPr": (
        "w:tblStyle",
        "w:tblpPr",
        "w:tblOverlap",
        "w:bidiVisual",
        "w:tblStyleRowBandSize",
        "w:tblStyleColBandSize",
        "w:tblW",
        "w:jc",
        "w:tblCellSpacing",
        "w:tblInd",
        "w:tblBorders",
        "w:shd",
        "w:tblLayout",
        "w:tblCellMar",
        "w:tblLook",
        "w:tblCaption",
        "w:tblDescription",
        "w:tblPrChange",
    ),
}


def test_generated_property_containers_follow_the_ooxml_child_order(tmp_path: Path) -> None:
    """Word's strict parser rejects an out-of-sequence property container.

    `w:pPr`, `w:tcPr` and `w:tblPr` are `xsd:sequence`, so appending a border or
    a shading element after python-docx has written a later-ordered sibling
    makes the document invalid. Word reports "unreadable content" and repairs it
    by discarding the property, so the thematic rule loses its line and table
    borders and header shading disappear. LibreOffice is lenient, which is why a
    rendered comparison does not catch it.
    """

    pytest.importorskip("numpy")
    markdown = tmp_path / "ruled.md"
    markdown.write_text(
        "Intro paragraph text here.\n\n---\n\n"
        "<table><tr><td>Head A</td><td>Head B</td></tr>"
        "<tr><td>one</td><td>two</td></tr></table>\n\nTail paragraph.\n",
        encoding="utf-8",
    )
    from PIL import ImageDraw

    image = Image.new("RGB", (620, 877), "white")
    draw = ImageDraw.Draw(image)
    for index in range(8):
        top = 60 + index * 70
        draw.rectangle((55, top, 500, top + 28), fill="black")
    layout = tmp_path / "ruled.png"
    image.save(layout)
    output = tmp_path / "ruled.docx"

    reconstruct_hybrid(markdown, layout, output=output)

    with zipfile.ZipFile(output) as package:
        document_xml = package.read("word/document.xml").decode("utf-8")

    inspected = 0
    for container, sequence in _PROPERTY_SEQUENCES.items():
        for match in re.finditer(rf"<{container}>(.*?)</{container}>", document_xml, re.S):
            present = [
                tag for tag in re.findall(r"<(w:[A-Za-z]+)", match.group(1)) if tag in sequence
            ]
            inspected += 1
            assert present == sorted(present, key=sequence.index), (
                f"{container} children out of schema order: {present}"
            )
    assert inspected >= 3
    # CT_Shd requires w:val; "clear" is the plain solid fill.
    assert 'w:val="clear"' in document_xml


def test_column_wrapping_never_breaks_a_protected_inline_span() -> None:
    """A hard wrap must not split inline math, code, or a URL.

    Every construct `parse_markdown_inline` protects is anchored to a single
    line, so a break inside one destroys it: the formula becomes literal dollar
    signs and a fragment of it is promoted to an unrelated equation. Plain prose
    must still wrap, otherwise the two-column layout loses its line budget.
    """

    def block(text: str) -> MarkdownBlock:
        return MarkdownBlock(id="md-1", index=0, kind=MarkdownBlockKind.PARAGRAPH, text=text)

    def math_of(text: str) -> list[str]:
        return [segment.value for segment in parse_markdown_inline(text) if segment.is_math]

    formula = r"Total energy is $E = mc^2 + \frac{1}{2} m v^2$ everywhere in the derivation."
    wrapped = _wrap_column_blocks([block(formula)], characters_per_line=28)[0]

    assert wrapped.text == formula
    assert math_of(wrapped.text) == math_of(formula)
    assert math_of(formula) == [r"E = mc^2 + \frac{1}{2} m v^2"]

    # A code span too long to fit one line, whose contents would otherwise be
    # re-read as TeX once a break lands inside it, is protected too.
    code = "Run `x_{1} and y_{2} and z_{3} and w_{4} here` now to finish the line."
    assert math_of(code) == []
    assert _wrap_column_blocks([block(code)], characters_per_line=28)[0].text == code

    prose = (
        "This is ordinary prose that is long enough to be wrapped across "
        "several lines by the column layout."
    )
    wrapped_prose = _wrap_column_blocks([block(prose)], characters_per_line=28)[0]

    assert "\n" in wrapped_prose.text
    assert wrapped_prose.text.replace("\n", " ") == prose


def test_image_table_pair_keeps_the_other_visuals_in_its_group() -> None:
    """A third figure in the group must still reach the document.

    The surviving flow was rebuilt by excluding every block whose kind is an
    image or a table, not just the two the pair renders, so any additional
    figure or table was absent from the pair, from `pre` and from `post`, and
    disappeared. No gate catches it: content projection skips image blocks, and
    visual slot coverage is measured on the plan rather than the render.
    """

    page = ScanPageLayout(
        number=1,
        width=1000,
        height=1400,
        pdf_width=612,
        pdf_height=792,
        content_bbox=PixelBox(x0=50, y0=50, x1=950, y1=1350),
        line_pitch=20.0,
        image=Image.new("RGB", (1000, 1400), "white"),
        metadata={"source_kind": "pdf"},
    )
    layout = ScanDocumentLayout(source="layout.pdf", pages=[page])
    kinds = [MarkdownBlockKind.IMAGE, MarkdownBlockKind.TABLE, MarkdownBlockKind.IMAGE]
    blocks = [
        MarkdownBlock(
            id=f"md-{index + 1}",
            index=index,
            kind=kind,
            text="",
            table_rows=[["a", "b"], ["c", "d"]] if kind is MarkdownBlockKind.TABLE else [],
        )
        for index, kind in enumerate(kinds)
    ]
    # The first figure and the table sit side by side; the second figure is below.
    boxes = {
        "md-1": PixelBox(x0=60, y0=100, x1=480, y1=400),
        "md-2": PixelBox(x0=520, y0=100, x1=940, y1=400),
        "md-3": PixelBox(x0=60, y0=500, x1=480, y1=800),
    }
    placements = {
        block.id: HybridBlockPlacement(
            block_id=block.id,
            block_index=block.index,
            page_number=1,
            source_bbox=boxes[block.id],
        )
        for block in blocks
    }

    document = WordDocument()
    rendered = _render_image_table_pair(
        document,
        blocks,
        placements,
        width=6.0,
        size=11.0,
        line_height=1.15,
        asset_bytes={},
        layout=layout,
    )

    assert rendered
    assert document.element.xml.count("<pic:pic") == 2


def _text_only_exam(tmp_path: Path) -> tuple[Path, Path]:
    """A Vietnamese exam with no figure anywhere, as reviewed Markdown + scan."""

    from PIL import ImageDraw

    markdown = tmp_path / "exam.md"
    markdown.write_text(
        "Câu 5: Tính giá trị của biểu thức sau\n\n"
        "A. 1\n\nB. 2\n\nC. 3\n\nD. 4\n\n"
        "Hình 2\n\n"
        "Câu 6: Cho hàm số y = f(x)\n\n"
        "A. 5\n\nB. 6\n",
        encoding="utf-8",
    )
    layout = tmp_path / "exam.png"
    image = Image.new("RGB", (620, 877), "white")
    draw = ImageDraw.Draw(image)
    for index in range(9):
        top = 60 + index * 60
        draw.rectangle((55, top, 515 - (index % 3) * 40, top + 26), fill="black")
    image.save(layout)
    return markdown, layout


def test_option_labels_survive_when_no_figure_is_rendered(tmp_path: Path) -> None:
    """Reviewed Markdown must not be deleted on the promise of absent pixels.

    The duplicate-annotation rule assumes a matched figure already carries the
    label's pixels. A text-only exam matches the same shape — a short
    unpunctuated line after the last option, before the next question — so the
    text was deleted with no figure reproducing it, and the
    `native_content_projection` gate, which still requires that text, failed.
    """

    pytest.importorskip("numpy")
    markdown, layout = _text_only_exam(tmp_path)
    content = parse_markdown_content(markdown)
    annotation = next(block for block in content.blocks if block.text == "Hình 2")

    # The shape still matches; only the missing figure withholds the deletion.
    assert _duplicate_figure_annotation_ids(content.blocks, figures_rendered=True) == {
        annotation.id
    }
    assert _duplicate_figure_annotation_ids(content.blocks, figures_rendered=False) == set()

    output = tmp_path / "exam.docx"
    reconstruct_hybrid(markdown, layout, output=output)
    with zipfile.ZipFile(output) as package:
        root = ElementTree.fromstring(package.read("word/document.xml"))
    namespace = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    visible = "".join(node.text or "" for node in root.iter(f"{namespace}t"))

    assert "Hình 2" in visible

    report = validate_hybrid(markdown, layout, output)
    projection = next(gate for gate in report.gates if gate.name == "native_content_projection")
    assert projection.passed


def test_structural_roles_and_masthead_rendering_remain_generic(tmp_path: Path) -> None:
    pytest.importorskip("numpy")
    markdown, layout = _mixed_masthead_fixture(tmp_path)
    parsed = parse_markdown_content(markdown)
    roles = {block.text: block.metadata.get("role") for block in parsed.blocks}

    assert roles["I. READING (4 points)"] == "section_heading"
    assert roles["Name:"] == roles["Candidate ID:"] == "form_field"
    assert roles["Read the passage:"] == "passage_lead"
    assert roles["(2) Second quoted paragraph remains editable and uses source wording."] == (
        "quoted_passage"
    )
    assert roles["(Source publication, page 5)"] == "attribution"
    assert roles["Answer the questions:"] == "question_lead"

    output = tmp_path / "mixed-layout.docx"
    reconstruct_hybrid(markdown, layout, output=output)
    with zipfile.ZipFile(output) as package:
        document_xml = package.read("word/document.xml")
        assert not any(name.startswith("word/media/") for name in package.namelist())
    root = ElementTree.fromstring(document_xml)
    namespace = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    texts = [node.text or "" for node in root.iter(f"{namespace}t")]
    visible = " ".join(texts)

    assert visible.index("ORGANIZATION NAME") < visible.index("OFFICIAL COPY")
    assert visible.index("OFFICIAL COPY") < visible.index("Subject: EDITABLE TEXT")
    assert root.find(f".//{namespace}tbl") is not None
    assert root.find(f".//{namespace}tabs/{namespace}tab[@{namespace}leader='dot']") is not None
    assert root.find(f".//{namespace}i") is not None
    question = next(
        paragraph
        for paragraph in root.iter(f"{namespace}p")
        if "Question 1." in "".join(node.text or "" for node in paragraph.iter(f"{namespace}t"))
    )
    properties = question.find(f"{namespace}pPr")
    assert properties is not None
    keep_next = properties.find(f"{namespace}keepNext")
    assert keep_next is None or keep_next.get(f"{namespace}val") in {"0", "false", "off"}


def test_separate_masthead_blocks_and_bottom_furniture_use_native_zones(
    tmp_path: Path,
) -> None:
    markdown = tmp_path / "separate-masthead.md"
    markdown.write_text(
        "# AUTHORITY NAME\n\n"
        "# OFFICIAL COPY\n\n"
        "(4 pages)\n\n"
        "# NATIONAL EXAMINATION 2026 Subject: MATHEMATICS\n\n"
        "Candidate name: ..... Candidate ID: .....\n\n"
        "Code: 0110\n\n"
        "I. QUESTIONS\n\n"
        "Question 1. Editable prompt remains full width.\n\n"
        "Page 1/4 - Code\n",
        encoding="utf-8",
    )
    content = parse_markdown_content(markdown)
    boxes = [
        PixelBox(x0=60, y0=30 + index * 34, x1=540, y1=55 + index * 34) for index in range(8)
    ] + [PixelBox(x0=440, y0=760, x1=555, y1=778)]
    page = ScanPageLayout(
        number=1,
        width=600,
        height=800,
        pdf_width=595,
        pdf_height=793,
        content_bbox=PixelBox(x0=35, y0=20, x1=565, y1=780),
        line_pitch=30,
        image=Image.new("RGB", (600, 800), "white"),
        metadata={
            "source_kind": "image",
            "column_count": 1,
            "header_column_count": 2,
            "header_divider": 270,
            "render_content_bbox": {"x0": 55, "y0": 20, "x1": 555, "y1": 780},
        },
    )
    scan = ScanDocumentLayout(source=str(tmp_path / "layout.png"), pages=[page])
    placements = [
        HybridBlockPlacement(
            block_id=block.id,
            block_index=block.index,
            page_number=1,
            source_bbox=boxes[index],
            source_rows=[boxes[index]],
            source_gap_before=0,
            match_score=1.0,
        )
        for index, block in enumerate(content.blocks)
    ]
    plan = HybridLayoutPlan(
        content_source=content.source,
        layout_source=scan.source,
        pages=[
            HybridPagePlan(
                number=1,
                pdf_width=page.pdf_width,
                pdf_height=page.pdf_height,
                raster_width=page.width,
                raster_height=page.height,
                content_bbox=page.content_bbox,
                line_pitch=page.line_pitch,
                placements=placements,
            )
        ],
    )

    payload = render_hybrid_docx(content, scan, plan, [])
    with zipfile.ZipFile(io.BytesIO(payload)) as package:
        root = ElementTree.fromstring(package.read("word/document.xml"))
        footer = ElementTree.fromstring(package.read("word/footer1.xml"))
    word = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    masthead = next(
        table
        for table in root.iter(f"{word}tbl")
        if (caption := table.find(f"{word}tblPr/{word}tblCaption")) is not None
        and caption.get(f"{word}val") == "docreconstruct:split-masthead"
    )
    cells = masthead.findall(f"{word}tr/{word}tc")
    assert len(cells) == 2
    left = " ".join(node.text or "" for node in cells[0].iter(f"{word}t"))
    right = " ".join(node.text or "" for node in cells[1].iter(f"{word}t"))
    assert "AUTHORITY NAME" in left and "Candidate name" in left
    assert "NATIONAL EXAMINATION" in right and "Code: 0110" in right
    body_text = " ".join(node.text or "" for node in root.iter(f"{word}t"))
    footer_text = " ".join(node.text or "" for node in footer.iter(f"{word}t"))
    assert "Page 1/4 - Code" not in body_text
    assert footer_text == "Page 1/4 - Code"
    footer_alignment = footer.find(f".//{word}jc")
    assert footer_alignment is not None
    assert footer_alignment.get(f"{word}val") == "right"

    left_footer_box = PixelBox(x0=45, y0=760, x1=160, y1=778)
    plan.pages[0].placements[-1] = (
        plan.pages[0]
        .placements[-1]
        .model_copy(update={"source_bbox": left_footer_box, "source_rows": [left_footer_box]})
    )
    left_payload = render_hybrid_docx(content, scan, plan, [])
    with zipfile.ZipFile(io.BytesIO(left_payload)) as package:
        left_footer = ElementTree.fromstring(package.read("word/footer1.xml"))
    left_alignment = left_footer.find(f".//{word}jc")
    assert left_alignment is not None
    assert left_alignment.get(f"{word}val") == "left"


def test_split_masthead_uses_source_columns_without_splitting_authority_block(
    tmp_path: Path,
) -> None:
    markdown = tmp_path / "positioned-masthead.md"
    markdown.write_text(
        "CENTRAL EDUCATION AUTHORITY\nOFFICIAL EXAMINATION\n(04 pages)\n\n"
        "# NATIONAL COMPLETION EXAMINATION 2026 Subject: PHYSICS Time: 50 minutes\n\n"
        "Candidate name: ..... Candidate ID: .....\n\n"
        "Code: 0204\n\n"
        "## PART I: Choose one answer\n\n"
        "Question 1. Editable prompt.\n",
        encoding="utf-8",
    )
    content = parse_markdown_content(markdown)
    source_boxes = [
        PixelBox(x0=80, y0=35, x1=265, y1=125),
        PixelBox(x0=320, y0=35, x1=555, y1=115),
        PixelBox(x0=55, y0=135, x1=345, y1=180),
        PixelBox(x0=445, y0=140, x1=555, y1=172),
        PixelBox(x0=55, y0=195, x1=550, y1=225),
        PixelBox(x0=55, y0=235, x1=550, y1=270),
    ]
    page = ScanPageLayout(
        number=1,
        width=600,
        height=800,
        pdf_width=595,
        pdf_height=793,
        content_bbox=PixelBox(x0=40, y0=20, x1=565, y1=780),
        line_pitch=30,
        image=Image.new("RGB", (600, 800), "white"),
        metadata={
            "source_kind": "image",
            "column_count": 1,
            "header_column_count": 2,
            "header_divider": 285,
            "render_content_bbox": {"x0": 50, "y0": 20, "x1": 560, "y1": 780},
        },
    )
    scan = ScanDocumentLayout(source=str(tmp_path / "layout.png"), pages=[page])
    placements = [
        HybridBlockPlacement(
            block_id=block.id,
            block_index=block.index,
            page_number=1,
            source_bbox=source_boxes[index],
            source_rows=[source_boxes[index]],
            source_gap_before=0,
            match_score=1.0,
            geometry_source="json_consensus",
        )
        for index, block in enumerate(content.blocks)
    ]
    plan = HybridLayoutPlan(
        content_source=content.source,
        layout_source=scan.source,
        pages=[
            HybridPagePlan(
                number=1,
                pdf_width=page.pdf_width,
                pdf_height=page.pdf_height,
                raster_width=page.width,
                raster_height=page.height,
                content_bbox=page.content_bbox,
                line_pitch=page.line_pitch,
                placements=placements,
            )
        ],
    )

    payload = render_hybrid_docx(content, scan, plan, [])
    with zipfile.ZipFile(io.BytesIO(payload)) as package:
        root = ElementTree.fromstring(package.read("word/document.xml"))
    word = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    masthead = next(
        table
        for table in root.iter(f"{word}tbl")
        if (caption := table.find(f"{word}tblPr/{word}tblCaption")) is not None
        and caption.get(f"{word}val") == "docreconstruct:split-masthead"
    )
    left_cell, right_cell = masthead.findall(f"{word}tr/{word}tc")
    left_text = " ".join(node.text or "" for node in left_cell.iter(f"{word}t"))
    right_text = " ".join(node.text or "" for node in right_cell.iter(f"{word}t"))

    authority = "CENTRAL EDUCATION AUTHORITY OFFICIAL EXAMINATION (04 pages)"
    assert authority in left_text
    assert "NATIONAL COMPLETION EXAMINATION" not in left_text
    assert left_text.index(authority) < left_text.index("Candidate name")
    assert "NATIONAL COMPLETION EXAMINATION" in right_text
    assert right_text.index("NATIONAL COMPLETION EXAMINATION") < right_text.index("Code: 0204")
    assert authority not in right_text

    authority_paragraph = next(
        paragraph
        for paragraph in left_cell.iter(f"{word}p")
        if "CENTRAL EDUCATION AUTHORITY"
        in "".join(node.text or "" for node in paragraph.iter(f"{word}t"))
    )
    assert authority_paragraph.find(f".//{word}br") is not None

    title_paragraph = next(
        paragraph
        for paragraph in right_cell.iter(f"{word}p")
        if "NATIONAL COMPLETION EXAMINATION"
        in "".join(node.text or "" for node in paragraph.iter(f"{word}t"))
    )
    title_alignment = title_paragraph.find(f"{word}pPr/{word}jc")
    assert title_alignment is not None
    assert title_alignment.get(f"{word}val") == "center"

    code_paragraph = next(
        paragraph
        for paragraph in right_cell.iter(f"{word}p")
        if "Code: 0204" in "".join(node.text or "" for node in paragraph.iter(f"{word}t"))
    )
    borders = code_paragraph.find(f"{word}pPr/{word}pBdr")
    assert borders is not None
    assert {
        edge.tag.rsplit("}", 1)[-1] for edge in borders if edge.get(f"{word}val") == "single"
    } == {"top", "left", "bottom", "right"}
    code_indent = code_paragraph.find(f"{word}pPr/{word}ind")
    assert code_indent is not None
    assert int(code_indent.get(f"{word}left", "0")) > 0
    assert int(code_indent.get(f"{word}right", "0")) > 0


def test_no_source_footer_does_not_materialize_empty_footer_part(tmp_path: Path) -> None:
    block = MarkdownBlock(
        id="page-1-body",
        index=0,
        kind=MarkdownBlockKind.PARAGRAPH,
        text="Editable first-page body.",
    )
    content = MarkdownContent(source=str(tmp_path / "content.md"), blocks=[block])
    page = ScanPageLayout(
        number=1,
        width=600,
        height=800,
        pdf_width=595,
        pdf_height=793,
        content_bbox=PixelBox(x0=35, y0=20, x1=565, y1=780),
        line_pitch=30,
        image=Image.new("RGB", (600, 800), "white"),
        metadata={"source_kind": "image", "column_count": 1},
    )
    scan = ScanDocumentLayout(source=str(tmp_path / "layout.png"), pages=[page])
    body_box = PixelBox(x0=60, y0=100, x1=500, y1=124)
    placement = HybridBlockPlacement(
        block_id=block.id,
        block_index=block.index,
        page_number=1,
        source_bbox=body_box,
        source_rows=[body_box],
        source_gap_before=0,
    )
    plan = HybridLayoutPlan(
        content_source=content.source,
        layout_source=scan.source,
        pages=[
            HybridPagePlan(
                number=1,
                pdf_width=page.pdf_width,
                pdf_height=page.pdf_height,
                raster_width=page.width,
                raster_height=page.height,
                content_bbox=page.content_bbox,
                line_pitch=page.line_pitch,
                placements=[placement],
            )
        ],
    )

    payload = render_hybrid_docx(content, scan, plan, [])
    with zipfile.ZipFile(io.BytesIO(payload)) as package:
        names = package.namelist()
        document = ElementTree.fromstring(package.read("word/document.xml"))
    word = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

    assert not any(name.startswith("word/footer") for name in names)
    assert document.find(f".//{word}footerReference") is None


def test_footer_to_no_footer_section_emits_empty_unlink(tmp_path: Path) -> None:
    blocks = [
        MarkdownBlock(
            id="page-1-body",
            index=0,
            kind=MarkdownBlockKind.PARAGRAPH,
            text="Editable first-page body.",
        ),
        MarkdownBlock(
            id="page-1-footer",
            index=1,
            kind=MarkdownBlockKind.PARAGRAPH,
            text="Page 1/2",
        ),
        MarkdownBlock(
            id="page-2-body",
            index=2,
            kind=MarkdownBlockKind.PARAGRAPH,
            text="Editable second-page body.",
        ),
    ]
    content = MarkdownContent(source=str(tmp_path / "content.md"), blocks=blocks)
    pages = [
        ScanPageLayout(
            number=number,
            width=600,
            height=800,
            pdf_width=595,
            pdf_height=793,
            content_bbox=PixelBox(x0=35, y0=20, x1=565, y1=780),
            line_pitch=30,
            image=Image.new("RGB", (600, 800), "white"),
            metadata={"source_kind": "image", "column_count": 1},
        )
        for number in (1, 2)
    ]
    scan = ScanDocumentLayout(source=str(tmp_path / "layout.png"), pages=pages)
    first_body = PixelBox(x0=60, y0=100, x1=500, y1=124)
    first_footer = PixelBox(x0=480, y0=760, x1=555, y1=780)
    second_body = PixelBox(x0=60, y0=100, x1=500, y1=124)
    placements = [
        HybridBlockPlacement(
            block_id="page-1-body",
            block_index=0,
            page_number=1,
            source_bbox=first_body,
            source_rows=[first_body],
            source_gap_before=0,
        ),
        HybridBlockPlacement(
            block_id="page-1-footer",
            block_index=1,
            page_number=1,
            source_bbox=first_footer,
            source_rows=[first_footer],
            source_gap_before=0,
        ),
        HybridBlockPlacement(
            block_id="page-2-body",
            block_index=2,
            page_number=2,
            source_bbox=second_body,
            source_rows=[second_body],
            source_gap_before=0,
        ),
    ]
    plan = HybridLayoutPlan(
        content_source=content.source,
        layout_source=scan.source,
        pages=[
            HybridPagePlan(
                number=page.number,
                pdf_width=page.pdf_width,
                pdf_height=page.pdf_height,
                raster_width=page.width,
                raster_height=page.height,
                content_bbox=page.content_bbox,
                line_pitch=page.line_pitch,
                placements=[
                    placement for placement in placements if placement.page_number == page.number
                ],
            )
            for page in pages
        ],
    )

    payload = render_hybrid_docx(content, scan, plan, [])
    with zipfile.ZipFile(io.BytesIO(payload)) as package:
        document = ElementTree.fromstring(package.read("word/document.xml"))
        footer_names = sorted(name for name in package.namelist() if name.startswith("word/footer"))
        footers = [ElementTree.fromstring(package.read(name)) for name in footer_names]
    word = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    footer_texts = [
        " ".join(node.text or "" for node in footer.iter(f"{word}t")) for footer in footers
    ]

    assert footer_texts == ["Page 1/2", ""]
    assert len(document.findall(f".//{word}footerReference")) == 2


def test_multilingual_page_fraction_footer_fallback_is_unique_per_section(
    tmp_path: Path,
) -> None:
    blocks = [
        MarkdownBlock(
            id="page-1-body",
            index=0,
            kind=MarkdownBlockKind.PARAGRAPH,
            text="Editable first-page body.",
        ),
        MarkdownBlock(
            id="page-1-link",
            index=1,
            kind=MarkdownBlockKind.PARAGRAPH,
            text="Open source: https://example.invalid/document",
        ),
        MarkdownBlock(
            id="page-1-number",
            index=2,
            kind=MarkdownBlockKind.PARAGRAPH,
            text="Page 1/2",
        ),
        MarkdownBlock(
            id="page-2-body",
            index=3,
            kind=MarkdownBlockKind.PARAGRAPH,
            text="Editable second-page body.",
        ),
        MarkdownBlock(
            id="page-2-link",
            index=4,
            kind=MarkdownBlockKind.PARAGRAPH,
            text="Open source: https://example.invalid/document",
        ),
        MarkdownBlock(
            id="page-2-number",
            index=5,
            kind=MarkdownBlockKind.PARAGRAPH,
            text="Страница 2/2",
        ),
        MarkdownBlock(
            id="page-2-banner",
            index=6,
            kind=MarkdownBlockKind.PARAGRAPH,
            text="Repeated editable top banner.",
        ),
    ]
    content = MarkdownContent(source=str(tmp_path / "content.md"), blocks=blocks)
    pages = [
        ScanPageLayout(
            number=number,
            width=600,
            height=800,
            pdf_width=595,
            pdf_height=793,
            content_bbox=PixelBox(x0=35, y0=20, x1=565, y1=780),
            line_pitch=30,
            image=Image.new("RGB", (600, 800), "white"),
            metadata={"source_kind": "image", "column_count": 1},
        )
        for number in (1, 2)
    ]
    scan = ScanDocumentLayout(source=str(tmp_path / "layout.png"), pages=pages)
    first_page_boxes = [
        PixelBox(x0=60, y0=100, x1=500, y1=124),
        PixelBox(x0=180, y0=744, x1=500, y1=762),
        PixelBox(x0=480, y0=764, x1=555, y1=780),
    ]
    placements = [
        HybridBlockPlacement(
            block_id=blocks[index].id,
            block_index=index,
            page_number=1,
            source_bbox=box,
            source_rows=[box],
            source_gap_before=0,
        )
        for index, box in enumerate(first_page_boxes)
    ]
    second_body = PixelBox(x0=60, y0=100, x1=500, y1=124)
    top_banner = PixelBox(x0=160, y0=28, x1=500, y1=48)
    placements.extend(
        [
            HybridBlockPlacement(
                block_id="page-2-body",
                block_index=3,
                page_number=2,
                source_bbox=second_body,
                source_rows=[second_body],
                source_gap_before=0,
            ),
            HybridBlockPlacement(block_id="page-2-link", block_index=4, page_number=2),
            HybridBlockPlacement(block_id="page-2-number", block_index=5, page_number=2),
            HybridBlockPlacement(
                block_id="page-2-banner",
                block_index=6,
                page_number=2,
                source_bbox=top_banner,
                source_rows=[top_banner],
                source_gap_before=0,
            ),
        ]
    )
    plan = HybridLayoutPlan(
        content_source=content.source,
        layout_source=scan.source,
        pages=[
            HybridPagePlan(
                number=page.number,
                pdf_width=page.pdf_width,
                pdf_height=page.pdf_height,
                raster_width=page.width,
                raster_height=page.height,
                content_bbox=page.content_bbox,
                line_pitch=page.line_pitch,
                placements=[
                    placement for placement in placements if placement.page_number == page.number
                ],
            )
            for page in pages
        ],
    )

    payload = render_hybrid_docx(content, scan, plan, [])
    with zipfile.ZipFile(io.BytesIO(payload)) as package:
        document = ElementTree.fromstring(package.read("word/document.xml"))
        footer_names = sorted(name for name in package.namelist() if name.startswith("word/footer"))
        footers = [ElementTree.fromstring(package.read(name)) for name in footer_names]
    word = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    body_text = " ".join(node.text or "" for node in document.iter(f"{word}t"))
    footer_texts = [
        " ".join(node.text or "" for node in footer.iter(f"{word}t")) for footer in footers
    ]

    assert "Page 1/2" not in body_text
    assert "Страница 2/2" not in body_text
    assert "https://example.invalid/document" not in body_text
    assert "Repeated editable top banner." in body_text
    assert footer_texts == [
        "Open source: https://example.invalid/document Page 1/2",
        "Open source: https://example.invalid/document Страница 2/2",
    ]
