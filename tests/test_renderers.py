from __future__ import annotations

import base64
import io
import json
import zipfile
from pathlib import Path

import pytest

from docreconstruct.ir import BBox, Document, Element, ElementStyle, ElementType, Page
from docreconstruct.renderers import (
    DOCXRenderer,
    HTMLRenderer,
    JSONRenderer,
    MarkdownRenderer,
    registry,
    render,
)


def _document(*elements: Element) -> Document:
    return Document(
        id="test-document",
        pages=[Page(id="page-1", number=1, width=600, height=800, elements=list(elements))],
    )


def _png_bytes() -> bytes:
    from PIL import Image

    output = io.BytesIO()
    Image.new("RGB", (2, 2), "red").save(output, format="PNG")
    return output.getvalue()


def test_json_is_deterministic_and_round_trips_canonical_ir() -> None:
    document = _document(
        Element(
            id="paragraph-1",
            type=ElementType.PARAGRAPH,
            bbox=BBox(x0=10, y0=20, x1=200, y1=60),
            text="Exact text Ω",
        )
    )
    renderer = JSONRenderer()

    first = renderer.render(document)
    second = renderer.render(document)

    assert first == second
    assert Document.model_validate_json(first) == document
    assert json.loads(first)["pages"][0]["elements"][0]["text"] == "Exact text Ω"


def test_html_is_fixed_page_escaped_and_self_contained() -> None:
    dangerous = '<script>alert("x")</script> & exact'
    document = _document(
        Element(
            id="text-1",
            type=ElementType.PARAGRAPH,
            bbox=BBox(x0=10, y0=20, x1=250, y1=70),
            text=dangerous,
            style=ElementStyle(font_size=12, font_weight=700),
        ),
        Element(
            id="table-1",
            type=ElementType.TABLE,
            bbox=BBox(x0=10, y0=100, x1=400, y1=220),
            metadata={"rows": [["A", "B"], ["<one>", "two"]]},
        ),
        Element(
            id="image-1",
            type=ElementType.IMAGE,
            bbox=BBox(x0=10, y0=240, x1=110, y1=340),
            metadata={
                "image_data": base64.b64encode(_png_bytes()).decode(),
                "mime_type": "image/png",
            },
        ),
    )

    output = HTMLRenderer().render(document)

    assert "width:600px;height:800px" in output
    assert dangerous not in output
    assert "&lt;script&gt;" in output
    assert "&lt;one&gt;" in output
    assert "data:image/png;base64," in output
    assert "http://" not in output and "https://" not in output


def test_local_image_dereference_is_disabled_by_default(tmp_path: Path) -> None:
    image_path = tmp_path / "private.png"
    image_bytes = _png_bytes()
    image_path.write_bytes(image_bytes)
    encoded = base64.b64encode(image_bytes).decode("ascii")
    document = _document(
        Element(
            id="private-image",
            type=ElementType.IMAGE,
            bbox=BBox(x0=0, y0=0, x1=20, y1=20),
            metadata={"image_ref": str(image_path)},
        )
    )

    safe_html = HTMLRenderer().render(document)
    opted_in_html = HTMLRenderer(
        allow_local_files=True,
        local_file_root=tmp_path,
    ).render(document)

    assert encoded not in safe_html
    assert "dr-image-placeholder" in safe_html
    assert encoded in opted_in_html

    if DOCXRenderer.is_available():
        safe_docx = DOCXRenderer().render(document)
        with zipfile.ZipFile(io.BytesIO(safe_docx)) as archive:
            assert not any(name.startswith("word/media/") for name in archive.namelist())


@pytest.mark.skipif(not DOCXRenderer.is_available(), reason="python-docx is optional")
def test_docx_groups_native_spans_and_preserves_styled_runs() -> None:
    from docx import Document as WordDocument

    document = _document(
        Element(
            id="page-1-text-1-1-1",
            type=ElementType.TEXT,
            bbox=BBox(x0=10, y0=10, x1=45, y1=20),
            text="Hello",
            reading_order=0,
            style=ElementStyle(font_size=10, font_weight=400),
        ),
        Element(
            id="page-1-text-1-1-2",
            type=ElementType.TEXT,
            bbox=BBox(x0=50, y0=10, x1=85, y1=20),
            text="world",
            reading_order=1,
            style=ElementStyle(font_size=10, font_weight=700),
        ),
        Element(
            id="page-1-text-1-2-1",
            type=ElementType.TEXT,
            bbox=BBox(x0=10, y0=24, x1=45, y1=34),
            text="again",
            reading_order=2,
            style=ElementStyle(font_size=10, font_weight=400),
        ),
        Element(
            id="heading-1",
            type=ElementType.HEADING,
            bbox=BBox(x0=10, y0=60, x1=200, y1=90),
            text="Heading",
            reading_order=3,
            metadata={"level": 2},
        ),
    )

    output = DOCXRenderer().render(document)
    reconstructed = WordDocument(io.BytesIO(output))
    nonempty = [paragraph for paragraph in reconstructed.paragraphs if paragraph.text]

    assert [paragraph.text for paragraph in nonempty] == ["Hello world again", "Heading"]
    world_run = next(run for run in nonempty[0].runs if run.text == "world")
    assert world_run.bold is True
    assert nonempty[1].style.name == "Heading 2"


@pytest.mark.skipif(not DOCXRenderer.is_available(), reason="python-docx is optional")
def test_docx_preserves_source_page_geometry_margins_and_boundaries() -> None:
    from docx import Document as WordDocument
    from docx.enum.section import WD_SECTION

    a4_width, a4_height = 595.2756, 841.8898
    source = Document(
        id="two-page-pdf",
        pages=[
            Page(
                id="page-1",
                number=1,
                width=a4_width,
                height=a4_height,
                source_type="scanned",
                metadata={"provider": "native_pdf"},
                elements=[
                    Element(
                        id="scan-1",
                        type=ElementType.IMAGE,
                        bbox=BBox(x0=0, y0=72, x1=a4_width, y1=a4_height - 72),
                        metadata={"image_data": base64.b64encode(_png_bytes()).decode()},
                    )
                ],
            ),
            Page(
                id="page-2",
                number=2,
                width=612,
                height=792,
                metadata={
                    "coordinate_unit": "pt",
                    "page_margins": {"left": 36, "top": 54, "right": 45, "bottom": 63},
                },
                elements=[
                    Element(
                        id="paragraph-2",
                        type=ElementType.PARAGRAPH,
                        bbox=BBox(x0=36, y0=54, x1=567, y1=80),
                        text="Second page",
                    )
                ],
            ),
        ],
    )

    reconstructed = WordDocument(io.BytesIO(DOCXRenderer().render(source)))

    assert len(reconstructed.sections) == 2
    first, second = reconstructed.sections
    assert first.page_width.pt == pytest.approx(a4_width, abs=0.1)
    assert first.page_height.pt == pytest.approx(a4_height, abs=0.1)
    assert first.left_margin.pt == pytest.approx(0, abs=0.1)
    assert first.top_margin.pt == pytest.approx(72, abs=0.1)
    assert first.right_margin.pt == pytest.approx(0, abs=0.1)
    assert first.bottom_margin.pt == pytest.approx(72, abs=0.1)
    assert second.page_width.pt == pytest.approx(612, abs=0.1)
    assert second.page_height.pt == pytest.approx(792, abs=0.1)
    assert second.left_margin.pt == pytest.approx(36, abs=0.1)
    assert second.top_margin.pt == pytest.approx(54, abs=0.1)
    assert second.right_margin.pt == pytest.approx(45, abs=0.1)
    assert second.bottom_margin.pt == pytest.approx(63, abs=0.1)
    assert second.start_type == WD_SECTION.NEW_PAGE

    picture = reconstructed.inline_shapes[0]
    assert picture.width.inches == pytest.approx(a4_width / 72, abs=0.01)


@pytest.mark.skipif(not DOCXRenderer.is_available(), reason="python-docx is optional")
def test_docx_image_line_box_is_not_clipped_by_text_line_height() -> None:
    from docx import Document as WordDocument

    source = _document(
        Element(
            id="figure-with-text-line-height",
            type=ElementType.FIGURE,
            bbox=BBox(x0=10, y0=20, x1=210, y1=120),
            style=ElementStyle(line_height=1),
            metadata={"image_data": base64.b64encode(_png_bytes()).decode()},
        )
    )

    reconstructed = WordDocument(io.BytesIO(DOCXRenderer().render(source)))

    assert len(reconstructed.inline_shapes) == 1
    picture_paragraph = next(
        paragraph for paragraph in reconstructed.paragraphs if paragraph._p.xpath(".//w:drawing")
    )
    assert picture_paragraph.paragraph_format.line_spacing is None
    assert picture_paragraph.paragraph_format.line_spacing_rule is None


def test_markdown_and_render_facade(tmp_path: Path) -> None:
    document = _document(
        Element(
            id="title-1",
            type=ElementType.TITLE,
            bbox=BBox(x0=0, y0=0, x1=100, y1=20),
            text="A title",
        ),
        Element(
            id="table-1",
            type=ElementType.TABLE,
            bbox=BBox(x0=0, y0=30, x1=100, y1=80),
            metadata={"rows": [["Name", "Value"], ["A", "1"]]},
        ),
    )

    markdown = MarkdownRenderer().render(document)
    destination = render(document, tmp_path / "artifact.json")

    assert markdown.startswith("# A title")
    assert "| Name | Value |" in markdown
    assert destination.is_file()
    assert registry.get("json").format == "json"


def test_table_html_is_converted_to_native_rows_with_conservative_spans() -> None:
    document = _document(
        Element(
            id="table-html",
            type=ElementType.TABLE,
            bbox=BBox(x0=0, y0=0, x1=300, y1=100),
            metadata={
                "table_html": (
                    "<table><tr><th rowspan='2'>A</th><th colspan='2'>B</th></tr>"
                    "<tr><td>C</td><td>D<script>ignored()</script></td></tr></table>"
                )
            },
        )
    )

    markdown = MarkdownRenderer().render(document)
    html_output = HTMLRenderer().render(document)

    # GFM has no span syntax, so Markdown keeps the flattened grid.
    assert "| A | B |  |" in markdown
    assert "|  | C | D |" in markdown
    assert "ignored" not in markdown
    # HTML can express the real shape, and flattening it produced boxes the
    # source never had: an empty header cell beside B, an empty cell under A.
    assert '<th rowspan="2">A</th><th colspan="2">B</th>' in html_output
    assert "<tr><td>C</td><td>D</td></tr>" in html_output
    assert "<th></th>" not in html_output


def _rotated_document(rotation: float) -> Document:
    return Document(
        id="rotated",
        pages=[
            Page(
                id="page-1",
                number=1,
                width=612,
                height=792,
                rotation=rotation,
                metadata={"coordinate_unit": "pt"},
                elements=[
                    Element(
                        id="heading",
                        type=ElementType.TEXT,
                        bbox=BBox(x0=72, y0=100, x1=272, y1=120),
                        text="Rotated heading",
                    )
                ],
            )
        ],
    )


@pytest.mark.parametrize(
    ("rotation", "width", "height", "box"),
    [
        (0.0, 612.0, 792.0, (72.0, 100.0)),
        (90.0, 792.0, 612.0, (792.0 - 120.0, 72.0)),
        (180.0, 612.0, 792.0, (612.0 - 272.0, 792.0 - 120.0)),
        (270.0, 792.0, 612.0, (100.0, 612.0 - 272.0)),
    ],
)
def test_page_rotation_reaches_both_plain_renderers(
    rotation: float, width: float, height: float, box: tuple[float, float]
) -> None:
    """Providers store unrotated boxes plus ``Page.rotation``.

    Both plain renderers read ``page.width``/``page.height`` directly, so a
    landscape scan stored the usual way (portrait MediaBox plus ``/Rotate 90``)
    came out as a portrait page with every element in the unrotated frame.
    """

    document = _rotated_document(rotation)

    from docx import Document as WordDocument

    section = WordDocument(io.BytesIO(DOCXRenderer().render(document))).sections[0]
    assert (round(section.page_width.pt), round(section.page_height.pt)) == (
        round(width),
        round(height),
    )
    assert (section.orientation == 1) is (width > height)

    markup = HTMLRenderer().render(document)
    assert f"width:{width:g}px;height:{height:g}px" in markup
    assert f"left:{box[0]:g}px;top:{box[1]:g}px" in markup


def test_unrecognized_page_rotation_leaves_geometry_alone() -> None:
    markup = HTMLRenderer().render(_rotated_document(37.0))

    assert "width:612px;height:792px" in markup
    assert "left:72px;top:100px" in markup


def test_webp_picture_is_transcoded_instead_of_aborting_the_render(tmp_path: Path) -> None:
    """python-docx rejects WEBP with a message-less UnrecognizedImageError.

    It escaped the renderer's ``(OSError, ValueError)`` guard entirely, so a
    single unsupported picture killed the whole document with a bare traceback.
    """

    from PIL import Image

    source = tmp_path / "scan.webp"
    Image.new("RGB", (120, 80), "steelblue").save(source, "WEBP")
    document = Document(
        id="webp",
        pages=[
            Page(
                id="page-1",
                number=1,
                width=200,
                height=200,
                elements=[
                    Element(
                        id="picture",
                        type=ElementType.IMAGE,
                        bbox=BBox(x0=0, y0=0, x1=120, y1=80),
                        metadata={"image": {"path": str(source), "mime_type": "image/webp"}},
                    )
                ],
            )
        ],
    )

    blob = DOCXRenderer(allow_local_files=True, local_file_root=str(tmp_path)).render(document)

    with zipfile.ZipFile(io.BytesIO(blob)) as archive:
        media = [name for name in archive.namelist() if name.startswith("word/media/")]
    assert media, "the picture should survive as an embedded part"


def test_local_image_path_is_accepted_by_the_docx_renderer(tmp_path: Path) -> None:
    """``add_picture`` treats a non-``str`` argument as an open stream.

    Passing the ``Path`` returned by ``_image_source`` raised AttributeError
    for every format, not just the ones Word cannot parse.
    """

    source = tmp_path / "scan.png"
    source.write_bytes(_png_bytes())
    document = Document(
        id="local",
        pages=[
            Page(
                id="page-1",
                number=1,
                width=200,
                height=200,
                elements=[
                    Element(
                        id="picture",
                        type=ElementType.IMAGE,
                        bbox=BBox(x0=0, y0=0, x1=40, y1=40),
                        metadata={"image": {"path": str(source), "mime_type": "image/png"}},
                    )
                ],
            )
        ],
    )

    blob = DOCXRenderer(allow_local_files=True, local_file_root=str(tmp_path)).render(document)

    with zipfile.ZipFile(io.BytesIO(blob)) as archive:
        assert any(name.startswith("word/media/") for name in archive.namelist())


def test_undecodable_picture_raises_a_named_renderer_error(tmp_path: Path) -> None:
    from docreconstruct.renderers.base import RendererError

    source = tmp_path / "broken.png"
    source.write_bytes(b"not an image at all")
    document = Document(
        id="broken",
        pages=[
            Page(
                id="page-1",
                number=1,
                width=200,
                height=200,
                elements=[
                    Element(
                        id="picture",
                        type=ElementType.IMAGE,
                        bbox=BBox(x0=0, y0=0, x1=40, y1=40),
                        metadata={"image": {"path": str(source), "mime_type": "image/png"}},
                    )
                ],
            )
        ],
    )

    with pytest.raises(RendererError, match="UnrecognizedImageError"):
        DOCXRenderer(allow_local_files=True, local_file_root=str(tmp_path)).render(document)


@pytest.mark.parametrize(
    ("label", "text", "expected"),
    [
        ("form feed", "page\x0cbreak", "pagebreak"),
        ("nul", "a\x00b", "ab"),
        ("vertical tab", "a\x0bb", "ab"),
        ("bell", "a\x07b", "ab"),
        ("tab is legal", "a\tb", "a\tb"),
    ],
)
def test_control_characters_do_not_abort_the_docx_render(
    label: str, text: str, expected: str
) -> None:
    """XML 1.0 cannot represent these at all, so lxml rejects the whole tree.

    A single form feed from OCR of a page-break glyph used to cost the entire
    document, with an uncaught ValueError rather than a renderer error. HTML
    and Markdown were never affected.
    """

    from docx import Document as WordDocument

    document = Document(
        id="control",
        pages=[
            Page(
                id="page-1",
                number=1,
                width=600,
                height=800,
                elements=[
                    Element(
                        id="before",
                        type=ElementType.TEXT,
                        bbox=BBox(x0=0, y0=0, x1=100, y1=20),
                        text="clean before",
                    ),
                    Element(
                        id="dirty",
                        type=ElementType.TEXT,
                        bbox=BBox(x0=0, y0=200, x1=100, y1=220),
                        text=text,
                    ),
                    Element(
                        id="after",
                        type=ElementType.TEXT,
                        bbox=BBox(x0=0, y0=400, x1=100, y1=420),
                        text="clean after",
                    ),
                ],
            )
        ],
    )

    rendered = WordDocument(io.BytesIO(DOCXRenderer().render(document)))
    body = "\n".join(paragraph.text for paragraph in rendered.paragraphs)

    assert expected in body
    # The neighbouring elements must survive too.
    assert "clean before" in body and "clean after" in body


def test_control_characters_survive_the_hybrid_docx_text_path() -> None:
    """The hybrid renderer is the primary reconstruction path and had the same hole."""

    from docx import Document as WordDocument

    from docreconstruct.reconstruction.hybrid_docx import _add_rich_text, _math_runs

    document = WordDocument()
    paragraph = document.add_paragraph()
    _add_rich_text(paragraph, "page\x0cbreak", size=11.0)
    assert paragraph.text == "pagebreak"

    math_paragraph = document.add_paragraph()
    _math_runs(math_paragraph, "x\x0c_{1}", size=11.0)
    assert math_paragraph.text == "x1"
