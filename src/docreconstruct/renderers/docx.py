"""Basic semantic Word renderer using the optional ``python-docx`` package."""

from __future__ import annotations

import base64
import binascii
import importlib.util
import io
import re
from pathlib import Path
from typing import Any

from ._utils import (
    allowed_local_path,
    bbox_tuple,
    element_metadata,
    element_style,
    element_text,
    element_type,
    enum_text,
    finite_number,
    mapping,
    ordered_elements,
    pages,
    table_rows,
    value,
)
from .base import OptionalDependencyError, Renderer, RendererError

_HEX_COLOR = re.compile(r"^#?([0-9a-fA-F]{6})$")
_NATIVE_SPAN_ID = re.compile(r"-text-(\d+)-(\d+)-(\d+)$")


def _require_docx() -> dict[str, Any]:
    try:
        from docx import Document as WordDocument
        from docx.enum.section import WD_SECTION
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt, RGBColor
    except (ImportError, ModuleNotFoundError) as exc:
        raise OptionalDependencyError(
            "DOCX rendering requires python-docx. Install it with "
            "`pip install 'docreconstruct[docx]'` or `pip install python-docx`."
        ) from exc
    return {
        "Document": WordDocument,
        "WD_SECTION": WD_SECTION,
        "WD_ALIGN_PARAGRAPH": WD_ALIGN_PARAGRAPH,
        "Inches": Inches,
        "Pt": Pt,
        "RGBColor": RGBColor,
    }


def _page_units_per_inch(page: Any, fallback_dpi: float) -> float:
    """Return the number of source-page coordinate units in one inch.

    Canonical page coordinates deliberately retain their provider's coordinate
    system. Native PDF coordinates are points, while raster/OCR coordinates are
    normally pixels. Explicit unit or DPI metadata wins over provider inference.
    """

    metadata = mapping(value(page, "metadata", None))
    raw_unit = metadata.get("coordinate_unit", metadata.get("unit", ""))
    unit = enum_text(raw_unit).strip().lower().replace("_", " ")
    unit_scales = {
        "pt": 72.0,
        "point": 72.0,
        "points": 72.0,
        "in": 1.0,
        "inch": 1.0,
        "inches": 1.0,
        "mm": 25.4,
        "millimeter": 25.4,
        "millimeters": 25.4,
        "cm": 2.54,
        "centimeter": 2.54,
        "centimeters": 2.54,
    }
    if unit in unit_scales:
        return unit_scales[unit]

    raw_dpi = metadata.get("dpi", metadata.get("image_dpi"))
    if isinstance(raw_dpi, (list, tuple)):
        raw_dpi = raw_dpi[0] if raw_dpi else None
    dpi = finite_number(raw_dpi, 0.0)
    if dpi > 0:
        return dpi

    provider = enum_text(metadata.get("provider", "")).strip().lower()
    if provider in {"native_pdf", "native pdf", "pdf"}:
        return 72.0
    return fallback_dpi


def _explicit_page_margins(page: Any) -> dict[str, float] | None:
    metadata = mapping(value(page, "metadata", None))
    raw = metadata.get("page_margins", metadata.get("margins"))
    if not isinstance(raw, dict):
        return None
    margins: dict[str, float] = {}
    for side in ("left", "top", "right", "bottom"):
        if raw.get(side) is None:
            return None
        amount = finite_number(raw[side], -1.0)
        if amount < 0:
            return None
        margins[side] = amount
    return margins


def _scan_page_margins(page: Any) -> dict[str, float] | None:
    """Infer scan whitespace from the union of positioned source elements."""

    source_type = enum_text(value(page, "source_type", "")).strip().lower()
    if source_type not in {"scanned", "image"}:
        return None
    width = max(1.0, finite_number(value(page, "width", 1.0), 1.0))
    height = max(1.0, finite_number(value(page, "height", 1.0), 1.0))
    boxes = [bbox_tuple(element) for element in ordered_elements(page)]
    boxes = [box for box in boxes if box[2] > box[0] and box[3] > box[1]]
    if not boxes:
        return None
    left = max(0.0, min(width, min(box[0] for box in boxes)))
    top = max(0.0, min(height, min(box[1] for box in boxes)))
    right_edge = max(0.0, min(width, max(box[2] for box in boxes)))
    bottom_edge = max(0.0, min(height, max(box[3] for box in boxes)))
    return {
        "left": left,
        "top": top,
        "right": max(0.0, width - right_edge),
        "bottom": max(0.0, height - bottom_edge),
    }


def _configure_section(
    section: Any,
    page: Any,
    api: dict[str, Any],
    *,
    fallback_dpi: float,
) -> float:
    """Apply source page geometry and return its coordinate units per inch."""

    units_per_inch = _page_units_per_inch(page, fallback_dpi)
    page_width = max(1.0, finite_number(value(page, "width", 1.0), 1.0))
    page_height = max(1.0, finite_number(value(page, "height", 1.0), 1.0))
    section.page_width = api["Inches"](page_width / units_per_inch)
    section.page_height = api["Inches"](page_height / units_per_inch)

    margins = _explicit_page_margins(page) or _scan_page_margins(page)
    if margins is not None:
        for side, amount in margins.items():
            setattr(section, f"{side}_margin", api["Inches"](amount / units_per_inch))
    return units_per_inch


def _image_source(
    element: Any,
    *,
    allow_local_files: bool = False,
    local_file_root: str | Path | None = None,
) -> io.BytesIO | Path | None:
    metadata = element_metadata(element)
    nested = metadata.get("image")
    image = dict(nested) if isinstance(nested, dict) else {}
    raw = image.get(
        "bytes",
        image.get("data", metadata.get("image_bytes", metadata.get("image_data"))),
    )
    if isinstance(raw, str):
        if raw.startswith("data:") and "," in raw:
            raw = raw.split(",", 1)[1]
        try:
            raw = base64.b64decode(raw, validate=True)
        except (ValueError, binascii.Error):
            raw = None
    if isinstance(raw, (bytes, bytearray, memoryview)):
        return io.BytesIO(bytes(raw))
    for key in ("path", "src", "image_ref"):
        path = allowed_local_path(
            image.get(key, metadata.get(key)),
            allow_local_files=allow_local_files,
            local_file_root=local_file_root,
        )
        if path is not None:
            return path
    return None


def _apply_text_style(run: Any, element: Any, api: dict[str, Any]) -> None:
    style = element_style(element)
    font = run.font
    if style.get("font_family"):
        font.name = str(style["font_family"])
    if style.get("font_size") is not None:
        size = finite_number(style["font_size"])
        if size > 0:
            font.size = api["Pt"](size)
    weight = style.get("font_weight")
    if weight is not None:
        if isinstance(weight, str):
            font.bold = weight.strip().lower() in {"bold", "semibold", "demibold", "heavy"}
        else:
            font.bold = finite_number(weight) >= 600
    if style.get("italic") is not None:
        font.italic = bool(style["italic"])
    if style.get("underline") is not None:
        font.underline = bool(style["underline"])
    color = style.get("color")
    if color and (match := _HEX_COLOR.fullmatch(str(color).strip())):
        font.color.rgb = api["RGBColor"].from_string(match.group(1).upper())


def _apply_paragraph_style(paragraph: Any, element: Any, api: dict[str, Any]) -> None:
    style = element_style(element)
    alignment = enum_text(style.get("alignment", "")).strip().lower()
    enum = api["WD_ALIGN_PARAGRAPH"]
    alignments = {
        "left": enum.LEFT,
        "center": enum.CENTER,
        "right": enum.RIGHT,
        "justify": enum.JUSTIFY,
    }
    if alignment in alignments:
        paragraph.alignment = alignments[alignment]
    if style.get("line_height") is not None:
        height = finite_number(style["line_height"])
        if height > 0:
            paragraph.paragraph_format.line_spacing = api["Pt"](height)
    metadata = element_metadata(element)
    for source, target in (
        ("space_before", "space_before"),
        ("space_after", "space_after"),
        ("left_indent", "left_indent"),
        ("right_indent", "right_indent"),
        ("first_line_indent", "first_line_indent"),
    ):
        if metadata.get(source) is not None:
            amount = finite_number(metadata[source])
            if amount >= 0 or source == "first_line_indent":
                setattr(paragraph.paragraph_format, target, api["Pt"](amount))


def _paragraph_for(document: Any, element: Any) -> Any:
    kind = element_type(element)
    metadata = element_metadata(element)
    if kind == "title":
        return document.add_paragraph(style="Title")
    if kind == "heading":
        level = int(max(1, min(9, finite_number(metadata.get("level", 2), 2.0))))
        return document.add_heading(level=level)
    if kind == "list_item":
        list_style = "List Number" if metadata.get("ordered") else "List Bullet"
        try:
            return document.add_paragraph(style=list_style)
        except KeyError:
            return document.add_paragraph()
    return document.add_paragraph()


def _native_span_key(element: Any) -> tuple[int, int, int] | None:
    match = _NATIVE_SPAN_ID.search(str(value(element, "id", "")))
    return tuple(int(group) for group in match.groups()) if match else None  # type: ignore[return-value]


def _same_line(left: Any, right: Any, page: Any) -> bool:
    left_key, right_key = _native_span_key(left), _native_span_key(right)
    if left_key is not None and right_key is not None:
        return left_key[:2] == right_key[:2]
    ax0, ay0, ax1, ay1 = bbox_tuple(left)
    bx0, by0, bx1, by1 = bbox_tuple(right)
    left_height = max(1.0, ay1 - ay0)
    right_height = max(1.0, by1 - by0)
    overlap = max(0.0, min(ay1, by1) - max(ay0, by0))
    center_delta = abs((ay0 + ay1 - by0 - by1) / 2)
    horizontal_gap = bx0 - ax1
    page_width = max(1.0, finite_number(value(page, "width", 1.0), 1.0))
    vertically_aligned = overlap / min(
        left_height, right_height
    ) >= 0.45 or center_delta <= 0.35 * max(left_height, right_height)
    return (
        vertically_aligned
        and bx0 >= ax0
        and horizontal_gap >= -0.15 * min(max(1.0, ax1 - ax0), max(1.0, bx1 - bx0))
        and horizontal_gap <= max(3 * max(left_height, right_height), 0.06 * page_width)
    )


def _styles_compatible(left: Any, right: Any) -> bool:
    left_style, right_style = element_style(left), element_style(right)
    for key in ("font_family", "italic", "alignment"):
        left_value, right_value = left_style.get(key), right_style.get(key)
        if left_value is not None and right_value is not None and left_value != right_value:
            return False
    left_size = finite_number(left_style.get("font_size"), 0.0)
    right_size = finite_number(right_style.get("font_size"), 0.0)
    if left_size and right_size and min(left_size, right_size) / max(left_size, right_size) < 0.85:
        return False
    left_weight = finite_number(left_style.get("font_weight"), 400.0) >= 600
    right_weight = finite_number(right_style.get("font_weight"), 400.0) >= 600
    return left_weight == right_weight


def _line_bounds(line: list[Any]) -> tuple[float, float, float, float]:
    boxes = [bbox_tuple(element) for element in line]
    return (
        min(box[0] for box in boxes),
        min(box[1] for box in boxes),
        max(box[2] for box in boxes),
        max(box[3] for box in boxes),
    )


def _adjacent_lines(left: list[Any], right: list[Any], page: Any) -> bool:
    left_key, right_key = _native_span_key(left[0]), _native_span_key(right[0])
    if (
        left_key is not None
        and right_key is not None
        and (left_key[0] != right_key[0] or right_key[1] != left_key[1] + 1)
    ):
        return False
    ax0, ay0, ax1, ay1 = _line_bounds(left)
    bx0, by0, bx1, by1 = _line_bounds(right)
    line_height = max(1.0, ay1 - ay0, by1 - by0)
    gap = by0 - ay1
    page_width = max(1.0, finite_number(value(page, "width", 1.0), 1.0))
    left_aligned = abs(ax0 - bx0) <= max(5.0, 0.6 * line_height)
    not_a_new_column = abs(ax0 - bx0) <= 0.08 * page_width
    return (
        -0.2 * line_height <= gap <= 1.35 * line_height
        and left_aligned
        and not_a_new_column
        and _styles_compatible(left[0], right[0])
    )


def _flow_groups(spans: list[Any], page: Any) -> list[list[list[Any]]]:
    """Conservatively group provider spans into lines and flowing paragraphs."""

    lines: list[list[Any]] = []
    for span in spans:
        if lines and _same_line(lines[-1][-1], span, page):
            lines[-1].append(span)
        else:
            lines.append([span])
    paragraphs: list[list[list[Any]]] = []
    for line in lines:
        if paragraphs and _adjacent_lines(paragraphs[-1][-1], line, page):
            paragraphs[-1].append(line)
        else:
            paragraphs.append([line])
    return paragraphs


def _span_separator(left: Any, right: Any, *, new_line: bool) -> str:
    left_text, right_text = element_text(left), element_text(right)
    if not left_text or not right_text or left_text[-1].isspace() or right_text[0].isspace():
        return ""
    if right_text[0] in ",.;:!?)]}%" or left_text[-1] in "([{$/":
        return ""
    if new_line:
        return "" if left_text.endswith(("-", "‐", "‑")) else " "
    _, _, left_x1, _ = bbox_tuple(left)
    right_x0, _, _, _ = bbox_tuple(right)
    gap = right_x0 - left_x1
    left_width = max(0.0, bbox_tuple(left)[2] - bbox_tuple(left)[0])
    visible_characters = max(1, len(left_text.strip()))
    average_character_width = left_width / visible_characters if left_width else 0.0
    return " " if gap > max(0.75, average_character_width * 0.25) else ""


def _render_text_flow(document: Any, spans: list[Any], page: Any, api: dict[str, Any]) -> None:
    for lines in _flow_groups(spans, page):
        first = lines[0][0]
        paragraph = document.add_paragraph()
        _apply_paragraph_style(paragraph, first, api)
        previous: Any | None = None
        for line_index, line in enumerate(lines):
            for span_index, span in enumerate(line):
                if previous is not None:
                    separator = _span_separator(
                        previous,
                        span,
                        new_line=line_index > 0 and span_index == 0,
                    )
                    if separator:
                        separator_run = paragraph.add_run(separator)
                        _apply_text_style(separator_run, previous, api)
                run = paragraph.add_run(element_text(span))
                _apply_text_style(run, span, api)
                previous = span


class DOCXRenderer(Renderer[bytes]):
    """Prefer native paragraphs, headings, tables, and pictures over screenshots."""

    format = "docx"
    extension = ".docx"
    media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    def __init__(
        self,
        *,
        image_dpi: float = 96.0,
        allow_local_files: bool = False,
        local_file_root: str | Path | None = None,
    ) -> None:
        self.image_dpi = max(1.0, finite_number(image_dpi, 96.0))
        self.allow_local_files = bool(allow_local_files)
        self.local_file_root = Path(local_file_root) if local_file_root is not None else None

    @classmethod
    def is_available(cls) -> bool:
        return importlib.util.find_spec("docx") is not None

    def render(self, source_document: Any) -> bytes:
        api = _require_docx()
        document = api["Document"]()
        metadata = mapping(value(source_document, "metadata", None))
        if metadata.get("title"):
            document.core_properties.title = str(metadata["title"])
        if metadata.get("author"):
            document.core_properties.author = str(metadata["author"])

        source_pages = pages(source_document)
        for page_index, page in enumerate(source_pages):
            if page_index == 0:
                section = document.sections[0]
            else:
                section = document.add_section(api["WD_SECTION"].NEW_PAGE)
            page_units_per_inch = _configure_section(
                section,
                page,
                api,
                fallback_dpi=self.image_dpi,
            )
            page_elements = ordered_elements(page)
            element_index = 0
            while element_index < len(page_elements):
                element = page_elements[element_index]
                kind = element_type(element)
                if kind == "text":
                    end = element_index + 1
                    while end < len(page_elements) and element_type(page_elements[end]) == "text":
                        end += 1
                    _render_text_flow(document, page_elements[element_index:end], page, api)
                    element_index = end
                    continue
                if kind == "table" and (rows := table_rows(element)):
                    columns = max((len(row) for row in rows), default=0)
                    if columns:
                        table = document.add_table(rows=len(rows), cols=columns)
                        table.style = "Table Grid"
                        for row_index, row in enumerate(rows):
                            for column_index, text in enumerate(row):
                                cell = table.cell(row_index, column_index)
                                cell.text = text
                        element_index += 1
                        continue
                if kind in {"image", "figure", "chart", "signature", "stamp"}:
                    source = _image_source(
                        element,
                        allow_local_files=self.allow_local_files,
                        local_file_root=self.local_file_root,
                    )
                    if source is not None:
                        paragraph = document.add_paragraph()
                        _apply_paragraph_style(paragraph, element, api)
                        paragraph.paragraph_format.space_before = api["Pt"](0)
                        paragraph.paragraph_format.space_after = api["Pt"](0)
                        # Exact line-height metadata is valid for text but clips
                        # inline pictures in Word/LibreOffice.  Let the picture
                        # establish the line box while retaining alignment.
                        paragraph.paragraph_format.line_spacing = None
                        paragraph.paragraph_format.line_spacing_rule = None
                        run = paragraph.add_run()
                        left, _, right, _ = bbox_tuple(element)
                        width_inches = max(0.0, right - left) / page_units_per_inch
                        kwargs = {"width": api["Inches"](width_inches)} if width_inches > 0 else {}
                        try:
                            run.add_picture(source, **kwargs)
                            element_index += 1
                            continue
                        except (OSError, ValueError) as exc:
                            element_id = value(element, "id", "<unknown>")
                            raise RendererError(
                                f"could not add image for element {element_id!r}: {exc}"
                            ) from exc
                    # Preserve the visible/alternative text when image bytes are
                    # unavailable; never fabricate a visual replacement.
                    text = element_text(element)
                    if not text:
                        text = str(element_metadata(element).get("alt", ""))
                    if not text:
                        element_index += 1
                        continue
                else:
                    text = element_text(element)

                paragraph = _paragraph_for(document, element)
                _apply_paragraph_style(paragraph, element, api)
                run = paragraph.add_run(text)
                _apply_text_style(run, element, api)
                element_index += 1

        output = io.BytesIO()
        document.save(output)
        return output.getvalue()


DocxRenderer = DOCXRenderer
