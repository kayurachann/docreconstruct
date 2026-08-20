"""Editable DOCX rendering for renderer-neutral hybrid layout plans."""

from __future__ import annotations

import io
import math
import re
import statistics
import textwrap
from collections.abc import Mapping, Sequence
from dataclasses import dataclass
from pathlib import Path
from typing import Any, cast

from docx import Document as WordDocument
from docx.document import Document as WordDocumentType
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import (
    WD_ALIGN_PARAGRAPH,
    WD_LINE_SPACING,
    WD_TAB_ALIGNMENT,
    WD_TAB_LEADER,
)
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from PIL import Image

from docreconstruct.ir import ElementStyle, TextAlignment
from docreconstruct.reconstruction.asset_matching import (
    AssetMatch,
    resolve_markdown_asset,
)
from docreconstruct.reconstruction.hybrid_planner import (
    HybridBlockPlacement,
    HybridLayoutPlan,
    apply_page_vertical_fit_budget,
    build_page_vertical_fit_budget,
    contains_tall_inline_math,
    equation_layout_units,
    source_row_reading_order,
)
from docreconstruct.reconstruction.markdown_content import (
    MarkdownBlock,
    MarkdownBlockKind,
    MarkdownContent,
)
from docreconstruct.reconstruction.markdown_inline import parse_markdown_inline
from docreconstruct.reconstruction.math_omml import append_omml, equation_row_count
from docreconstruct.reconstruction.scan_layout import PixelBox, ScanDocumentLayout, ScanRegionKind

_FONT = "Times New Roman"
_HAN_PATTERN = re.compile(r"[\u3400-\u4dbf\u4e00-\u9fff\uf900-\ufaff]")
_KANA_PATTERN = re.compile(r"[\u3040-\u30ff\u31f0-\u31ff]")
_HANGUL_PATTERN = re.compile(r"[\u1100-\u11ff\u3130-\u318f\uac00-\ud7af]")


@dataclass(frozen=True)
class _PlacementGeometry:
    """Physical geometry retained from one source block placement."""

    gap_before: float
    height: float
    left_indent: float
    right_indent: float
    row_heights: tuple[float, ...]
    row_gaps: tuple[float, ...]

    @property
    def row_spacing(self) -> int | None:
        if not self.row_gaps:
            return None
        return max(1, round(statistics.median(self.row_gaps)))


def _east_asia_profile(text: str) -> tuple[str, str] | None:
    """Return a Word font/language pair for the East Asian script in a run."""

    if _KANA_PATTERN.search(text):
        return "Yu Mincho", "ja-JP"
    if _HANGUL_PATTERN.search(text):
        return "Batang", "ko-KR"
    if _HAN_PATTERN.search(text):
        return "SimSun", "zh-CN"
    return None


def _east_asia_heading_profile(text: str) -> tuple[str, str] | None:
    """Return the restrained sans/Hei counterpart for a CJK heading."""

    if _KANA_PATTERN.search(text):
        return "Yu Gothic", "ja-JP"
    if _HANGUL_PATTERN.search(text):
        return "Malgun Gothic", "ko-KR"
    if _HAN_PATTERN.search(text):
        return "SimHei", "zh-CN"
    return None


def _page_render_content_bbox(page: Any) -> PixelBox:
    """Return the OCR-free printable frame selected by scan analysis."""

    value = page.metadata.get("render_content_bbox")
    if isinstance(value, dict):
        try:
            return PixelBox.model_validate(value)
        except ValueError:
            pass
    return page.content_bbox


def _document_paper_color(layout: ScanDocumentLayout) -> str:
    """Return the median native page background colour across source pages."""

    channels: list[tuple[int, int, int]] = []
    for page in layout.pages:
        value = page.metadata.get("paper_color")
        if isinstance(value, str) and re.fullmatch(r"[0-9A-Fa-f]{6}", value):
            channels.append(
                (
                    int(value[0:2], 16),
                    int(value[2:4], 16),
                    int(value[4:6], 16),
                )
            )
    if not channels:
        return "FFFFFF"
    return "".join(
        f"{round(statistics.median(channel[index] for channel in channels)):02X}"
        for index in range(3)
    )


def _set_native_page_background(document: WordDocumentType, color: str) -> None:
    """Set editable Word page colour; never insert a full-page scan image."""

    if color.upper() == "FFFFFF":
        return
    root = document._element
    background = root.find(qn("w:background"))
    if background is None:
        background = OxmlElement("w:background")
        root.insert(0, background)
    background.set(qn("w:color"), color.upper())


def _placement_geometry_points(
    placement: HybridBlockPlacement | None,
    layout: ScanDocumentLayout | None,
    *,
    available_width_points: float | None = None,
) -> _PlacementGeometry | None:
    """Return source-derived block geometry in physical points.

    Page-wide flows use the scan content box as their horizontal coordinate
    system.  Narrow layout cells use the detected source column when one is
    available, preventing a right-column block from inheriting a page-wide
    left indent inside its already-positioned cell.
    """

    if placement is None or layout is None or placement.source_bbox is None:
        return None
    page = layout.pages[placement.page_number - 1]
    horizontal_scale = page.pdf_width / page.width
    vertical_scale = (
        page.pdf_height / page.height
        if page.metadata.get("source_kind") == "image"
        else horizontal_scale
    )
    render_frame = _page_render_content_bbox(page)
    horizontal_origin = float(render_frame.x0)
    horizontal_end = float(render_frame.x1)
    full_width_points = (horizontal_end - horizontal_origin) * horizontal_scale
    horizontal_points_per_pixel = horizontal_scale
    if available_width_points is not None and available_width_points < full_width_points * 0.82:
        column_boxes = page.metadata.get("column_boxes")
        center = (placement.source_bbox.x0 + placement.source_bbox.x1) / 2.0
        candidates = (
            [box for box in column_boxes if isinstance(box, list) and len(box) == 4]
            if isinstance(column_boxes, list)
            else []
        )
        if candidates:
            selected = min(
                candidates,
                key=lambda box: (
                    0.0
                    if float(box[0]) <= center <= float(box[2])
                    else min(abs(center - float(box[0])), abs(center - float(box[2])))
                ),
            )
            horizontal_origin = float(selected[0])
            horizontal_end = float(selected[2])
            horizontal_points_per_pixel = available_width_points / max(
                1.0,
                horizontal_end - horizontal_origin,
            )
    gap = float(placement.source_gap_before or 0) * vertical_scale
    height = placement.source_bbox.height * vertical_scale
    left = max(
        0.0,
        (placement.source_bbox.x0 - horizontal_origin) * horizontal_points_per_pixel,
    )
    right = max(
        0.0,
        (horizontal_end - placement.source_bbox.x1) * horizontal_points_per_pixel,
    )
    if available_width_points is not None and left + right >= available_width_points:
        normalization = max(1.0, left + right)
        maximum_indents = max(0.0, available_width_points - 1.0)
        left *= maximum_indents / normalization
        right *= maximum_indents / normalization
    rows = source_row_reading_order(page, placement.source_rows)
    row_heights = tuple(row.height * vertical_scale for row in rows)
    row_gaps = tuple(
        max(0.0, following.y0 - previous.y1) * vertical_scale
        for previous, following in zip(rows, rows[1:], strict=False)
    )
    return _PlacementGeometry(
        gap_before=gap,
        height=height,
        left_indent=left,
        right_indent=right,
        row_heights=row_heights,
        row_gaps=row_gaps,
    )


def _set_font(
    run: Any,
    size: float,
    *,
    bold: bool | None = None,
    italic: bool | None = None,
    east_asia_heading: bool = False,
) -> None:
    run.font.name = _FONT
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    properties = run._element.get_or_add_rPr()
    fonts = properties.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        properties.insert(0, fonts)
    for attribute in ("ascii", "hAnsi", "cs"):
        fonts.set(qn(f"w:{attribute}"), _FONT)
    text = str(run.text or "")
    east_asia = _east_asia_heading_profile(text) if east_asia_heading else _east_asia_profile(text)
    if east_asia is not None:
        east_asia_font, east_asia_language = east_asia
        fonts.set(qn("w:eastAsia"), east_asia_font)
        language = properties.find(qn("w:lang"))
        if language is None:
            language = OxmlElement("w:lang")
            properties.append(language)
        language.set(qn("w:eastAsia"), east_asia_language)


def _apply_evidence_style(
    paragraph: Paragraph,
    style: ElementStyle | None,
    *,
    allow_alignment: bool,
) -> None:
    """Apply high-confidence JSON style hints without changing Markdown text."""

    if style is None:
        return
    if allow_alignment and style.alignment is not None:
        alignment = {
            TextAlignment.LEFT: WD_ALIGN_PARAGRAPH.LEFT,
            TextAlignment.CENTER: WD_ALIGN_PARAGRAPH.CENTER,
            TextAlignment.RIGHT: WD_ALIGN_PARAGRAPH.RIGHT,
            TextAlignment.JUSTIFY: WD_ALIGN_PARAGRAPH.JUSTIFY,
        }.get(style.alignment)
        if alignment is not None:
            paragraph.alignment = alignment
    color = style.color.removeprefix("#") if isinstance(style.color, str) else None
    valid_color = color.upper() if color and re.fullmatch(r"[0-9A-Fa-f]{6}", color) else None
    family = style.font_family.strip() if style.font_family else None
    for run in paragraph.runs:
        if style.font_weight is not None and style.font_weight >= 600:
            run.bold = True
        if style.italic is True:
            run.italic = True
        if style.underline is True:
            run.underline = True
        if valid_color is not None:
            run.font.color.rgb = RGBColor.from_string(valid_color)
        if family:
            run.font.name = family
            properties = run._element.get_or_add_rPr()
            fonts = properties.rFonts
            if fonts is None:
                fonts = OxmlElement("w:rFonts")
                properties.insert(0, fonts)
            for attribute in ("ascii", "hAnsi", "cs"):
                fonts.set(qn(f"w:{attribute}"), family)


def _format_paragraph(
    paragraph: Paragraph,
    *,
    size: float,
    line_height: float,
    alignment: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.JUSTIFY,
    before: float = 0.0,
    after: float = 0.0,
    left_indent: float = 0.0,
    right_indent: float | None = None,
    keep_next: bool = False,
) -> Paragraph:
    formatting = paragraph.paragraph_format
    formatting.space_before = Pt(before)
    formatting.space_after = Pt(after)
    formatting.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    formatting.line_spacing = Pt(line_height)
    formatting.left_indent = Pt(left_indent)
    if right_indent is not None:
        formatting.right_indent = Pt(right_indent)
    formatting.keep_with_next = keep_next
    paragraph.alignment = alignment
    for run in paragraph.runs:
        _set_font(run, size)
    return paragraph


def _set_paragraph_mark_math_format(paragraph: Paragraph, size: float) -> None:
    """Pin the display-math paragraph fallback to the expression base size."""

    paragraph_properties = paragraph._p.get_or_add_pPr()
    properties = paragraph_properties.find(qn("w:rPr"))
    if properties is None:
        properties = OxmlElement("w:rPr")
        section_properties = paragraph_properties.find(qn("w:sectPr"))
        if section_properties is None:
            paragraph_properties.append(properties)
        else:
            paragraph_properties.insert(
                list(paragraph_properties).index(section_properties),
                properties,
            )
    fonts = properties.find(qn("w:rFonts"))
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        properties.insert(0, fonts)
    for attribute in ("ascii", "hAnsi", "cs"):
        fonts.set(qn(f"w:{attribute}"), "Cambria Math")
    half_points = str(max(2, round(size * 2)))
    for tag in ("w:sz", "w:szCs"):
        node = properties.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            properties.append(node)
        node.set(qn("w:val"), half_points)


def _math_runs(paragraph: Paragraph, value: str, *, size: float, bold: bool = False) -> None:
    value = (
        value.replace("\\%", "%")
        .replace("\\times", "×")
        .replace("\\rightarrow", "→")
        .replace("\\leftrightarrow", "↔")
    )
    cursor = 0
    pattern = re.compile(r"([_^])\{([^{}]+)\}")
    for match in pattern.finditer(value):
        if match.start() > cursor:
            run = paragraph.add_run(value[cursor : match.start()])
            _set_font(run, size, bold=bold, italic=True)
        run = paragraph.add_run(match.group(2))
        _set_font(run, size * 0.82, bold=bold, italic=True)
        run.font.superscript = match.group(1) == "^"
        run.font.subscript = match.group(1) == "_"
        cursor = match.end()
    if cursor < len(value):
        run = paragraph.add_run(value[cursor:])
        _set_font(run, size, bold=bold, italic=True)


def _add_rich_text(
    paragraph: Paragraph,
    text: str,
    *,
    size: float,
    bold_prefix: bool = False,
    east_asia_heading: bool = False,
) -> None:
    if bold_prefix:
        match = re.match(
            r"^((?:(?:Câu|Question|Q\.?|Bài|Item)\s*[\w.-]+\s*[:.)]"
            r"|[A-Z]\s*[-\u2012-\u2015]\s*\d+|[A-Da-d][.)]))\s*",
            text,
            flags=re.IGNORECASE,
        )
        if match:
            run = paragraph.add_run(match.group(1) + " ")
            _set_font(run, size, bold=True)
            text = text[match.end() :]
    for segment in parse_markdown_inline(text):
        if segment.is_math:
            append_omml(paragraph, segment.value, font_size=size)
        else:
            run = paragraph.add_run(segment.value)
            _set_font(run, size, east_asia_heading=east_asia_heading)


def _new_paragraph(
    parent: WordDocumentType | _Cell,
    text: str,
    *,
    size: float,
    line_height: float,
    kind: MarkdownBlockKind = MarkdownBlockKind.PARAGRAPH,
    role: str | None = None,
    available_width_points: float | None = None,
    placement: HybridBlockPlacement | None = None,
    layout: ScanDocumentLayout | None = None,
) -> Paragraph:
    paragraph = parent.add_paragraph()
    heading = kind is MarkdownBlockKind.HEADING
    option = kind is MarkdownBlockKind.OPTION
    list_item = kind is MarkdownBlockKind.LIST_ITEM
    structural_heading = role in {"section_heading", "passage_lead", "question_lead"}
    quoted_passage = role == "quoted_passage"
    attribution = role == "attribution"
    rendered_size = size + (0.8 if role == "section_heading" else 0.2 if heading else 0)
    left_indent = 14.0 if list_item else 3.0 if option else 0.0
    source_geometry = _placement_geometry_points(
        placement,
        layout,
        available_width_points=available_width_points,
    )
    source_before = None
    right_indent = None
    if source_geometry is not None:
        source_before = (
            0.0
            if isinstance(parent, _Cell)
            else min(source_geometry.gap_before, line_height * 1.35)
        )
        measured_row_height = (
            statistics.median(source_geometry.row_heights)
            if source_geometry.row_heights
            else rendered_size * 1.08
        )
        if (
            placement is not None
            and layout is not None
            and not contains_tall_inline_math(text)
            and kind
            in {
                MarkdownBlockKind.PARAGRAPH,
                MarkdownBlockKind.OPTION,
                MarkdownBlockKind.LIST_ITEM,
                MarkdownBlockKind.CODE,
            }
        ):
            page = layout.pages[placement.page_number - 1]
            page_scale = (
                page.pdf_height / page.height
                if page.metadata.get("source_kind") == "image"
                else page.pdf_width / page.width
            )
            ordinary_row_ceiling = page.line_pitch * page_scale * 1.50
            ordinary_rows = [
                height for height in source_geometry.row_heights if height <= ordinary_row_ceiling
            ]
            if ordinary_rows:
                measured_row_height = statistics.median(ordinary_rows)
            elif source_geometry.row_heights:
                # Provider evidence can merge several baselines around a
                # figure into one tall bbox. It is a paragraph envelope, not
                # a native Word line-height instruction.
                measured_row_height = rendered_size * 1.08
        if heading and (
            measured_row_height > rendered_size * 1.22
            or source_geometry.height >= max(rendered_size * 2.8, line_height * 3.6)
        ):
            # A one-line display masthead can legitimately be an order of
            # magnitude larger than body text.  Preserve that editable scale
            # when positioned evidence supplies one tall row, while fitting
            # the wording into its measured horizontal envelope.  Multi-row
            # headings retain the restrained cap because their union height
            # is not a font-size measurement.
            visible_characters = max(1, len(re.sub(r"\s+", "", text)))
            single_display_row = len(source_geometry.row_heights) <= 1
            graphical_envelope = (
                visible_characters <= 40
                and source_geometry.height >= max(rendered_size * 2.8, line_height * 3.6)
                and source_geometry.height >= measured_row_height * 3.6
            )
            display_masthead = (
                single_display_row
                and measured_row_height >= max(rendered_size * 2.8, line_height * 2.4)
            ) or graphical_envelope
            if display_masthead:
                display_height = (
                    source_geometry.height if graphical_envelope else measured_row_height
                )
                frame_width = (
                    max(
                        1.0,
                        available_width_points
                        - source_geometry.left_indent
                        - source_geometry.right_indent,
                    )
                    if available_width_points is not None
                    else display_height * max(2.0, len(text) * 0.62)
                )
                width_fit = frame_width / (visible_characters * 0.62)
                rendered_size = max(
                    rendered_size,
                    min(112.0, display_height * 0.80, width_fit),
                )
            else:
                rendered_size = max(
                    rendered_size,
                    min(rendered_size * 2.15, measured_row_height * 0.82),
                )
        # Ink envelopes are evidence for baseline rhythm, not native Word text
        # frames.  In particular, a paragraph separated by a figure can have a
        # tall union bbox even though every source row is ordinary body text.
        line_height = max(
            line_height,
            rendered_size * 1.08,
            min(measured_row_height, line_height * 1.35),
        )
        # Ordinary prose must retain the full detected body column.  Glyph ink
        # does not extend to both paragraph edges, and treating x0/x1 as Word
        # indents caused short footer/question lines to collapse into a narrow
        # strip.  Source indents remain meaningful for preformatted code and
        # explicit list nesting only.
        if kind is MarkdownBlockKind.CODE:
            left_indent += source_geometry.left_indent
            right_indent = source_geometry.right_indent
        elif (
            heading
            and not isinstance(parent, _Cell)
            and placement is not None
            and layout is not None
            and layout.pages[placement.page_number - 1].metadata.get("column_count", 1) > 1
        ):
            # A spanning newspaper headline is rendered in the full-width
            # prefix, where its source left/right envelope expresses the
            # intended span.  Prose inside native column cells deliberately
            # continues to use the complete cell width.
            left_indent = max(left_indent, source_geometry.left_indent)
            right_indent = source_geometry.right_indent
    if attribution and available_width_points:
        left_indent = available_width_points * 0.40
    centered_cell_heading = bool(
        heading
        and isinstance(parent, _Cell)
        and source_geometry is not None
        and available_width_points is not None
        and min(source_geometry.left_indent, source_geometry.right_indent)
        >= available_width_points * 0.08
        and abs(source_geometry.left_indent - source_geometry.right_indent)
        <= available_width_points * 0.22
    )
    _format_paragraph(
        paragraph,
        size=rendered_size,
        line_height=line_height,
        alignment=(
            WD_ALIGN_PARAGRAPH.RIGHT
            if attribution
            else WD_ALIGN_PARAGRAPH.CENTER
            if centered_cell_heading
            else WD_ALIGN_PARAGRAPH.LEFT
            if heading or structural_heading or role == "form_field"
            else WD_ALIGN_PARAGRAPH.JUSTIFY
        ),
        before=(
            source_before
            if source_before is not None
            else 0.35
            if heading or structural_heading
            else 0.12
        ),
        after=0.08 if heading or structural_heading else 0,
        left_indent=left_indent,
        right_indent=right_indent,
        keep_next=heading or structural_heading,
    )
    _add_rich_text(
        paragraph,
        text,
        size=rendered_size,
        east_asia_heading=heading or structural_heading,
        bold_prefix=option
        or bool(
            re.match(
                r"^(?:(?:Câu|Question|Q\.?|Bài|Item)\s|[A-Z]\s*[-\u2012-\u2015]\s*\d+)",
                text,
                re.I,
            )
        ),
    )
    if heading or structural_heading:
        for run in paragraph.runs:
            run.bold = True
    if quoted_passage or attribution:
        for run in paragraph.runs:
            run.italic = True
    if quoted_passage:
        paragraph.paragraph_format.first_line_indent = Pt(29)
    if list_item:
        paragraph.paragraph_format.first_line_indent = Pt(-8)
    if re.search(r"\\(?:frac|dfrac|tfrac|int|sum|prod|sqrt|lfloor|lceil)\b", text):
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
        paragraph.paragraph_format.line_spacing = Pt(line_height)
    _apply_evidence_style(
        paragraph,
        placement.evidence_style if placement is not None else None,
        allow_alignment=not (heading or structural_heading or attribution),
    )
    return paragraph


def _new_equation(
    parent: WordDocumentType | _Cell,
    latex: str,
    *,
    size: float,
    line_height: float,
    placement: HybridBlockPlacement | None = None,
    layout: ScanDocumentLayout | None = None,
    available_width_points: float | None = None,
) -> Paragraph:
    paragraph = parent.add_paragraph()
    row_count = equation_row_count(latex)
    tall = bool(re.search(r"\\(?:frac|dfrac|tfrac|int|sum|prod|sqrt)\b", latex))
    row_floor = max(
        line_height * (1.62 if tall else 1.18),
        size + (8.0 if tall else 4.0),
    )
    source_geometry = _placement_geometry_points(
        placement,
        layout,
        available_width_points=available_width_points,
    )
    source_before = 0.0
    source_left = 0.0
    source_right = None
    row_spacing = None
    if source_geometry is not None:
        source_before = 0.0 if isinstance(parent, _Cell) else source_geometry.gap_before
        source_left = source_geometry.left_indent
        source_right = source_geometry.right_indent
        if source_geometry.row_heights:
            # The paragraph line is a floor for one visual row, not a box for
            # the entire equation array.  The Office Math array owns its total
            # height and receives the measured inter-row spacing below.
            row_floor = max(row_floor, max(source_geometry.row_heights))
        elif row_count == 1:
            row_floor = max(row_floor, source_geometry.height)
        row_spacing = source_geometry.row_spacing if row_count > 1 else None
    _format_paragraph(
        paragraph,
        size=size,
        line_height=row_floor,
        alignment=(
            WD_ALIGN_PARAGRAPH.LEFT if source_geometry is not None else WD_ALIGN_PARAGRAPH.CENTER
        ),
        before=source_before,
        after=0,
        left_indent=source_left,
        right_indent=source_right,
    )
    # Office Math arrays determine their own total height.  An exact one-line
    # box clips or overlaps aligned derivations, so retain only a per-row floor.
    if row_count > 1 or tall or source_geometry is not None:
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
        paragraph.paragraph_format.line_spacing = Pt(row_floor)
        paragraph.paragraph_format.keep_together = True
    _set_paragraph_mark_math_format(paragraph, size)
    append_omml(
        paragraph,
        latex,
        # Keep display and inline mathematics on the same document-wide base
        # size.  OMML itself performs the semantic enlargement/reduction for
        # operators, fractions, limits, and scripts.
        font_size=size,
        display=True,
        justification="left" if source_geometry is not None else "centerGroup",
        row_spacing=row_spacing,
    )
    return paragraph


def _set_cell_margins(cell: _Cell, *, value: int = 32) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for side in ("top", "start", "bottom", "end"):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value if side in {"start", "end"} else 0))
        node.set(qn("w:type"), "dxa")


def _set_table_borders(table: Table, *, visible: bool, size: int = 5) -> None:
    properties = table._tbl.tblPr
    borders = properties.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        properties.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single" if visible else "nil")
        if visible:
            node.set(qn("w:sz"), str(size))
            node.set(qn("w:color"), "000000")


def _set_table_widths(table: Table, widths: Sequence[float]) -> None:
    table.autofit = False
    properties = table._tbl.tblPr
    layout = properties.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        properties.append(layout)
    layout.set(qn("w:type"), "fixed")
    twips = [max(1, int(round(width * 1440))) for width in widths]
    table_width = properties.first_child_found_in("w:tblW")
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        properties.append(table_width)
    table_width.set(qn("w:type"), "dxa")
    table_width.set(qn("w:w"), str(sum(twips)))
    grid = table._tbl.tblGrid
    columns = list(grid)
    while len(columns) < len(twips):
        column = OxmlElement("w:gridCol")
        grid.append(column)
        columns.append(column)
    for column, twip_width in zip(columns, twips, strict=False):
        column.set(qn("w:w"), str(twip_width))
    for row in table.rows:
        for index, width_inches in enumerate(widths):
            if index < len(row.cells):
                row.cells[index].width = Inches(width_inches)
                cell_width = row.cells[index]._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
                if cell_width is not None:
                    cell_width.set(qn("w:type"), "dxa")
                    cell_width.set(qn("w:w"), str(twips[index]))


def _clear_cell(cell: _Cell, *, size: float, line_height: float) -> None:
    cell.text = ""
    _format_paragraph(
        cell.paragraphs[0],
        size=size,
        line_height=line_height,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
    )
    _set_cell_margins(cell)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def _set_row_minimum_height(row: Any, height_points: float) -> None:
    """Give a native Word table row a non-clipping source-derived floor."""

    properties = row._tr.get_or_add_trPr()
    height = properties.find(qn("w:trHeight"))
    if height is None:
        height = OxmlElement("w:trHeight")
        properties.append(height)
    height.set(qn("w:val"), str(max(1, round(height_points * 20))))
    height.set(qn("w:hRule"), "atLeast")


def _borderless_table(
    parent: WordDocumentType | _Cell,
    rows: int,
    columns: int,
    widths: Sequence[float],
    *,
    size: float,
    line_height: float,
) -> Table:
    table = parent.add_table(rows=rows, cols=columns)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(table, visible=False)
    _set_table_widths(table, widths)
    for row in table.rows:
        properties = row._tr.get_or_add_trPr()
        if properties.find(qn("w:cantSplit")) is None:
            properties.append(OxmlElement("w:cantSplit"))
        for cell in row.cells:
            _clear_cell(cell, size=size, line_height=line_height)
    return table


def _add_native_table(
    parent: WordDocumentType | _Cell,
    block: MarkdownBlock,
    *,
    available_width: float,
    size: float,
    line_height: float,
    placement: HybridBlockPlacement | None = None,
    layout: ScanDocumentLayout | None = None,
) -> Table | None:
    rows = block.table_rows
    columns = max((len(row) for row in rows), default=0)
    if not rows or not columns:
        return None
    table = parent.add_table(rows=len(rows), cols=columns)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(table, visible=True, size=4)
    _set_table_widths(table, [available_width / columns] * columns)
    source_geometry = _placement_geometry_points(
        placement,
        layout,
        available_width_points=available_width * 72.0,
    )
    source_row_floor = source_geometry.height / len(rows) if source_geometry is not None else 0.0
    row_floor = max(
        7.4,
        line_height * 0.76,
        min(source_row_floor, line_height * 1.85),
    )
    for row_index, row in enumerate(rows):
        row_properties = table.rows[row_index]._tr.get_or_add_trPr()
        if row_properties.find(qn("w:cantSplit")) is None:
            row_properties.append(OxmlElement("w:cantSplit"))
        _set_row_minimum_height(table.rows[row_index], row_floor)
        for column_index in range(columns):
            cell = table.cell(row_index, column_index)
            _clear_cell(cell, size=size * 0.78, line_height=row_floor)
            value = row[column_index] if column_index < len(row) else ""
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _add_rich_text(paragraph, value, size=size * 0.78)
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
            paragraph.paragraph_format.line_spacing = Pt(row_floor)
            paragraph.paragraph_format.keep_together = True
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_index == 0:
                for run in paragraph.runs:
                    run.bold = True
                properties = cell._tc.get_or_add_tcPr()
                shading = OxmlElement("w:shd")
                shading.set(qn("w:fill"), "F2F2F2")
                properties.append(shading)
    return table


def _add_picture(
    parent: WordDocumentType | _Cell,
    data: bytes,
    *,
    width_inches: float,
    alignment: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.CENTER,
) -> Paragraph:
    paragraph = parent.add_paragraph()
    formatting = paragraph.paragraph_format
    formatting.space_before = Pt(0)
    formatting.space_after = Pt(0)
    formatting.line_spacing = None
    formatting.line_spacing_rule = None
    paragraph.alignment = alignment
    paragraph.add_run().add_picture(io.BytesIO(data), width=Inches(max(0.15, width_inches)))
    return paragraph


def _crop_bytes(layout: ScanDocumentLayout, page_number: int, bbox: PixelBox) -> bytes:
    page = layout.pages[page_number - 1]
    crop = page.image.crop((bbox.x0, bbox.y0, bbox.x1, bbox.y1))
    output = io.BytesIO()
    crop.save(output, format="PNG", optimize=True)
    return output.getvalue()


def _source_figure_bytes(
    asset_bytes: dict[str, bytes],
    block_id: str,
    layout: ScanDocumentLayout,
    page_number: int,
    bbox: PixelBox,
) -> bytes:
    """Prefer saved asset bytes only when their shape agrees with geometry."""

    data = asset_bytes.get(block_id)
    if data is not None:
        try:
            with Image.open(io.BytesIO(data)) as image:
                intrinsic_aspect = image.width / max(1, image.height)
            source_aspect = bbox.width / max(1, bbox.height)
            aspect_factor = max(
                intrinsic_aspect / max(source_aspect, 1e-6),
                source_aspect / max(intrinsic_aspect, 1e-6),
            )
            if aspect_factor <= 1.35:
                return data
        except OSError:
            pass
    # Provider URLs can point at a padded or differently scaled preview even
    # when JSON owns an exact source bbox. Recropping that authoritative box
    # preserves the physical aspect and prevents a Word row from inflating.
    return _crop_bytes(layout, page_number, bbox)


def _top_boundary_ink_rate(data: bytes) -> float:
    try:
        image = Image.open(io.BytesIO(data)).convert("L")
    except OSError:
        return 0.0
    depth = max(1, min(image.size) // 100)
    pixels = image.crop((0, 0, image.width, depth)).tobytes()
    return sum(value < 180 for value in pixels) / max(1, len(pixels))


def _trim_embedded_figure_tail(text: str) -> str:
    """Hide a tiny OCR tail duplicated by an immediately following figure."""

    match = re.match(r"^(.+?[.!?])\s+([^.!?]{1,32})$", text)
    if not match:
        return text
    tail = match.group(2).split()
    alphabetic = sum(character.isalpha() for character in match.group(2))
    if 1 <= len(tail) <= 5 and alphabetic <= 8 and any(len(token) == 1 for token in tail):
        return match.group(1)
    return text


def _source_option_grid(
    blocks: Sequence[MarkdownBlock],
    placements: dict[str, HybridBlockPlacement],
    layout: ScanDocumentLayout,
) -> tuple[int, float] | None:
    """Recover a row-major option grid from distinct source slot boxes."""

    option_placements = [placements.get(block.id) for block in blocks]
    if not option_placements or any(
        placement is None or placement.source_bbox is None for placement in option_placements
    ):
        return None
    resolved = [cast(HybridBlockPlacement, placement) for placement in option_placements]
    if len({placement.page_number for placement in resolved}) != 1:
        return None
    coordinates = {
        (
            cast(PixelBox, placement.source_bbox).x0,
            cast(PixelBox, placement.source_bbox).y0,
            cast(PixelBox, placement.source_bbox).x1,
            cast(PixelBox, placement.source_bbox).y1,
        )
        for placement in resolved
    }
    # A shared coarse bbox plus several source rows does not identify which
    # option owns which column. Preserve the established content fallback.
    if len(coordinates) != len(resolved):
        return None

    row_groups: list[list[HybridBlockPlacement]] = []
    for placement in resolved:
        box = cast(PixelBox, placement.source_bbox)
        if not row_groups:
            row_groups.append([placement])
            continue
        previous_boxes = [cast(PixelBox, item.source_bbox) for item in row_groups[-1]]
        overlap = max(
            max(0, min(box.y1, previous.y1) - max(box.y0, previous.y0))
            / max(1, min(box.height, previous.height))
            for previous in previous_boxes
        )
        if overlap >= 0.55:
            row_groups[-1].append(placement)
        else:
            row_groups.append([placement])
    columns = len(row_groups[0])
    if columns < 1 or any(len(row) != columns for row in row_groups):
        return None
    if columns > 1 and any(
        any(
            cast(PixelBox, left.source_bbox).x0 >= cast(PixelBox, right.source_bbox).x0
            for left, right in zip(row, row[1:], strict=False)
        )
        for row in row_groups
    ):
        return None
    page = layout.pages[resolved[0].page_number - 1]
    vertical_scale = (
        page.pdf_height / page.height
        if page.metadata.get("source_kind") == "image"
        else page.pdf_width / page.width
    )
    boxes = [cast(PixelBox, placement.source_bbox) for placement in resolved]
    height = (max(box.y1 for box in boxes) - min(box.y0 for box in boxes)) * vertical_scale
    return columns, height


def _render_options(
    parent: WordDocumentType | _Cell,
    blocks: list[MarkdownBlock],
    placements: dict[str, HybridBlockPlacement],
    *,
    width: float,
    size: float,
    line_height: float,
    layout: ScanDocumentLayout,
) -> None:
    longest = max((len(block.text) for block in blocks), default=0)
    maximum_fraction_count = max(
        (len(re.findall(r"\\(?:frac|dfrac|tfrac)\b", block.text)) for block in blocks),
        default=0,
    )
    first_placement = placements.get(blocks[0].id) if blocks else None
    source_geometry = _placement_geometry_points(
        first_placement,
        layout,
        available_width_points=width * 72.0,
    )
    source_grid = _source_option_grid(blocks, placements, layout)
    source_row_count = len(source_geometry.row_heights) if source_geometry is not None else 0
    has_nary = any(re.search(r"\\(?:int|oint|sum|prod)\b", block.text) for block in blocks)
    columns = (
        source_grid[0]
        if source_grid is not None
        else 2
        if len(blocks) == 4 and source_row_count >= 2 and has_nary and maximum_fraction_count == 0
        else 2
        if (
            len(blocks) == 4
            and source_geometry is None
            and has_nary
            and longest <= 115
            and maximum_fraction_count <= 1
        )
        else 1
        if len(blocks) == 4 and source_geometry is None and has_nary
        else len(blocks)
        if len(blocks) in {2, 3, 4} and longest <= 55 and maximum_fraction_count <= 1
        else 2
        if len(blocks) == 4 and longest <= 115
        else 1
    )
    if columns > 1:
        rows = math.ceil(len(blocks) / columns)
        table = _borderless_table(
            parent,
            rows,
            columns,
            [width / columns] * columns,
            size=size,
            line_height=line_height,
        )
        tall_math = any(contains_tall_inline_math(block.text) for block in blocks)
        natural_floor = max(
            line_height,
            size * (1.82 if tall_math else 1.18),
        )
        source_total = (
            source_grid[1]
            if source_grid is not None
            else source_geometry.height
            if source_geometry is not None
            else 0.0
        )
        row_floor = max(natural_floor, source_total / max(1, rows))
        row_floor = min(row_floor, line_height * (2.35 if tall_math else 1.55))
        source_before = (
            min(source_geometry.gap_before, line_height * 1.20)
            if source_geometry is not None
            else 0.0
        )
        for row_index, row in enumerate(table.rows):
            _set_row_minimum_height(row, row_floor + (source_before if row_index == 0 else 0.0))
        for index, block in enumerate(blocks):
            paragraph = table.cell(index // columns, index % columns).paragraphs[0]
            _add_rich_text(paragraph, block.text, size=size * 0.96, bold_prefix=True)
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
            paragraph.paragraph_format.line_spacing = Pt(row_floor)
            if index < columns and source_before:
                paragraph.paragraph_format.space_before = Pt(source_before)
            paragraph.paragraph_format.keep_together = True
            table.cell(
                index // columns, index % columns
            ).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        return
    for block in blocks:
        _new_paragraph(
            parent,
            block.text,
            size=size,
            line_height=line_height,
            kind=MarkdownBlockKind.OPTION,
        )


def _add_thematic_rule(
    parent: WordDocumentType | _Cell,
    *,
    line_height: float,
) -> None:
    """Render Markdown thematic breaks as native editable paragraph borders."""

    paragraph = parent.add_paragraph()
    _format_paragraph(
        paragraph,
        size=2.0,
        line_height=max(2.0, line_height * 0.22),
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        before=min(1.5, line_height * 0.10),
        after=min(1.5, line_height * 0.10),
    )
    properties = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    borders.append(bottom)
    properties.append(borders)


def _render_linear(
    parent: WordDocumentType | _Cell,
    blocks: list[MarkdownBlock],
    placements: dict[str, HybridBlockPlacement],
    *,
    width: float,
    size: float,
    line_height: float,
    asset_bytes: dict[str, bytes],
    layout: ScanDocumentLayout,
) -> None:
    index = 0
    while index < len(blocks):
        block = blocks[index]
        if block.kind is MarkdownBlockKind.OPTION:
            end = index + 1
            while end < len(blocks) and blocks[end].kind is MarkdownBlockKind.OPTION:
                end += 1
            option_blocks = list(blocks[index:end])
            if end < len(blocks) and blocks[end].kind is MarkdownBlockKind.IMAGE:
                last = option_blocks[-1]
                trimmed = _trim_embedded_figure_tail(last.text)
                if trimmed != last.text:
                    option_blocks[-1] = last.model_copy(update={"text": trimmed})
            _render_options(
                parent,
                option_blocks,
                placements,
                width=width,
                size=size,
                line_height=line_height,
                layout=layout,
            )
            index = end
            continue
        if block.kind is MarkdownBlockKind.TABLE:
            _add_native_table(
                parent,
                block,
                available_width=width,
                size=size,
                line_height=line_height,
                placement=placements.get(block.id),
                layout=layout,
            )
        elif block.kind is MarkdownBlockKind.IMAGE:
            placement = placements[block.id]
            bbox = placement.source_bbox
            if bbox:
                page = layout.pages[placement.page_number - 1]
                picture_width = bbox.width / page.width * page.pdf_width / 72.0
                data = _source_figure_bytes(
                    asset_bytes,
                    block.id,
                    layout,
                    placement.page_number,
                    bbox,
                )
                display_width = min(width, picture_width)
                left_outdent = 0.0
                right_outdent = 0.0
                if not isinstance(parent, _Cell):
                    render_frame = _page_render_content_bbox(page)
                    horizontal_scale = page.pdf_width / page.width
                    left_outdent = max(0.0, render_frame.x0 - bbox.x0) * horizontal_scale
                    right_outdent = max(0.0, bbox.x1 - render_frame.x1) * horizontal_scale
                    display_width = min(
                        page.pdf_width / 72.0,
                        width + (left_outdent + right_outdent) / 72.0,
                        picture_width,
                    )
                paragraph = _add_picture(
                    parent,
                    data,
                    width_inches=display_width,
                    alignment=(
                        WD_ALIGN_PARAGRAPH.LEFT
                        if left_outdent or right_outdent
                        else WD_ALIGN_PARAGRAPH.CENTER
                    ),
                )
                if left_outdent:
                    paragraph.paragraph_format.left_indent = Pt(-left_outdent)
                if right_outdent:
                    paragraph.paragraph_format.right_indent = Pt(-right_outdent)
        elif block.kind is MarkdownBlockKind.RULE:
            _add_thematic_rule(parent, line_height=line_height)
        elif block.kind is MarkdownBlockKind.CODE:
            _new_paragraph(
                parent,
                block.text,
                size=size * 0.9,
                line_height=line_height,
                kind=block.kind,
                role=str(block.metadata.get("role", "")) or None,
                available_width_points=width * 72.0,
                placement=placements.get(block.id),
                layout=layout,
            )
        elif block.kind is MarkdownBlockKind.EQUATION:
            _new_equation(
                parent,
                block.text,
                size=size,
                line_height=line_height,
                placement=placements.get(block.id),
                layout=layout,
                available_width_points=width * 72.0,
            )
        else:
            _new_paragraph(
                parent,
                block.text,
                size=size,
                line_height=line_height,
                kind=block.kind,
                role=str(block.metadata.get("role", "")) or None,
                available_width_points=width * 72.0,
                placement=placements.get(block.id),
                layout=layout,
            )
        index += 1


def _overlap_ratio(left: PixelBox, right: PixelBox) -> float:
    overlap = max(0, min(left.y1, right.y1) - max(left.y0, right.y0))
    return overlap / max(1, min(left.height, right.height))


def _horizontal_image_row(
    blocks: Sequence[MarkdownBlock],
    placements: dict[str, HybridBlockPlacement],
) -> list[tuple[MarkdownBlock, PixelBox]]:
    """Return same-page images that form one disjoint, overlapping row."""

    candidates: list[tuple[MarkdownBlock, PixelBox]] = []
    page_numbers: set[int] = set()
    for block in blocks:
        if block.kind is not MarkdownBlockKind.IMAGE:
            continue
        placement = placements.get(block.id)
        if placement is None or placement.source_bbox is None:
            return []
        page_numbers.add(placement.page_number)
        candidates.append((block, placement.source_bbox))
    if len(candidates) < 2 or len(page_numbers) != 1:
        return []
    candidates.sort(key=lambda item: item[1].x0)
    boxes = [box for _block, box in candidates]
    if any(left.x1 > right.x0 for left, right in zip(boxes, boxes[1:], strict=False)):
        return []
    common_overlap = max(0, min(box.y1 for box in boxes) - max(box.y0 for box in boxes))
    if common_overlap / max(1, min(box.height for box in boxes)) < 0.55:
        return []
    return candidates


def _add_horizontal_image_row(
    parent: WordDocumentType | _Cell,
    images: Sequence[tuple[MarkdownBlock, PixelBox]],
    *,
    frame: PixelBox,
    width: float,
    size: float,
    line_height: float,
    asset_bytes: dict[str, bytes],
    layout: ScanDocumentLayout,
    page_number: int,
) -> Table:
    """Render independently editable image objects in one native table row."""

    source_widths = [float(max(1, images[0][1].x0 - frame.x0))]
    for index, (_block, box) in enumerate(images):
        source_widths.append(float(box.width))
        if index + 1 < len(images):
            source_widths.append(float(max(1, images[index + 1][1].x0 - box.x1)))
    source_widths.append(float(max(1, frame.x1 - images[-1][1].x1)))
    normalization = max(1.0, sum(source_widths))
    widths = [width * value / normalization for value in source_widths]
    table = _borderless_table(
        parent,
        1,
        len(widths),
        widths,
        size=size,
        line_height=line_height,
    )
    caption = OxmlElement("w:tblCaption")
    caption.set(qn("w:val"), "docreconstruct:horizontal-visual-row")
    table._tbl.tblPr.append(caption)
    page = layout.pages[page_number - 1]
    for image_index, (block, box) in enumerate(images):
        column = image_index * 2 + 1
        data = _source_figure_bytes(asset_bytes, block.id, layout, page_number, box)
        source_width = box.width / page.width * page.pdf_width / 72.0
        _add_picture(
            table.cell(0, column),
            data,
            width_inches=min(widths[column], source_width),
        )
    return table


def _render_horizontal_visual_group(
    parent: WordDocumentType | _Cell,
    blocks: list[MarkdownBlock],
    placements: dict[str, HybridBlockPlacement],
    *,
    width: float,
    size: float,
    line_height: float,
    asset_bytes: dict[str, bytes],
    layout: ScanDocumentLayout,
) -> bool:
    """Render a full-width same-group image band without vertical stacking."""

    images = _horizontal_image_row(blocks, placements)
    if not images:
        return False
    image_ids = {block.id for block, _box in images}
    image_indices = [index for index, block in enumerate(blocks) if block.id in image_ids]
    if image_indices != list(range(min(image_indices), max(image_indices) + 1)):
        return False
    page_number = placements[images[0][0].id].page_number
    page = layout.pages[page_number - 1]
    union = PixelBox(
        x0=min(box.x0 for _block, box in images),
        y0=min(box.y0 for _block, box in images),
        x1=max(box.x1 for _block, box in images),
        y1=max(box.y1 for _block, box in images),
    )
    flow = [block for block in blocks if block.id not in image_ids]
    left_fraction = (union.x0 - page.content_bbox.x0) / max(1, page.content_bbox.width)
    if flow and left_fraction >= 0.34 and union.width <= page.content_bbox.width * 0.58:
        return False
    first = min(image_indices)
    last = max(image_indices)
    _render_linear(
        parent,
        blocks[:first],
        placements,
        width=width,
        size=size,
        line_height=line_height,
        asset_bytes=asset_bytes,
        layout=layout,
    )
    _add_horizontal_image_row(
        parent,
        images,
        frame=page.content_bbox,
        width=width,
        size=size,
        line_height=line_height,
        asset_bytes=asset_bytes,
        layout=layout,
        page_number=page_number,
    )
    _render_linear(
        parent,
        blocks[last + 1 :],
        placements,
        width=width,
        size=size,
        line_height=line_height,
        asset_bytes=asset_bytes,
        layout=layout,
    )
    return True


def _render_image_table_pair(
    parent: WordDocumentType | _Cell,
    blocks: list[MarkdownBlock],
    placements: dict[str, HybridBlockPlacement],
    *,
    width: float,
    size: float,
    line_height: float,
    asset_bytes: dict[str, bytes],
    layout: ScanDocumentLayout,
) -> bool:
    image = next((block for block in blocks if block.kind is MarkdownBlockKind.IMAGE), None)
    table_block = next((block for block in blocks if block.kind is MarkdownBlockKind.TABLE), None)
    if image is None or table_block is None:
        return False
    image_box = placements[image.id].source_bbox
    table_box = placements[table_block.id].source_bbox
    if not image_box or not table_box or _overlap_ratio(image_box, table_box) < 0.25:
        return False
    page = layout.pages[placements[image.id].page_number - 1]
    midpoint = page.content_bbox.x0 + page.content_bbox.width / 2
    image_center = (image_box.x0 + image_box.x1) / 2
    table_center = (table_box.x0 + table_box.x1) / 2
    if image_center >= midpoint or table_center <= midpoint:
        return False
    first_visual = min(blocks.index(image), blocks.index(table_block))
    pre = [
        block
        for block in blocks[:first_visual]
        if block.kind not in {MarkdownBlockKind.IMAGE, MarkdownBlockKind.TABLE}
    ]
    post = [
        block
        for block in blocks[first_visual:]
        if block.kind not in {MarkdownBlockKind.IMAGE, MarkdownBlockKind.TABLE}
    ]
    _render_linear(
        parent,
        pre,
        placements,
        width=width,
        size=size,
        line_height=line_height,
        asset_bytes=asset_bytes,
        layout=layout,
    )
    left_fraction = max(
        0.32, min(0.68, (image_box.x1 - page.content_bbox.x0) / page.content_bbox.width)
    )
    pair = _borderless_table(
        parent,
        1,
        2,
        [width * left_fraction, width * (1 - left_fraction)],
        size=size,
        line_height=line_height,
    )
    image_width = image_box.width / page.width * page.pdf_width / 72.0
    data = _source_figure_bytes(asset_bytes, image.id, layout, page.number, image_box)
    _add_picture(pair.cell(0, 0), data, width_inches=min(width * left_fraction, image_width))
    _add_native_table(
        pair.cell(0, 1),
        table_block,
        available_width=width * (1 - left_fraction),
        size=size,
        line_height=line_height,
        placement=placements.get(table_block.id),
        layout=layout,
    )
    _render_linear(
        parent,
        post,
        placements,
        width=width,
        size=size,
        line_height=line_height,
        asset_bytes=asset_bytes,
        layout=layout,
    )
    return True


def _render_side_visual_group(
    parent: WordDocumentType | _Cell,
    blocks: list[MarkdownBlock],
    placements: dict[str, HybridBlockPlacement],
    *,
    width: float,
    size: float,
    line_height: float,
    asset_bytes: dict[str, bytes],
    layout: ScanDocumentLayout,
) -> bool:
    visuals = [block for block in blocks if block.kind is MarkdownBlockKind.IMAGE]
    if not visuals:
        tables = [block for block in blocks if block.kind is MarkdownBlockKind.TABLE]
        if len(tables) != 1:
            return False
        visual = tables[0]
    else:
        visual = visuals[0]
    optional_boxes = [placements[block.id].source_bbox for block in visuals or [visual]]
    if not optional_boxes or any(box is None for box in optional_boxes):
        return False
    boxes = [cast(PixelBox, box) for box in optional_boxes]
    page_number = placements[visual.id].page_number
    page = layout.pages[page_number - 1]
    union = PixelBox(
        x0=min(box.x0 for box in boxes),
        y0=min(box.y0 for box in boxes),
        x1=max(box.x1 for box in boxes),
        y1=max(box.y1 for box in boxes),
    )
    left_fraction = (union.x0 - page.content_bbox.x0) / max(1, page.content_bbox.width)
    if left_fraction < 0.34 or union.width > page.content_bbox.width * 0.58:
        return False
    left_fraction = max(0.38, min(0.76, left_fraction))
    flow_blocks = [block for block in blocks if block not in (visuals or [visual])]
    if visuals:
        flow_blocks = [
            block.model_copy(update={"text": _trim_embedded_figure_tail(block.text)})
            if block.kind is MarkdownBlockKind.OPTION
            else block
            for block in flow_blocks
        ]
    pair = _borderless_table(
        parent,
        1,
        2,
        [width * left_fraction, width * (1 - left_fraction)],
        size=size,
        line_height=line_height,
    )
    _render_linear(
        pair.cell(0, 0),
        flow_blocks,
        placements,
        width=width * left_fraction,
        size=size,
        line_height=line_height,
        asset_bytes=asset_bytes,
        layout=layout,
    )
    right = pair.cell(0, 1)
    if visuals:
        horizontal = _horizontal_image_row(visuals, placements)
        if horizontal:
            _add_horizontal_image_row(
                right,
                horizontal,
                frame=union,
                width=width * (1 - left_fraction),
                size=size,
                line_height=line_height,
                asset_bytes=asset_bytes,
                layout=layout,
                page_number=page_number,
            )
        else:
            for block in visuals:
                placement = placements[block.id]
                box = placement.source_bbox
                if box:
                    data = _source_figure_bytes(
                        asset_bytes,
                        block.id,
                        layout,
                        page_number,
                        box,
                    )
                    if 0.04 <= _top_boundary_ink_rate(data) <= 0.75:
                        expanded = PixelBox(
                            x0=box.x0,
                            y0=max(page.content_bbox.y0, box.y0 - round(page.line_pitch)),
                            x1=box.x1,
                            y1=box.y1,
                        )
                        data = _crop_bytes(layout, page_number, expanded)
                    picture_width = box.width / page.width * page.pdf_width / 72.0
                    _add_picture(
                        right,
                        data,
                        width_inches=min(width * (1 - left_fraction), picture_width),
                    )
    else:
        _add_native_table(
            right,
            visual,
            available_width=width * (1 - left_fraction),
            size=size,
            line_height=line_height,
            placement=placements.get(visual.id),
            layout=layout,
        )
    return True


def _render_shared_table_cluster(
    parent: WordDocumentType | _Cell,
    blocks: list[MarkdownBlock],
    placements: dict[str, HybridBlockPlacement],
    *,
    width: float,
    size: float,
    line_height: float,
    asset_bytes: dict[str, bytes],
    layout: ScanDocumentLayout,
) -> bool:
    tables = [block for block in blocks if block.kind is MarkdownBlockKind.TABLE]
    if len(tables) != 2:
        return False
    first_box = placements[tables[0].id].source_bbox
    second_box = placements[tables[1].id].source_bbox
    if not first_box or first_box != second_box:
        return False
    first_index = blocks.index(tables[0])
    second_index = blocks.index(tables[1])
    pre = blocks[:first_index]
    between = blocks[first_index + 1 : second_index]
    post = blocks[second_index + 1 :]
    _render_linear(
        parent,
        pre,
        placements,
        width=width,
        size=size,
        line_height=line_height,
        asset_bytes=asset_bytes,
        layout=layout,
    )
    first_columns = max((len(row) for row in tables[0].table_rows), default=1)
    second_columns = max((len(row) for row in tables[1].table_rows), default=1)
    left_fraction = max(0.22, min(0.48, first_columns / (first_columns + second_columns)))
    pair = _borderless_table(
        parent,
        1,
        2,
        [width * left_fraction, width * (1 - left_fraction)],
        size=size,
        line_height=line_height,
    )
    _add_native_table(
        pair.cell(0, 0),
        tables[0],
        available_width=width * left_fraction,
        size=size,
        line_height=line_height,
        placement=placements.get(tables[0].id),
        layout=layout,
    )
    if between:
        _render_linear(
            pair.cell(0, 1),
            between[-1:],
            placements,
            width=width * (1 - left_fraction),
            size=size * 0.9,
            line_height=line_height * 0.9,
            asset_bytes=asset_bytes,
            layout=layout,
        )
    _add_native_table(
        pair.cell(0, 1),
        tables[1],
        available_width=width * (1 - left_fraction),
        size=size,
        line_height=line_height,
        placement=placements.get(tables[1].id),
        layout=layout,
    )
    if len(between) > 1:
        _render_linear(
            pair.cell(0, 0),
            between[:-1],
            placements,
            width=width * left_fraction,
            size=size * 0.9,
            line_height=line_height * 0.9,
            asset_bytes=asset_bytes,
            layout=layout,
        )
    _render_linear(
        parent,
        post,
        placements,
        width=width,
        size=size,
        line_height=line_height,
        asset_bytes=asset_bytes,
        layout=layout,
    )
    return True


def _render_multiple_side_tables(
    parent: WordDocumentType | _Cell,
    blocks: list[MarkdownBlock],
    placements: dict[str, HybridBlockPlacement],
    *,
    width: float,
    size: float,
    line_height: float,
    asset_bytes: dict[str, bytes],
    layout: ScanDocumentLayout,
) -> bool:
    """Recreate a text-left/data-right region containing several native tables."""

    tables = [block for block in blocks if block.kind is MarkdownBlockKind.TABLE]
    if len(tables) < 2:
        return False
    boxes = [placements[block.id].source_bbox for block in tables]
    if any(box is None for box in boxes):
        return False
    table_boxes = [box for box in boxes if box is not None]
    page_number = placements[tables[0].id].page_number
    if any(placements[block.id].page_number != page_number for block in tables):
        return False
    page = layout.pages[page_number - 1]
    side_threshold = page.content_bbox.x0 + page.content_bbox.width * 0.34
    if any(
        box.x0 < side_threshold or box.width > page.content_bbox.width * 0.68 for box in table_boxes
    ):
        return False
    first_index = min(blocks.index(block) for block in tables)
    last_index = max(blocks.index(block) for block in tables)
    if first_index == 0:
        return False
    post_start = len(blocks)
    for index in range(last_index + 1, len(blocks)):
        if blocks[index].kind is MarkdownBlockKind.OPTION:
            post_start = index
            break
    left_blocks = blocks[:first_index]
    right_blocks = blocks[first_index:post_start]
    post_blocks = blocks[post_start:]
    left_fraction = (min(box.x0 for box in table_boxes) - page.content_bbox.x0) / max(
        1, page.content_bbox.width
    )
    left_fraction = max(0.34, min(0.60, left_fraction))
    pair = _borderless_table(
        parent,
        1,
        2,
        [width * left_fraction, width * (1 - left_fraction)],
        size=size,
        line_height=line_height,
    )
    _render_linear(
        pair.cell(0, 0),
        left_blocks,
        placements,
        width=width * left_fraction,
        size=size,
        line_height=line_height,
        asset_bytes=asset_bytes,
        layout=layout,
    )
    _render_linear(
        pair.cell(0, 1),
        right_blocks,
        placements,
        width=width * (1 - left_fraction),
        size=size,
        line_height=line_height,
        asset_bytes=asset_bytes,
        layout=layout,
    )
    if post_blocks:
        _render_linear(
            parent,
            post_blocks,
            placements,
            width=width,
            size=size,
            line_height=line_height,
            asset_bytes=asset_bytes,
            layout=layout,
        )
    return True


def _render_group(
    parent: WordDocumentType | _Cell,
    blocks: list[MarkdownBlock],
    placements: dict[str, HybridBlockPlacement],
    *,
    width: float,
    size: float,
    line_height: float,
    asset_bytes: dict[str, bytes],
    layout: ScanDocumentLayout,
) -> None:
    if _render_horizontal_visual_group(
        parent,
        blocks,
        placements,
        width=width,
        size=size,
        line_height=line_height,
        asset_bytes=asset_bytes,
        layout=layout,
    ):
        return
    if _render_image_table_pair(
        parent,
        blocks,
        placements,
        width=width,
        size=size,
        line_height=line_height,
        asset_bytes=asset_bytes,
        layout=layout,
    ):
        return
    if _render_shared_table_cluster(
        parent,
        blocks,
        placements,
        width=width,
        size=size,
        line_height=line_height,
        asset_bytes=asset_bytes,
        layout=layout,
    ):
        return
    if _render_multiple_side_tables(
        parent,
        blocks,
        placements,
        width=width,
        size=size,
        line_height=line_height,
        asset_bytes=asset_bytes,
        layout=layout,
    ):
        return
    if _render_side_visual_group(
        parent,
        blocks,
        placements,
        width=width,
        size=size,
        line_height=line_height,
        asset_bytes=asset_bytes,
        layout=layout,
    ):
        return
    _render_linear(
        parent,
        blocks,
        placements,
        width=width,
        size=size,
        line_height=line_height,
        asset_bytes=asset_bytes,
        layout=layout,
    )


def _header_crop(page: Any) -> PixelBox | None:
    """Retain an unmatched first-page masthead without copying a full scan."""

    bands = sorted(page.line_bands)
    if len(bands) < 4:
        return None
    end = None
    for index in range(2, min(len(bands) - 1, 9)):
        if bands[index + 1][0] - bands[index][1] > page.line_pitch * 1.02:
            end = bands[index][1] + int(page.line_pitch * 0.42)
            break
    if end is None or end >= page.height * 0.22:
        return None
    return PixelBox(
        x0=page.content_bbox.x0,
        y0=page.content_bbox.y0,
        x1=page.content_bbox.x1,
        y1=min(page.content_bbox.y1, end),
    )


def _compact_empty_paragraphs_in_cells(document: WordDocumentType) -> None:
    """Remove cell padding paragraphs without producing invalid WordprocessingML.

    A table cell whose final child is a nested table must retain a trailing
    paragraph.  Removing that paragraph makes office applications repair the
    file differently and can collapse a fixed two-column layout into one flow.
    """

    for cell in document._element.body.xpath(".//w:tc"):
        content = [child for child in cell if child.tag != qn("w:tcPr")]
        for child in list(content):
            if child.tag != qn("w:p"):
                continue
            if child.xpath(".//w:t[normalize-space(.)] | .//w:drawing | .//m:oMath"):
                continue
            is_required_trailing = child is content[-1] and any(
                sibling.tag == qn("w:tbl") for sibling in content[:-1]
            )
            if not is_required_trailing and len(content) > 1:
                cell.remove(child)
                content.remove(child)
                continue
            properties = child.get_or_add_pPr()
            spacing = properties.find(qn("w:spacing"))
            if spacing is None:
                spacing = OxmlElement("w:spacing")
                properties.append(spacing)
            spacing.set(qn("w:before"), "0")
            spacing.set(qn("w:after"), "0")
            spacing.set(qn("w:lineRule"), "exact")
            spacing.set(qn("w:line"), "20")


def _places_figure_pixels(content: MarkdownContent, layout: ScanDocumentLayout) -> bool:
    """Report whether this job renders any original figure pixels at all."""

    if any(block.kind is MarkdownBlockKind.IMAGE for block in content.blocks):
        return True
    return any(
        region.kind is not ScanRegionKind.TABLE for page in layout.pages for region in page.regions
    )


def _duplicate_figure_annotation_ids(
    blocks: Sequence[MarkdownBlock],
    *,
    figures_rendered: bool,
) -> set[str]:
    """Find OCR labels interleaved between options and already present in figures.

    OCR-to-Markdown tools sometimes emit short labels from a nearby figure in
    the middle of an A-D option sequence.  The matched original figure already
    contains those pixels, so rendering the labels again would duplicate and
    misplace them.  This structural rule is document-agnostic and never alters
    ordinary prose or a complete option.

    That justification only holds when some figure is actually placed.  A
    text-only exam still matches the surrounding shape — a short unpunctuated
    line after the final option, before the next question — so without
    ``figures_rendered`` the rule deleted reviewed Markdown that no raster
    reproduced, and the ``native_content_projection`` gate, which still counts
    that text as required, then failed the job.
    """

    duplicate_ids: set[str] = set()
    if not figures_rendered:
        return duplicate_ids
    index = 1
    while index < len(blocks) - 1:
        if blocks[index - 1].kind is not MarkdownBlockKind.OPTION:
            index += 1
            continue
        end = index
        while (
            end < len(blocks)
            and blocks[end].kind is MarkdownBlockKind.PARAGRAPH
            and len(blocks[end].text) <= 160
            and not re.search(r"[.!?;:]\s*$", blocks[end].text)
        ):
            end += 1
        if (
            end > index
            and end < len(blocks)
            and blocks[end].kind is MarkdownBlockKind.OPTION
            and blocks[index - 1].group_id
            and blocks[index - 1].group_id == blocks[end].group_id
        ):
            duplicate_ids.update(block.id for block in blocks[index:end])
        index = max(index + 1, end)
    for index in range(1, len(blocks) - 1):
        block = blocks[index]
        if (
            block.kind is MarkdownBlockKind.PARAGRAPH
            and len(block.text) <= 80
            and not re.search(r"[.!?;:]\s*$", block.text)
            and blocks[index + 1].starts_group
            and block.group_id
            and block.group_id == blocks[index - 1].group_id
            and blocks[index - 1].kind
            in {MarkdownBlockKind.OPTION, MarkdownBlockKind.IMAGE, MarkdownBlockKind.TABLE}
        ):
            duplicate_ids.add(block.id)
    return duplicate_ids


def _body_groups(blocks: Sequence[MarkdownBlock]) -> list[list[MarkdownBlock]]:
    groups: list[list[MarkdownBlock]] = []
    for block in blocks:
        if not groups or (block.starts_group and groups[-1]):
            groups.append([block])
        else:
            groups[-1].append(block)
    return groups


def _column_weight(blocks: Sequence[MarkdownBlock], *, characters_per_line: float) -> float:
    total = 0.0
    for block in blocks:
        if block.kind is MarkdownBlockKind.EQUATION:
            total += equation_layout_units(block.text)
        elif block.kind is MarkdownBlockKind.TABLE:
            total += max(1.5, len(block.table_rows) * 1.15)
        elif block.kind is MarkdownBlockKind.IMAGE:
            total += 6.0
        elif block.kind is MarkdownBlockKind.RULE:
            total += 0.35
        else:
            visible = re.sub(r"\$([^$]+)\$", r"\1", block.text)
            total += max(1.0, math.ceil(len(visible) / characters_per_line))
            if block.kind is MarkdownBlockKind.LIST_ITEM:
                total += 0.15
    return total


def _split_plain_block_for_column(
    block: MarkdownBlock,
    *,
    target_weight: float,
    characters_per_line: float,
) -> tuple[MarkdownBlock, MarkdownBlock] | None:
    """Split editable prose at a word boundary for a native column break."""

    if block.kind is not MarkdownBlockKind.PARAGRAPH or "$" in block.text or target_weight < 1.0:
        return None
    target_character = round(target_weight * characters_per_line)
    boundaries = [match.start() for match in re.finditer(r"\s+", block.text)]
    eligible = [
        boundary
        for boundary in boundaries
        if characters_per_line <= boundary <= len(block.text) - characters_per_line
    ]
    if not eligible:
        return None
    split_at = min(eligible, key=lambda boundary: abs(boundary - target_character))
    prefix = block.text[:split_at].rstrip()
    suffix = block.text[split_at:].lstrip()
    if not prefix or not suffix:
        return None
    return (
        block.model_copy(update={"text": prefix}),
        block.model_copy(update={"text": suffix}),
    )


def _balanced_editable_column_streams(
    blocks: Sequence[MarkdownBlock],
    *,
    widths: Sequence[float],
    target_heights: Sequence[float],
) -> list[list[MarkdownBlock]]:
    """Flow ordered native prose through similarly sized newspaper columns.

    Markdown parsers normally preserve reading order, but a long paragraph can
    cross a physical column boundary.  A cell-per-block assignment cannot
    express that.  For ungrouped text-only layouts this routine preserves the
    exact wording while splitting only at whitespace and balancing cumulative
    text weight by the source column-height proportions.
    """

    if (
        len(widths) < 2
        or len(widths) != len(target_heights)
        or not blocks
        or any(
            block.kind
            in {
                MarkdownBlockKind.IMAGE,
                MarkdownBlockKind.TABLE,
                MarkdownBlockKind.EQUATION,
                MarkdownBlockKind.CODE,
                MarkdownBlockKind.OPTION,
            }
            for block in blocks
        )
    ):
        return []
    reference_characters = max(24.0, statistics.median(widths) * 17.0)
    total_weight = _column_weight(blocks, characters_per_line=reference_characters)
    total_height = sum(max(1.0, value) for value in target_heights)
    targets = [
        total_weight * max(1.0, height) / max(1.0, total_height) for height in target_heights
    ]
    pending = list(blocks)
    streams: list[list[MarkdownBlock]] = []
    for column_index in range(len(widths) - 1):
        characters_per_line = max(24.0, widths[column_index] * 17.0)
        target = targets[column_index]
        stream: list[MarkdownBlock] = []
        current = 0.0
        while pending:
            block = pending[0]
            weight = _column_weight([block], characters_per_line=characters_per_line)
            if current + weight <= target + 0.25:
                stream.append(pending.pop(0))
                current += weight
                continue
            remaining = target - current
            split = _split_plain_block_for_column(
                block,
                target_weight=remaining,
                characters_per_line=characters_per_line,
            )
            if split is not None:
                prefix, suffix = split
                stream.append(prefix)
                pending[0] = suffix
            elif not stream:
                stream.append(pending.pop(0))
            break
        streams.append(stream)
    streams.append(pending)
    return streams if all(streams) else []


def _dialogue_story_boundary(blocks: Sequence[MarkdownBlock]) -> int | None:
    """Find a short byline immediately followed by a sustained dialogue."""

    def is_dialogue_turn(block: MarkdownBlock) -> bool:
        return block.kind is MarkdownBlockKind.LIST_ITEM or bool(
            re.match(r"^\s*[-\u2012-\u2015]", block.text)
        )

    for index in range(1, len(blocks) - 3):
        if len(blocks[index].text) > 80 or is_dialogue_turn(blocks[index]):
            continue
        following = blocks[index + 1 : index + 7]
        if (
            following
            and is_dialogue_turn(following[0])
            and sum(is_dialogue_turn(block) for block in following) >= 3
        ):
            return index
    return None


def _split_column_groups(
    groups: list[list[MarkdownBlock]],
    *,
    left_width: float,
    right_width: float,
) -> tuple[list[MarkdownBlock], list[MarkdownBlock]]:
    if len(groups) < 2:
        flattened = [block for group in groups for block in group]
        midpoint = max(1, len(flattened) // 2)
        return flattened[:midpoint], flattened[midpoint:]
    left_characters = max(24.0, left_width * 12.5)
    right_characters = max(24.0, right_width * 12.5)
    best_index = 1
    best_cost = math.inf
    for index in range(1, len(groups)):
        left_blocks = [block for group in groups[:index] for block in group]
        right_blocks = [block for group in groups[index:] for block in group]
        left_weight = _column_weight(left_blocks, characters_per_line=left_characters)
        right_weight = _column_weight(right_blocks, characters_per_line=right_characters)
        cost = abs(left_weight - right_weight) / max(1.0, left_weight + right_weight)
        if cost < best_cost:
            best_cost = cost
            best_index = index
    return (
        [block for group in groups[:best_index] for block in group],
        [block for group in groups[best_index:] for block in group],
    )


def _render_header_blocks(
    document: WordDocumentType,
    blocks: Sequence[MarkdownBlock],
    *,
    size: float,
    line_height: float,
) -> float:
    consumed = 0.0
    for index, block in enumerate(blocks):
        if block.kind is MarkdownBlockKind.EQUATION:
            _new_equation(document, block.text, size=size, line_height=line_height)
            consumed += line_height * equation_layout_units(block.text)
            continue
        paragraph = document.add_paragraph()
        title = index == 0 and block.kind is MarkdownBlockKind.HEADING
        text = block.text
        title_lines = 1
        if title and len(text) > 86:
            preferred = round(len(text) * 0.72)
            candidates = [match.start() for match in re.finditer(r"\s+", text)]
            if candidates:
                split = min(candidates, key=lambda position: abs(position - preferred))
                text = text[:split].rstrip() + "\n" + text[split:].lstrip()
                title_lines = 2
        title_increment = 2.7 if title else 0.15
        paragraph_height = line_height + (2.4 if title else 0.3)
        after = 4.0 if index == len(blocks) - 1 else 1.0
        _format_paragraph(
            paragraph,
            size=size + title_increment,
            line_height=paragraph_height,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            before=0,
            after=after,
            keep_next=True,
        )
        _add_rich_text(paragraph, text, size=size + title_increment)
        if title:
            for run in paragraph.runs:
                run.bold = True
        consumed += paragraph_height * title_lines + after
    return consumed


def _wrap_column_blocks(
    blocks: Sequence[MarkdownBlock],
    *,
    characters_per_line: int,
) -> list[MarkdownBlock]:
    result: list[MarkdownBlock] = []
    for block in blocks:
        if block.kind not in {MarkdownBlockKind.PARAGRAPH, MarkdownBlockKind.LIST_ITEM}:
            result.append(block)
            continue
        lines = textwrap.wrap(
            block.text,
            width=max(28, characters_per_line),
            expand_tabs=False,
            replace_whitespace=False,
            drop_whitespace=True,
            break_long_words=False,
            break_on_hyphens=False,
        )
        result.append(block.model_copy(update={"text": "\n".join(lines)}))
    return result


def _add_vertical_spacer(document: WordDocumentType, points: float) -> None:
    if points <= 1.0:
        return
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(max(0.0, points - 1.0))
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    paragraph.paragraph_format.line_spacing = Pt(1.0)
    run = paragraph.add_run("\u00a0")
    _set_font(run, 1.0)


def _render_two_column_page(
    document: WordDocumentType,
    blocks: list[MarkdownBlock],
    placements: dict[str, HybridBlockPlacement],
    *,
    width: float,
    size: float,
    line_height: float,
    asset_bytes: dict[str, bytes],
    layout: ScanDocumentLayout,
    page_number: int,
) -> bool:
    first_group = next((index for index, block in enumerate(blocks) if block.starts_group), None)
    if first_group is None or first_group == len(blocks):
        return False
    page = layout.pages[page_number - 1]
    boxes = page.metadata.get("column_boxes")
    gutter = page.metadata.get("column_gutter")
    if not (
        isinstance(boxes, list)
        and len(boxes) == 2
        and all(isinstance(box, list) and len(box) == 4 for box in boxes)
        and isinstance(gutter, list)
        and len(gutter) == 2
    ):
        return False
    content_width = max(1, page.content_bbox.width)
    left_fraction = max(0.25, (float(boxes[0][2]) - float(boxes[0][0])) / content_width)
    right_fraction = max(0.25, (float(boxes[1][2]) - float(boxes[1][0])) / content_width)
    gap_fraction = max(0.018, (float(gutter[1]) - float(gutter[0])) / content_width)
    normalization = left_fraction + gap_fraction + right_fraction
    left_width = width * left_fraction / normalization
    gap_width = width * gap_fraction / normalization
    right_width = width * right_fraction / normalization
    header_height = _render_header_blocks(
        document,
        blocks[:first_group],
        size=size,
        line_height=line_height,
    )
    page_scale = page.pdf_width / page.width
    raster_offset = max(0.0, (page.pdf_height - page.height * page_scale) / 2)
    target_body_top = raster_offset + float(boxes[0][1]) * page_scale
    top_margin = document.sections[-1].top_margin
    section_top = float(top_margin.pt) if top_margin is not None else 0.0
    _add_vertical_spacer(document, target_body_top - section_top - header_height)
    groups = _body_groups(blocks[first_group:])
    left_blocks, right_blocks = _split_column_groups(
        groups,
        left_width=left_width,
        right_width=right_width,
    )
    bottoms = page.metadata.get("column_content_bottoms")
    if isinstance(bottoms, list) and len(bottoms) == 2:
        left_target = (float(bottoms[0]) - float(boxes[0][1])) / page.line_pitch
        right_target = (float(bottoms[1]) - float(boxes[1][1])) / page.line_pitch
    else:
        left_target = right_target = 0.0
    left_characters = round(max(34.0, left_width * 17.0))
    right_characters = round(max(34.0, right_width * 17.0))
    left_units = _column_weight(left_blocks, characters_per_line=left_characters)
    right_units = _column_weight(right_blocks, characters_per_line=right_characters)
    # Three line-units model Word's fixed table/paragraph overhead.  The
    # remaining source capacity controls only leading, preserving natural line
    # widths while matching unequal column depths.
    left_line_scale = max(0.82, min(1.24, left_target / (left_units + 3.0))) if left_target else 1.0
    right_line_scale = (
        max(0.82, min(1.24, right_target / (right_units + 3.0))) if right_target else 1.0
    )
    left_blocks = _wrap_column_blocks(
        left_blocks,
        characters_per_line=left_characters,
    )
    right_blocks = _wrap_column_blocks(
        right_blocks,
        characters_per_line=right_characters,
    )
    table = _borderless_table(
        document,
        1,
        3,
        [left_width, gap_width, right_width],
        size=size,
        line_height=line_height,
    )
    for column, column_blocks, column_width, column_line_height in (
        (0, left_blocks, left_width, line_height * left_line_scale),
        (2, right_blocks, right_width, line_height * right_line_scale),
    ):
        cursor = 0
        while cursor < len(column_blocks):
            group_id = column_blocks[cursor].group_id
            end = cursor + 1
            while end < len(column_blocks) and column_blocks[end].group_id == group_id:
                end += 1
            _render_group(
                table.cell(0, column),
                column_blocks[cursor:end],
                placements,
                width=column_width,
                size=size,
                line_height=column_line_height,
                asset_bytes=asset_bytes,
                layout=layout,
            )
            cursor = end
    return True


def _source_column_boxes(page: Any) -> list[PixelBox]:
    """Return validated, left-to-right source columns for a 2--4 column page."""

    count = page.metadata.get("column_count")
    values = page.metadata.get("column_boxes")
    if isinstance(count, bool) or not isinstance(count, int) or count not in {2, 3, 4}:
        return []
    if not isinstance(values, list) or len(values) != count:
        return []
    boxes: list[PixelBox] = []
    for value in values:
        try:
            if isinstance(value, dict):
                box = PixelBox.model_validate(value)
            elif isinstance(value, list) and len(value) == 4:
                box = PixelBox(
                    x0=round(float(value[0])),
                    y0=round(float(value[1])),
                    x1=round(float(value[2])),
                    y1=round(float(value[3])),
                )
            else:
                return []
        except (TypeError, ValueError):
            return []
        boxes.append(box)
    boxes.sort(key=lambda box: box.x0)
    if any(left.x1 >= right.x0 for left, right in zip(boxes, boxes[1:], strict=False)):
        return []
    return boxes


def _column_table_widths(boxes: Sequence[PixelBox], width: float) -> list[float]:
    """Project source column and gutter widths into one fixed Word table row."""

    source_widths: list[float] = []
    for index, box in enumerate(boxes):
        if index:
            source_widths.append(float(max(1, box.x0 - boxes[index - 1].x1)))
        source_widths.append(float(box.width))
    normalization = max(1.0, sum(source_widths))
    return [width * source_width / normalization for source_width in source_widths]


def _source_column_ink_capacities(
    page: Any,
    boxes: Sequence[PixelBox],
    *,
    body_top: float,
    bottoms: Sequence[float],
) -> list[float]:
    """Estimate relative editable story capacity from source foreground mass."""

    paper = str(page.metadata.get("paper_color", "FFFFFF"))
    try:
        paper_luminance = sum(int(paper[index : index + 2], 16) for index in (0, 2, 4)) / 3
    except (TypeError, ValueError):
        paper_luminance = 255.0
    threshold = max(96, min(190, round(paper_luminance - 48)))
    capacities: list[float] = []
    for index, box in enumerate(boxes):
        top = max(box.y0, round(body_top))
        bottom = min(box.y1, round(bottoms[index]))
        if bottom <= top:
            return []
        histogram = page.image.crop((box.x0, top, box.x1, bottom)).convert("L").histogram()
        capacities.append(float(sum(histogram[:threshold])))
    if not capacities or min(capacities) < 32:
        return []
    # Reject a diagram-dominated outlier; ordinary newspaper columns can vary
    # in weight, but a tenfold foreground jump is not text-flow capacity.
    if max(capacities) / min(capacities) > 3.5:
        return []
    # Foreground mass also reflects font weight and scan darkness, not just
    # word volume.  Square-root damping keeps that style evidence useful
    # without starving a lighter neighboring column of editable content.
    return [math.sqrt(capacity) for capacity in capacities]


def _source_column_index(box: PixelBox, columns: Sequence[PixelBox]) -> int:
    """Choose the source column with the strongest horizontal intersection."""

    overlaps = [max(0, min(box.x1, column.x1) - max(box.x0, column.x0)) for column in columns]
    if max(overlaps, default=0) > 0:
        return max(range(len(columns)), key=overlaps.__getitem__)
    center = (box.x0 + box.x1) / 2.0
    return min(
        range(len(columns)),
        key=lambda index: abs(center - (columns[index].x0 + columns[index].x1) / 2.0),
    )


def _render_block_stream(
    parent: WordDocumentType | _Cell,
    blocks: Sequence[MarkdownBlock],
    placements: dict[str, HybridBlockPlacement],
    *,
    width: float,
    size: float,
    line_height: float,
    asset_bytes: dict[str, bytes],
    layout: ScanDocumentLayout,
) -> None:
    """Render an ordered block stream while preserving explicit semantic groups."""

    cursor = 0
    while cursor < len(blocks):
        group_id = blocks[cursor].group_id
        end = cursor + 1
        if group_id is not None:
            while end < len(blocks) and blocks[end].group_id == group_id:
                end += 1
        _render_group(
            parent,
            list(blocks[cursor:end]),
            placements,
            width=width,
            size=size,
            line_height=line_height,
            asset_bytes=asset_bytes,
            layout=layout,
        )
        cursor = end


def _render_multi_column_page(
    document: WordDocumentType,
    blocks: list[MarkdownBlock],
    placements: dict[str, HybridBlockPlacement],
    *,
    width: float,
    size: float,
    line_height: float,
    asset_bytes: dict[str, bytes],
    layout: ScanDocumentLayout,
    page_number: int,
) -> bool:
    """Render a geometry-assigned 2--4 column page without semantic group markers."""

    page = layout.pages[page_number - 1]
    columns = _source_column_boxes(page)
    if not columns or not blocks:
        return False

    body_top = min(column.y0 for column in columns)
    raw_bottoms = page.metadata.get("column_content_bottoms")
    bottoms = (
        [float(value) for value in raw_bottoms]
        if isinstance(raw_bottoms, list)
        and len(raw_bottoms) == len(columns)
        and all(isinstance(value, int | float) for value in raw_bottoms)
        else [float(column.y1) for column in columns]
    )
    body_bottom = max(bottoms)
    tolerance = max(2.0, page.line_pitch * 0.35)
    median_column_width = statistics.median(column.width for column in columns)
    leading_anchor_bottom = max(
        (
            placement.source_bbox.y1
            for block in blocks
            if block.kind is MarkdownBlockKind.IMAGE
            and (placement := placements.get(block.id)) is not None
            and placement.source_bbox is not None
            and placement.source_bbox.y0 <= body_top + page.line_pitch
        ),
        default=body_top,
    )
    spanning_deadline = max(body_top, leading_anchor_bottom) + page.line_pitch * 4.0
    assignments: list[str | int | None] = []
    for block in blocks:
        placement = placements.get(block.id)
        box = placement.source_bbox if placement is not None else None
        if box is None:
            assignments.append(None)
        elif box.y1 <= body_top + tolerance:
            assignments.append("prefix")
        elif box.y0 >= body_bottom + tolerance:
            assignments.append("suffix")
        elif box.width >= median_column_width * 1.45 and box.y0 < spanning_deadline:
            assignments.append("prefix")
        else:
            assignments.append(_source_column_index(box, columns))

    if not any(isinstance(assignment, int) for assignment in assignments):
        return False
    # A geometry-less annotation inherits the nearest known flow.  This keeps
    # content editable without inventing coordinates and is only a fallback;
    # source-assigned placements remain authoritative whenever present.
    for index, assignment in enumerate(assignments):
        if assignment is not None:
            continue
        previous = next(
            (
                assignments[candidate]
                for candidate in range(index - 1, -1, -1)
                if assignments[candidate] is not None
            ),
            None,
        )
        following = next(
            (
                assignments[candidate]
                for candidate in range(index + 1, len(assignments))
                if assignments[candidate] is not None
            ),
            None,
        )
        assignments[index] = previous if previous is not None else following

    prefix = [
        block
        for block, assignment in zip(blocks, assignments, strict=True)
        if assignment == "prefix"
    ]
    suffix = [
        block
        for block, assignment in zip(blocks, assignments, strict=True)
        if assignment == "suffix"
    ]
    column_blocks: list[list[MarkdownBlock]] = [[] for _ in columns]
    for block, assignment in zip(blocks, assignments, strict=True):
        if isinstance(assignment, int):
            column_blocks[assignment].append(block)
    if not any(column_blocks):
        return False

    body_placement_tops = [
        box.y0
        for block, assignment in zip(blocks, assignments, strict=True)
        if isinstance(assignment, int)
        and (placement := placements.get(block.id)) is not None
        and (box := placement.source_bbox) is not None
    ]
    effective_body_top = max(
        body_top,
        min(body_placement_tops, default=body_top),
    )
    table_widths = _column_table_widths(columns, width)
    editable_widths = [table_widths[index * 2] for index in range(len(columns))]
    body_sequence = [
        block
        for block, assignment in zip(blocks, assignments, strict=True)
        if isinstance(assignment, int)
    ]
    body_geometry_coverage = sum(
        1
        for block in body_sequence
        if (placement := placements.get(block.id)) is not None and placement.source_bbox is not None
    ) / max(1, len(body_sequence))
    geometry_materially_covers_body = body_geometry_coverage >= 0.50
    target_heights = [max(1.0, bottom - effective_body_top) for bottom in bottoms]
    ink_capacities = _source_column_ink_capacities(
        page,
        columns,
        body_top=effective_body_top,
        bottoms=bottoms,
    )
    flow_capacities = ink_capacities or target_heights
    balanced_streams: list[list[MarkdownBlock]] = []
    if (
        len(columns) >= 3
        and not geometry_materially_covers_body
        and not any(block.starts_group for block in body_sequence)
    ):
        dialogue_boundary = _dialogue_story_boundary(body_sequence)
        if dialogue_boundary is not None:
            leading_streams = _balanced_editable_column_streams(
                body_sequence[:dialogue_boundary],
                widths=editable_widths[:-1],
                target_heights=flow_capacities[:-1],
            )
            if len(leading_streams) == len(columns) - 1:
                balanced_streams = [
                    *leading_streams,
                    list(body_sequence[dialogue_boundary:]),
                ]
        if not balanced_streams:
            balanced_streams = _balanced_editable_column_streams(
                body_sequence,
                widths=editable_widths,
                target_heights=flow_capacities,
            )
    if balanced_streams:
        column_blocks = balanced_streams

    def source_order(block: MarkdownBlock) -> tuple[int, int]:
        placement = placements.get(block.id)
        source_box = placement.source_bbox if placement is not None else None
        return (
            source_box.y0 if source_box is not None else page.height + block.index,
            block.index,
        )

    if not balanced_streams:
        for stream in column_blocks:
            stream.sort(key=source_order)

    _render_block_stream(
        document,
        prefix,
        placements,
        width=width,
        size=size,
        line_height=line_height,
        asset_bytes=asset_bytes,
        layout=layout,
    )
    table = _borderless_table(
        document,
        1,
        len(table_widths),
        table_widths,
        size=size,
        line_height=line_height,
    )
    # The body grid is deliberately one logical row so each native Word cell
    # represents a source column.  It must nevertheless be allowed to split
    # at paragraph boundaries: an unsplittable row that is even slightly
    # taller than the space below a masthead is moved outside the page (and
    # LibreOffice clips it completely).  Ordinary compact option/table rows
    # still retain ``cantSplit`` via ``_borderless_table``.
    body_row_properties = table.rows[0]._tr.get_or_add_trPr()
    cannot_split = body_row_properties.find(qn("w:cantSplit"))
    if cannot_split is not None:
        body_row_properties.remove(cannot_split)
    caption = OxmlElement("w:tblCaption")
    caption.set(qn("w:val"), f"docreconstruct:body-columns-{len(columns)}")
    table._tbl.tblPr.append(caption)
    for index, stream in enumerate(column_blocks):
        if not stream:
            continue
        source_target = max(0.0, bottoms[index] - effective_body_top) / max(1.0, page.line_pitch)
        characters_per_line = max(24.0, table_widths[index * 2] * 17.0)
        content_units = _column_weight(stream, characters_per_line=characters_per_line)
        line_scale = (
            max(0.72, min(1.0, source_target / (content_units + 1.5))) if source_target else 1.0
        )
        _render_block_stream(
            table.cell(0, index * 2),
            stream,
            placements,
            width=table_widths[index * 2],
            size=size,
            line_height=line_height * line_scale,
            asset_bytes=asset_bytes,
            layout=layout,
        )
    _render_block_stream(
        document,
        suffix,
        placements,
        width=width,
        size=size,
        line_height=line_height,
        asset_bytes=asset_bytes,
        layout=layout,
    )
    return True


def _masthead_code_field(text: str) -> bool:
    """Return whether a short label/value pair is a compact document code.

    A code value is deliberately recognized by shape rather than language. It
    must be one compact token after the only colon, which excludes ordinary
    fields such as a duration followed by its unit.
    """

    if text.count(":") != 1:
        return False
    value = text.rpartition(":")[2].strip()
    return bool(re.fullmatch(r"[A-Z0-9][A-Z0-9./_-]{1,15}", value, flags=re.IGNORECASE))


def _masthead_emphasis(block: MarkdownBlock) -> bool:
    """Recognize display masthead text without relying on document wording."""

    if block.kind is MarkdownBlockKind.HEADING:
        return True
    display_text = re.sub(r"\s*\([^()]*\)\s*$", "", block.text).strip()
    letters = [character for character in display_text if character.isalpha()]
    return bool(
        len(letters) >= 5
        and sum(character.isupper() for character in letters) / len(letters) >= 0.86
    )


def _masthead_fallback_side(
    block: MarkdownBlock,
    *,
    index: int,
    primary_heading_index: int | None,
) -> str:
    """Classify geometry-ambiguous furniture by reusable structural shape."""

    text = block.text.strip()
    role = str(block.metadata.get("role", ""))
    if primary_heading_index is not None:
        if index < primary_heading_index:
            return "left"
        if index == primary_heading_index:
            return "right"
    if role == "form_field" or (
        block.kind is MarkdownBlockKind.PARAGRAPH and (text.endswith(":") or text.count(":") >= 2)
    ):
        return "left"
    if _masthead_code_field(text):
        return "right"
    if block.kind is MarkdownBlockKind.HEADING:
        return "left" if primary_heading_index is None and index == 0 else "right"
    if text.startswith("(") and text.endswith(")"):
        return "left"
    if ":" in text:
        return "right"
    return "left" if _masthead_emphasis(block) or len(text) <= 42 else "right"


def _masthead_geometry_side(
    placement: HybridBlockPlacement | None,
    *,
    render_frame: PixelBox,
    divider: float,
) -> str | None:
    """Assign a masthead block when its source box clearly favors one side."""

    if placement is None or placement.source_bbox is None:
        return None
    box = placement.source_bbox
    left_overlap = max(
        0.0,
        min(float(box.x1), divider) - max(float(box.x0), float(render_frame.x0)),
    )
    right_overlap = max(
        0.0,
        min(float(box.x1), float(render_frame.x1)) - max(float(box.x0), divider),
    )
    covered = left_overlap + right_overlap
    if covered <= 0:
        return None
    dominance = max(left_overlap, right_overlap) / covered
    # A nearly page-wide synthetic/estimated box is not column evidence. The
    # threshold still accepts a form row that modestly crosses the divider.
    if dominance < 0.62:
        return None
    return "left" if left_overlap > right_overlap else "right"


def _sort_masthead_entries(
    entries: list[tuple[MarkdownBlock, str, bool]],
    placements: dict[str, HybridBlockPlacement],
) -> list[tuple[MarkdownBlock, str, bool]]:
    """Keep each source column in physical top-to-bottom order when known."""

    if not entries or any(
        (placement := placements.get(block.id)) is None or placement.source_bbox is None
        for block, _text, _emphasis in entries
    ):
        return sorted(entries, key=lambda entry: entry[0].index)
    return sorted(
        entries,
        key=lambda entry: (
            cast(PixelBox, placements[entry[0].id].source_bbox).y0,
            entry[0].index,
        ),
    )


def _render_masthead_cell(
    cell: _Cell,
    entries: Sequence[tuple[MarkdownBlock, str, bool]],
    *,
    side: str,
    width_inches: float,
    size: float,
    line_height: float,
    placements: dict[str, HybridBlockPlacement],
    source_x0: float,
    source_x1: float,
) -> None:
    cell.text = ""
    _set_cell_margins(cell, value=18)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    for index, (block, text, title_fragment) in enumerate(entries):
        paragraph = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
        role = str(block.metadata.get("role", ""))
        punctuation_fields = text.count(":") >= 1 and block.kind is MarkdownBlockKind.PARAGRAPH
        form_field = role == "form_field" or (side == "left" and punctuation_fields)
        placement = placements.get(block.id)
        box = placement.source_bbox if placement is not None else None
        source_width = max(1.0, source_x1 - source_x0)
        clipped_width = (
            max(
                0.0,
                min(source_x1, float(box.x1)) - max(source_x0, float(box.x0)),
            )
            if box is not None
            else source_width
        )
        code_field = (
            side == "right"
            and punctuation_fields
            and _masthead_code_field(text)
            and clipped_width / source_width <= 0.55
        )
        if form_field and text.count(":") >= 2:
            text = re.sub(r"(?<=\.{3})\s+(?=[A-ZÀ-Ỹ])", "\n ", text, count=1)
        if side == "left" and title_fragment:
            text = re.sub(r"\s+(?=\([^()]*\)\s*$)", "\n", text, count=1)
        rendered_size = size * 0.86 + (0.25 if title_fragment else 0.0)
        _format_paragraph(
            paragraph,
            size=rendered_size,
            line_height=line_height,
            alignment=(
                WD_ALIGN_PARAGRAPH.LEFT
                if form_field and not code_field
                else WD_ALIGN_PARAGRAPH.CENTER
            ),
            before=0,
            after=0,
            keep_next=True,
        )
        _add_rich_text(paragraph, text, size=rendered_size)
        bold = (
            title_fragment
            or block.kind is MarkdownBlockKind.HEADING
            or (len(text) <= 42 and not text.startswith("(") and not form_field)
        )
        italic = (not title_fragment) and (
            text.startswith("(") or (side == "right" and len(text) >= 36)
        )
        for run in paragraph.runs:
            if bold:
                run.bold = True
            if italic:
                run.italic = True
        if code_field:
            if box is not None:
                clipped_x0 = max(source_x0, min(source_x1, float(box.x0)))
                clipped_x1 = max(source_x0, min(source_x1, float(box.x1)))
                if clipped_x1 > clipped_x0:
                    paragraph.paragraph_format.left_indent = Pt(
                        width_inches * 72.0 * (clipped_x0 - source_x0) / source_width
                    )
                    paragraph.paragraph_format.right_indent = Pt(
                        width_inches * 72.0 * (source_x1 - clipped_x1) / source_width
                    )
            properties = paragraph._p.get_or_add_pPr()
            borders = properties.find(qn("w:pBdr"))
            if borders is None:
                borders = OxmlElement("w:pBdr")
                properties.append(borders)
            for edge in ("top", "left", "bottom", "right"):
                border = borders.find(qn(f"w:{edge}"))
                if border is None:
                    border = OxmlElement(f"w:{edge}")
                    borders.append(border)
                border.set(qn("w:val"), "single")
                border.set(qn("w:sz"), "6")
                border.set(qn("w:space"), "2")
                border.set(qn("w:color"), "000000")
        if form_field and text.rstrip().endswith(":"):
            tab_position = max(18.0, width_inches * 72.0 - 8.0)
            paragraph.paragraph_format.tab_stops.add_tab_stop(
                Pt(tab_position),
                WD_TAB_ALIGNMENT.RIGHT,
                WD_TAB_LEADER.DOTS,
            )
            leader = paragraph.add_run("\t\ufeff")
            _set_font(leader, rendered_size)


def _masthead_entries(
    preamble: Sequence[MarkdownBlock],
    placements: dict[str, HybridBlockPlacement],
    *,
    render_frame: PixelBox,
    divider: float,
) -> (
    tuple[
        list[tuple[MarkdownBlock, str, bool]],
        list[tuple[MarkdownBlock, str, bool]],
    ]
    | None
):
    """Classify separate Markdown masthead blocks into two native zones.

    OCR/Markdown providers commonly serialize the left authority column first
    and the right examination title second.  The old renderer only handled a
    single merged title.  This classifier uses document structure and generic
    form punctuation; it never recognizes or rewrites the source wording.
    """

    if not preamble:
        return None
    headings = [
        (index, block)
        for index, block in enumerate(preamble)
        if block.kind is MarkdownBlockKind.HEADING
    ]
    primary_index = (
        max(headings, key=lambda item: len(item[1].text))[0] if len(headings) >= 2 else None
    )

    left_entries: list[tuple[MarkdownBlock, str, bool]] = []
    right_entries: list[tuple[MarkdownBlock, str, bool]] = []
    for index, block in enumerate(preamble):
        text = block.text
        side = _masthead_geometry_side(
            placements.get(block.id),
            render_frame=render_frame,
            divider=divider,
        ) or _masthead_fallback_side(
            block,
            index=index,
            primary_heading_index=primary_index,
        )
        entry = (block, text, _masthead_emphasis(block))
        (left_entries if side == "left" else right_entries).append(entry)

    # Retain a native two-zone masthead even when every provider box was a
    # page-wide estimate. Move whole blocks only; never invent text fragments.
    if not left_entries and len(right_entries) >= 2:
        left_entries.append(right_entries.pop(0))
    if not right_entries and len(left_entries) >= 2:
        candidate = max(
            range(1, len(left_entries)),
            key=lambda item: (
                left_entries[item][0].kind is MarkdownBlockKind.HEADING,
                len(left_entries[item][1]),
            ),
        )
        right_entries.append(left_entries.pop(candidate))
    if not left_entries or not right_entries:
        return None
    return (
        _sort_masthead_entries(left_entries, placements),
        _sort_masthead_entries(right_entries, placements),
    )


def _render_split_masthead(
    document: WordDocumentType,
    blocks: list[MarkdownBlock],
    placements: dict[str, HybridBlockPlacement],
    page: Any,
    *,
    width: float,
    size: float,
    line_height: float,
) -> tuple[bool, list[MarkdownBlock]]:
    """Render a two-zone masthead above an independently single-column body."""

    if page.metadata.get("header_column_count") != 2 or page.metadata.get("column_count") != 1:
        return False, blocks
    section_index = next(
        (
            index
            for index, block in enumerate(blocks)
            if block.metadata.get("role") == "section_heading"
        ),
        None,
    )
    if section_index is None or section_index < 4:
        return False, blocks
    preamble = blocks[:section_index]
    divider = page.metadata.get("header_divider")
    if not isinstance(divider, (int, float)):
        return False, blocks
    render_frame = _page_render_content_bbox(page)
    left_fraction = (float(divider) - render_frame.x0) / max(1, render_frame.width)
    left_fraction = max(0.25, min(0.55, left_fraction))
    entries = _masthead_entries(
        preamble,
        placements,
        render_frame=render_frame,
        divider=float(divider),
    )
    if entries is None:
        return False, blocks
    left_entries, right_entries = entries

    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(table, visible=False)
    caption = OxmlElement("w:tblCaption")
    caption.set(qn("w:val"), "docreconstruct:split-masthead")
    table._tbl.tblPr.append(caption)
    widths = [width * left_fraction, width * (1.0 - left_fraction)]
    _set_table_widths(table, widths)
    row_properties = table.rows[0]._tr.get_or_add_trPr()
    row_properties.append(OxmlElement("w:cantSplit"))
    for column_index, (column_entries, side) in enumerate(
        ((left_entries, "left"), (right_entries, "right"))
    ):
        _render_masthead_cell(
            table.cell(0, column_index),
            column_entries,
            side=side,
            width_inches=widths[column_index],
            size=size,
            line_height=line_height,
            placements=placements,
            source_x0=(float(render_frame.x0) if side == "left" else float(divider)),
            source_x1=(float(divider) if side == "left" else float(render_frame.x1)),
        )
    return True, blocks[section_index:]


def _partition_source_footer(
    blocks: list[MarkdownBlock],
    placements: dict[str, HybridBlockPlacement],
    page: Any,
) -> tuple[list[MarkdownBlock], list[MarkdownBlock]]:
    """Move short trailing bottom furniture out of the editable body flow."""

    threshold = page.content_bbox.y0 + page.content_bbox.height * 0.95
    split = len(blocks)
    while split > 0:
        block = blocks[split - 1]
        placement = placements.get(block.id)
        box = placement.source_bbox if placement is not None else None
        if (
            box is None
            or box.y0 < threshold
            or block.kind not in {MarkdownBlockKind.PARAGRAPH, MarkdownBlockKind.HEADING}
            or len(block.text) > 120
        ):
            break
        split -= 1
    footer_ids = {block.id for block in blocks[split:]}

    # Saved OCR can omit or weaken the bbox for page furniture, especially
    # when a repeated top banner follows it in Markdown reading order.  A
    # short multilingual label containing the current n/m page fraction is a
    # safe fallback near the tail; ordinary answer fractions are excluded by
    # kind, current-page validation, and the compact label envelope.
    fraction_pattern = re.compile(r"(?<!\d)(\d{1,4})\s*/\s*(\d{1,4})(?!\d)")
    search_start = max(0, len(blocks) - 3)
    for index in range(len(blocks) - 1, search_start - 1, -1):
        block = blocks[index]
        if (
            block.kind not in {MarkdownBlockKind.PARAGRAPH, MarkdownBlockKind.HEADING}
            or len(block.text) > 80
            or "$" in block.text
        ):
            continue
        match = fraction_pattern.search(block.text)
        if match is None:
            continue
        current, total = (int(value) for value in match.groups())
        if current != page.number or not 1 <= current <= total:
            continue
        prefix = block.text[: match.start()].strip(" \t:：—–-()[]")
        suffix = block.text[match.end() :].strip(" \t:：—–-()[]")
        compact_label = bool(prefix) and len(prefix) <= 32 and len(prefix.split()) <= 4
        placement = placements.get(block.id)
        box = placement.source_bbox if placement is not None else None
        low_on_page = box is not None and box.y0 >= (
            page.content_bbox.y0 + page.content_bbox.height * 0.88
        )
        if not low_on_page and (
            not compact_label or len(suffix) > 24 or re.search(r"[?;=]", prefix + suffix)
        ):
            continue
        footer_ids.add(block.id)
        if index > 0:
            preceding = blocks[index - 1]
            if (
                preceding.kind in {MarkdownBlockKind.PARAGRAPH, MarkdownBlockKind.HEADING}
                and len(preceding.text) <= 160
                and re.search(r"(?:https?://|www\.)", preceding.text, flags=re.IGNORECASE)
            ):
                footer_ids.add(preceding.id)
        break

    return (
        [block for block in blocks if block.id not in footer_ids],
        [block for block in blocks if block.id in footer_ids],
    )


def _render_source_footer(
    section: Any,
    blocks: Sequence[MarkdownBlock],
    placements: dict[str, HybridBlockPlacement],
    page: Any,
    *,
    size: float,
    line_height: float,
    clear_inherited: bool = False,
) -> None:
    if not blocks and not clear_inherited:
        return
    footer = section.footer
    footer.is_linked_to_previous = False
    paragraph = footer.paragraphs[0]
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)
    if not blocks:
        return
    rendered_size = max(8.0, size * 0.82)
    _format_paragraph(
        paragraph,
        size=rendered_size,
        line_height=max(rendered_size * 1.12, line_height * 0.82),
        alignment=WD_ALIGN_PARAGRAPH.RIGHT,
        before=0,
        after=0,
    )
    for index, block in enumerate(blocks):
        if index:
            paragraph.add_run().add_break()
        _add_rich_text(paragraph, block.text, size=rendered_size)
    boxes = [
        placements[block.id].source_bbox
        for block in blocks
        if block.id in placements and placements[block.id].source_bbox is not None
    ]
    if boxes:
        render_frame = _page_render_content_bbox(page)
        centers = [(cast(PixelBox, box).x0 + cast(PixelBox, box).x1) / 2.0 for box in boxes]
        center_fraction = (statistics.median(centers) - render_frame.x0) / max(
            1.0, render_frame.width
        )
        paragraph.alignment = (
            WD_ALIGN_PARAGRAPH.LEFT
            if center_fraction < 0.34
            else WD_ALIGN_PARAGRAPH.RIGHT
            if center_fraction > 0.66
            else WD_ALIGN_PARAGRAPH.CENTER
        )
        vertical_scale = page.pdf_height / page.height
        bottom_gap = (page.height - max(cast(PixelBox, box).y1 for box in boxes)) * vertical_scale
        section.footer_distance = Pt(max(4.0, min(24.0, bottom_gap)))


def render_hybrid_docx(
    content: MarkdownContent,
    layout: ScanDocumentLayout,
    plan: HybridLayoutPlan,
    asset_matches: list[AssetMatch],
    *,
    asset_payloads: Mapping[str, bytes] | None = None,
    render_input_sha256: str | None = None,
) -> bytes:
    """Render a hybrid plan as native Word paragraphs/tables plus real figures."""

    document = WordDocument()
    _set_native_page_background(document, _document_paper_color(layout))
    document.core_properties.title = Path(content.source).stem
    document.core_properties.subject = (
        "Editable reconstruction from Markdown content and scan layout"
    )
    document.core_properties.author = "docreconstruct"
    if render_input_sha256 is not None:
        if re.fullmatch(r"[0-9a-f]{64}", render_input_sha256) is None:
            raise ValueError("render_input_sha256 must be a lowercase SHA-256 digest")
        document.core_properties.identifier = (
            f"docreconstruct-render-input-sha256:{render_input_sha256}"
        )
    document.core_properties.comments = (
        "Paragraphs and tables are native Word objects. "
        "Raster objects are limited to source figures and unmatched non-text source fragments."
    )
    settings = document.settings.element
    compatibility = settings.find(qn("w:compat"))
    if compatibility is None:
        compatibility = OxmlElement("w:compat")
        settings.append(compatibility)
    if compatibility.find(qn("w:doNotExpandShiftReturn")) is None:
        compatibility.append(OxmlElement("w:doNotExpandShiftReturn"))
    block_by_id = {block.id: block for block in content.blocks}
    placement_by_id = {
        placement.block_id: placement for page in plan.pages for placement in page.placements
    }
    asset_match_by_id = {match.block_id: match for match in asset_matches}
    prepared_asset_mode = asset_payloads is not None
    asset_bytes: dict[str, bytes] = dict(asset_payloads or {})
    duplicate_figure_annotations = _duplicate_figure_annotation_ids(
        content.blocks,
        figures_rendered=_places_figure_pixels(content, layout),
    )
    markdown_directory = Path(content.source).parent
    for block in () if prepared_asset_mode else content.image_blocks:
        if block.id in asset_bytes:
            continue
        match = asset_match_by_id.get(block.id)
        if match is None or not match.resolved:
            continue
        try:
            resolved = resolve_markdown_asset(block, markdown_directory=markdown_directory)
        except (OSError, TimeoutError, ValueError):
            resolved = None
        if resolved:
            asset_bytes[block.id] = resolved.data

    horizontal_scales = [page.pdf_width / page.width for page in layout.pages]
    vertical_scales = [
        page.pdf_height / page.height
        if page.metadata.get("source_kind") == "image"
        else horizontal_scale
        for page, horizontal_scale in zip(layout.pages, horizontal_scales, strict=True)
    ]
    render_frames = [_page_render_content_bbox(page) for page in layout.pages]
    left_points = [
        frame.x0 * scale for frame, scale in zip(render_frames, horizontal_scales, strict=True)
    ]
    right_points = [
        (page.width - frame.x1) * scale
        for page, frame, scale in zip(
            layout.pages,
            render_frames,
            horizontal_scales,
            strict=True,
        )
    ]
    top_points = []
    bottom_points = []
    for page, horizontal_scale, vertical_scale in zip(
        layout.pages,
        horizontal_scales,
        vertical_scales,
        strict=True,
    ):
        offset = (
            0.0
            if page.metadata.get("source_kind") == "image"
            else max(0.0, (page.pdf_height - page.height * horizontal_scale) / 2)
        )
        top_points.append(offset + page.content_bbox.y0 * vertical_scale)
        bottom_points.append(page.pdf_height - (offset + page.content_bbox.y1 * vertical_scale))
    left_margin = statistics.median(left_points)
    right_margin = statistics.median(right_points)
    top_margin = statistics.median(top_points)
    # Ignore a final intentionally sparse page when deriving document margins.
    dense_bottoms = sorted(bottom_points)[: max(1, math.ceil(len(bottom_points) * 0.8))]
    bottom_margin = statistics.median(dense_bottoms)
    pitch_points = statistics.median(
        page.line_pitch * scale for page, scale in zip(layout.pages, vertical_scales, strict=True)
    )
    # Baseline pitch includes leading; native font size is normally about
    # three-quarters of that pitch.  Treating the two as nearly equal made
    # photographed exam sheets wrap and expand far beyond their source rows.
    body_size = max(8.6, min(12.0, pitch_points * 0.76))
    line_height = max(body_size + 1.2, pitch_points)

    effective_footer_has_content = False
    for page_index, page_plan in enumerate(plan.pages):
        section = (
            document.sections[0] if page_index == 0 else document.add_section(WD_SECTION.NEW_PAGE)
        )
        source_page = layout.pages[page_index]
        page_top_margin = top_margin
        first_block_index = min(
            (placement.block_index for placement in page_plan.placements),
            default=0,
        )
        leading_image_tops = [
            placement.source_bbox.y0
            for placement in page_plan.placements
            if placement.source_bbox is not None
            and placement.block_index <= first_block_index + 2
            and block_by_id[placement.block_id].kind is MarkdownBlockKind.IMAGE
        ]
        if leading_image_tops:
            source_offset = (
                0.0
                if source_page.metadata.get("source_kind") == "image"
                else max(
                    0.0,
                    (source_page.pdf_height - source_page.height * horizontal_scales[page_index])
                    / 2,
                )
            )
            page_top_margin = min(
                page_top_margin,
                source_offset + min(leading_image_tops) * vertical_scales[page_index],
            )
        section.page_width = Pt(page_plan.pdf_width)
        section.page_height = Pt(page_plan.pdf_height)
        section.left_margin = Pt(left_margin)
        section.right_margin = Pt(right_margin)
        section.top_margin = Pt(page_top_margin)
        # Let a source page that demonstrably uses more of the paper retain
        # that space, while avoiding a sparse final page inflating every page.
        page_bottom_margin = min(bottom_margin, bottom_points[page_index])
        section.bottom_margin = Pt(page_bottom_margin)
        section.header_distance = Pt(12)
        section.footer_distance = Pt(12)
        available_width = (page_plan.pdf_width - left_margin - right_margin) / 72.0
        if (
            page_index == 0
            and source_page.metadata.get("source_kind") == "pdf"
            and source_page.metadata.get("header_column_count") != 2
        ):
            header = _header_crop(source_page)
            if header:
                _add_picture(
                    document,
                    _crop_bytes(layout, 1, header),
                    width_inches=available_width,
                )
        blocks = [
            block_by_id[placement.block_id]
            for placement in page_plan.placements
            if placement.block_id not in duplicate_figure_annotations
        ]
        blocks, footer_blocks = _partition_source_footer(
            blocks,
            placement_by_id,
            source_page,
        )
        body_ids = {block.id for block in blocks}
        body_placements = [
            placement for placement in page_plan.placements if placement.block_id in body_ids
        ]
        budget_options: dict[str, Any] = {}
        if any(
            placement.source_bbox is None or placement.source_gap_before is None
            for placement in body_placements
        ):
            budget_options["blocks"] = blocks
        if source_page.text_lines:
            budget_options.update({"blocks": blocks, "line_height_points": line_height})
        vertical_budget = build_page_vertical_fit_budget(
            source_page,
            body_placements,
            printable_height_points=(page_plan.pdf_height - page_top_margin - page_bottom_margin),
            font_size_points=body_size,
            **budget_options,
        )
        page_body_size = body_size * vertical_budget.font_size_scale
        page_line_height = line_height * vertical_budget.line_height_scale
        fitted_placements = apply_page_vertical_fit_budget(
            source_page,
            body_placements,
            vertical_budget,
        )
        body_placement_by_id = dict(placement_by_id)
        body_placement_by_id.update(
            {placement.block_id: placement for placement in fitted_placements}
        )
        _render_source_footer(
            section,
            footer_blocks,
            placement_by_id,
            source_page,
            size=page_body_size,
            line_height=page_line_height,
            clear_inherited=effective_footer_has_content and not footer_blocks,
        )
        effective_footer_has_content = bool(footer_blocks)
        masthead_rendered, blocks = _render_split_masthead(
            document,
            blocks,
            body_placement_by_id,
            source_page,
            width=available_width,
            size=page_body_size,
            line_height=page_line_height,
        )
        if masthead_rendered:
            _add_vertical_spacer(document, page_line_height * 0.05)
        rendered_columns = False
        column_count = source_page.metadata.get("column_count")
        if column_count == 2:
            rendered_columns = _render_two_column_page(
                document,
                blocks,
                body_placement_by_id,
                width=available_width,
                size=page_body_size,
                line_height=page_line_height,
                asset_bytes=asset_bytes,
                layout=layout,
                page_number=page_index + 1,
            )
        if not rendered_columns and column_count in {2, 3, 4}:
            rendered_columns = _render_multi_column_page(
                document,
                blocks,
                body_placement_by_id,
                width=available_width,
                size=page_body_size,
                line_height=page_line_height,
                asset_bytes=asset_bytes,
                layout=layout,
                page_number=page_index + 1,
            )
        if not rendered_columns:
            # Preserve source order while grouping semantically related blocks.
            cursor = 0
            while cursor < len(blocks):
                group_id = blocks[cursor].group_id
                if group_id is None:
                    end = cursor + 1
                else:
                    end = cursor + 1
                    while end < len(blocks) and blocks[end].group_id == group_id:
                        end += 1
                _render_group(
                    document,
                    blocks[cursor:end],
                    body_placement_by_id,
                    width=available_width,
                    size=page_body_size,
                    line_height=page_line_height,
                    asset_bytes=asset_bytes,
                    layout=layout,
                )
                cursor = end
        if blocks and blocks[-1].kind is MarkdownBlockKind.HEADING:
            for paragraph in reversed(document.paragraphs):
                if paragraph.text == blocks[-1].text:
                    paragraph.paragraph_format.keep_with_next = False
                    break
    _compact_empty_paragraphs_in_cells(document)
    for paragraph in document.paragraphs:
        if not paragraph.text and not paragraph._p.xpath(".//w:drawing | .//m:oMath"):
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            paragraph.paragraph_format.line_spacing = Pt(1)
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


__all__ = ["render_hybrid_docx"]
